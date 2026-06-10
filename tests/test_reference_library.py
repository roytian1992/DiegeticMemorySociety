from __future__ import annotations

import json
from pathlib import Path

from dms.cli import main
from dms.llm import FakeReferenceFactPropertyClient, FakeReferenceKGClient
from dms.llm.client import LLMResult
from dms.reference_library import (
    ChromaReferenceIndexConfig,
    ReferenceFactPropertyExtractionConfig,
    ReferenceKGExtractionConfig,
    ReferenceKnowledgeImportConfig,
    ReferenceKnowledgeQuery,
    ReferenceLibraryIngestConfig,
    build_chroma_reference_index,
    extract_reference_facts_properties,
    extract_reference_kg,
    get_reference_asset_counts,
    ingest_reference_library,
    import_reference_knowledge,
    query_reference_knowledge,
)
from tests.helpers import write_jsonl


def _read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def test_ingest_reference_library_chunks_mixed_formats(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "world.md").write_text(
        "# 世界观\n550A位于数字生命研究室。\n\n# 风格\n对白要短促，信息密度高。\n",
        encoding="utf-8",
    )
    (refs_dir / "profiles.txt").write_text("刘培强：年轻飞行员。\n张鹏：教官。", encoding="utf-8")
    (refs_dir / "timeline.json").write_text(
        json.dumps(
            {
                "documents": [
                    {
                        "title": "时间线",
                        "chunks": [
                            {"heading": "2044", "content": "2044 年发生太空电梯危机。"},
                            {"heading": "地点", "content": "加蓬基地负责重要发射任务。"},
                        ],
                    }
                ]
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    write_jsonl(refs_dir / "notes.jsonl", [{"title": "笔记", "content": "作者笔记：不要把设定讲成说明书。"}])
    output_dir = tmp_path / "library"

    summary = ingest_reference_library(
        ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=output_dir, max_chunk_chars=80)
    )

    raw_docs = _read_jsonl(output_dir / "raw_documents.jsonl")
    chunks = _read_jsonl(output_dir / "reference_chunks.jsonl")
    assert summary["file_count"] == 4
    assert summary["raw_document_count"] == 5
    assert len(raw_docs) == 5
    assert summary["reference_chunk_count"] == len(chunks)
    assert {chunk["format"] for chunk in chunks} >= {"md", "txt", "json", "jsonl"}
    assert any(chunk["heading"] == "世界观" and "550A" in chunk["content"] for chunk in chunks)


def test_ingest_reference_library_keeps_long_boundaryless_text_intact(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    long_text = "A" * 220
    (refs_dir / "long.txt").write_text(long_text, encoding="utf-8")
    output_dir = tmp_path / "library"

    summary = ingest_reference_library(
        ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=output_dir, max_chunk_chars=40)
    )

    chunks = _read_jsonl(output_dir / "reference_chunks.jsonl")
    assert summary["reference_chunk_count"] == 1
    assert chunks[0]["content"] == long_text


def test_ingest_reference_library_chunks_docx_headings(tmp_path: Path) -> None:
    from docx import Document

    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    doc = Document()
    doc.add_heading("人物设定", level=1)
    doc.add_heading("韩临", level=2)
    doc.add_paragraph("本名：韩临", style="List Bullet")
    doc.add_paragraph("出生：2004 年", style="List Bullet")
    doc.add_heading("时间线", level=2)
    doc.add_paragraph("2033 年，雪去世。")
    doc.save(refs_dir / "profiles.docx")
    output_dir = tmp_path / "library"

    summary = ingest_reference_library(
        ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=output_dir, max_chunk_chars=120)
    )

    raw_docs = _read_jsonl(output_dir / "raw_documents.jsonl")
    chunks = _read_jsonl(output_dir / "reference_chunks.jsonl")
    assert summary["file_count"] == 1
    assert raw_docs[0]["format"] == "docx"
    assert raw_docs[0]["raw"]["parser"] == "python-docx"
    assert any(chunk["heading"] == "人物设定 / 韩临" and "- 本名：韩临" in chunk["content"] for chunk in chunks)
    assert any(chunk["heading"] == "人物设定 / 时间线" and "2033 年" in chunk["content"] for chunk in chunks)


def test_cli_reference_ingest_and_extract_kg_fake(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "style.md").write_text("# 人物\n刘培强第一次接触550A训练。\n", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"

    ingest_code = main(
        [
            "ingest-reference-library",
            str(refs_dir),
            "--output-dir",
            str(library_dir),
            "--max-chunk-chars",
            "120",
        ]
    )
    extract_code = main(
        [
            "extract-reference-kg",
            str(library_dir),
            "--output-dir",
            str(kg_dir),
            "--no-dry-run",
            "--provider",
            "fake",
        ]
    )

    assert ingest_code == 0
    assert extract_code == 0
    entities = _read_jsonl(kg_dir / "entities.jsonl")
    assert {entity["entity_name"] for entity in entities} >= {"刘培强", "550A"}


def test_lightrag_style_reference_knowledge_assets_and_source_filter(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text(
        "# 人物\n刘培强第一次接触550A训练。张鹏指导刘培强保持稳定。\n",
        encoding="utf-8",
    )
    (refs_dir / "timeline.md").write_text(
        "# 时间线\n2044 年爆发太空电梯危机。\n",
        encoding="utf-8",
    )
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    kg_summary = extract_reference_kg(
        ReferenceKGExtractionConfig(library_dir=library_dir, output_dir=kg_dir, dry_run=False),
        llm_client=FakeReferenceKGClient(),
    )
    import_summary = import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, db_path=db_path, reset=True)
    )
    counts = get_reference_asset_counts(db_path)
    all_result = query_reference_knowledge(
        ReferenceKnowledgeQuery(db_path=db_path, query="刘培强 550A 张鹏", top_k=8)
    )
    timeline_doc_id = next(
        row["doc_id"]
        for row in _read_jsonl(library_dir / "raw_documents.jsonl")
        if row["file_name"] == "timeline.md"
    )
    profiles_doc_id = next(
        row["doc_id"]
        for row in _read_jsonl(library_dir / "raw_documents.jsonl")
        if row["file_name"] == "profiles.md"
    )
    filtered_result = query_reference_knowledge(
        ReferenceKnowledgeQuery(
            db_path=db_path,
            query="刘培强 550A 张鹏",
            source_doc_ids=(timeline_doc_id,),
            top_k=8,
        )
    )

    assert kg_summary["entity_extraction_count"] >= 4
    assert import_summary["asset_model"] == "source_local_external_reference_v1"
    assert counts["reference_full_docs"] == 2
    assert counts["reference_source_scopes"] == 2
    assert counts["reference_llm_response_cache"] == 2
    assert counts["reference_text_chunks"] == 2
    assert counts["reference_source_local_entities"] >= 4
    assert counts["reference_source_local_relationships"] >= 2
    assert counts["reference_entity_clusters"] >= 4
    assert counts["reference_atomic_facts"] >= counts["reference_source_local_entities"]
    assert counts["reference_vector_documents"] >= counts["reference_text_chunks"] + counts["reference_source_local_entities"]
    assert all_result["asset_model"] == "source_local_external_reference_v1"
    assert all_result["mode"] == "source_local"
    assert all_result["evidence_board"]["source_local_entities"]
    assert all_result["atomic_facts"]
    assert all_result["relationship_facts"]
    assert all_result["evidence_board"]["relationship_facts"] == all_result["relationship_facts"]
    assert all(prop["property_name"] != "entity_type" for prop in all_result["entity_properties"])
    assert {entity["entity_name"] for entity in all_result["entities"]} >= {"刘培强", "550A"}
    assert [entity["entity_name"] for entity in all_result["entities"]].count("刘培强") == 1
    relationship_by_id = {item["source_local_relation_id"]: item for item in all_result["relationships"]}
    relationship_fact = next(
        item
        for item in all_result["relationship_facts"]
        if item["source_local_relation_id"] in relationship_by_id
    )
    relationship = relationship_by_id[relationship_fact["source_local_relation_id"]]
    assert relationship["evidence_chunks"]
    assert relationship_fact["evidence_chunks"]
    assert relationship_fact["source_local_relation_id"] == relationship["source_local_relation_id"]
    assert "刘培强" in relationship_fact["statement"]
    assert "刘培强" in relationship_fact["evidence_chunks"][0]["content"]

    assert filtered_result["source_filter"]["mode"] == "files"
    assert all(timeline_doc_id in (item.get("source_doc_ids") or [item.get("full_doc_id")]) for item in filtered_result["entities"] + filtered_result["relationships"] + filtered_result["chunks"])

    profile_filtered_result = query_reference_knowledge(
        ReferenceKnowledgeQuery(
            db_path=db_path,
            query="刘培强",
            source_doc_ids=(profiles_doc_id,),
            top_k=8,
        )
    )
    assert "刘培强" in {entity["entity_name"] for entity in profile_filtered_result["entities"]}
    assert all(profiles_doc_id in (item.get("source_doc_ids") or [item.get("full_doc_id")]) for item in profile_filtered_result["entities"] + profile_filtered_result["relationships"] + profile_filtered_result["chunks"])

    chroma_dir = tmp_path / "reference_chroma"
    index_summary = build_chroma_reference_index(
        ChromaReferenceIndexConfig(
            db_path=db_path,
            persist_dir=chroma_dir,
            reset=True,
            embedding_provider="hash",
        )
    )
    chroma_filtered_result = query_reference_knowledge(
        ReferenceKnowledgeQuery(
            db_path=db_path,
            query="刘培强",
            source_paths=("profiles.md",),
            chroma_dir=chroma_dir,
            top_k=8,
            embedding_provider="hash",
        )
    )
    assert index_summary["document_count"] == counts["reference_vector_documents"]
    assert "刘培强" in {entity["entity_name"] for entity in chroma_filtered_result["entities"]}
    assert all(profiles_doc_id in (item.get("source_doc_ids") or [item.get("full_doc_id")]) for item in chroma_filtered_result["entities"] + chroma_filtered_result["relationships"] + chroma_filtered_result["chunks"])


def test_reference_facts_properties_extraction_imports_structured_properties(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("# 人物\n刘培强第一次接触550A训练。张鹏指导刘培强保持稳定。\n", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    facts_dir = tmp_path / "facts"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    extract_reference_kg(
        ReferenceKGExtractionConfig(library_dir=library_dir, output_dir=kg_dir, dry_run=False),
        llm_client=FakeReferenceKGClient(),
    )
    facts_summary = extract_reference_facts_properties(
        ReferenceFactPropertyExtractionConfig(library_dir=library_dir, kg_dir=kg_dir, output_dir=facts_dir, dry_run=False),
        llm_client=FakeReferenceFactPropertyClient(),
    )
    import_summary = import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, facts_dir=facts_dir, db_path=db_path, reset=True)
    )
    result = query_reference_knowledge(ReferenceKnowledgeQuery(db_path=db_path, query="刘培强 profile_note", top_k=8))

    assert facts_summary["atomic_fact_count"] > 0
    assert facts_summary["entity_property_count"] > 0
    assert facts_summary["selection"]["min_entity_degree"] == 2
    assert facts_summary["selection"]["total_jobs"] >= facts_summary["eligible_entity_cluster_count"]
    assert import_summary["extracted_atomic_facts"] == facts_summary["atomic_fact_count"]
    assert import_summary["extracted_entity_properties"] == facts_summary["entity_property_count"]
    assert get_reference_asset_counts(db_path)["reference_entity_properties"] > 0
    assert any(prop["entity_name"] == "刘培强" for prop in result["entity_properties"])
    assert any(fact["fact_type"].endswith("_atomic_fact") for fact in result["atomic_facts"] + result["relationship_facts"])


def test_reference_query_can_disable_and_bind_fact_properties_per_entity(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text(
        "# 人物\n刘培强第一次接触550A训练。刘培强在模拟训练中保持情绪稳定。张鹏指导刘培强。\n",
        encoding="utf-8",
    )
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    facts_dir = tmp_path / "facts"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    chunk = chunks[0]
    kg_dir.mkdir()
    write_jsonl(
        kg_dir / "entities.jsonl",
        [
            {
                "extraction_id": "ent_lpq",
                "entity_name": "刘培强",
                "entity_type": "character",
                "description": "刘培强是参与训练的角色。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
            {
                "extraction_id": "ent_zp",
                "entity_name": "张鹏",
                "entity_type": "character",
                "description": "张鹏指导刘培强。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
        ],
    )
    write_jsonl(
        kg_dir / "relationships.jsonl",
        [
            {
                "extraction_id": "rel_lpq_zp",
                "src_id": "刘培强",
                "tgt_id": "张鹏",
                "keywords": "训练指导",
                "description": "张鹏指导刘培强进行训练。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
                "relation_id": "rel_lpq_zp",
            }
        ],
    )
    write_jsonl(kg_dir / "trace.jsonl", [{"chunk_id": chunk["chunk_id"], "doc_id": chunk["doc_id"], "status": "manual"}])

    source_local_entity_id = __import__(
        "dms.reference_library", fromlist=["_reference_source_local_entity_id"]
    )._reference_source_local_entity_id(chunk["source_scope_id"], "刘培强")
    cluster_id = __import__(
        "dms.reference_library", fromlist=["_reference_entity_cluster_id"]
    )._reference_entity_cluster_id("刘培强", "character")
    facts_dir.mkdir()
    fact_statements = [
        "刘培强在模拟训练中保持情绪稳定。",
        "刘培强面对张鹏指导时能稳定执行动作。",
        "刘培强第一次接触550A训练。",
        "刘培强记录训练日志。",
        "刘培强听取基地广播。",
    ]
    write_jsonl(
        facts_dir / "atomic_facts.jsonl",
        [
            {
                "fact_id": f"manual_fact_{index}",
                "source_scope_id": chunk["source_scope_id"],
                "source_local_entity_id": source_local_entity_id,
                "source_local_relation_id": "",
                "cluster_id": cluster_id,
                "subject": "刘培强",
                "predicate": "states",
                "object": "",
                "statement": statement,
                "fact_type": "entity_atomic_fact",
                "confidence": 0.8,
                "source_chunk_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "evidence": statement,
            }
            for index, statement in enumerate(fact_statements, start=1)
        ],
    )
    write_jsonl(
        facts_dir / "entity_properties.jsonl",
        [
            {
                "property_id": "manual_prop_1",
                "source_scope_id": chunk["source_scope_id"],
                "source_local_entity_id": source_local_entity_id,
                "cluster_id": cluster_id,
                "entity_name": "刘培强",
                "property_name": "训练状态",
                "property_value": "情绪稳定",
                "statement": "刘培强.训练状态: 情绪稳定",
                "confidence": 0.8,
                "source_chunk_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "evidence": "刘培强在模拟训练中保持情绪稳定。",
            }
        ],
    )
    import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, facts_dir=facts_dir, db_path=db_path, reset=True)
    )

    disabled = query_reference_knowledge(
        ReferenceKnowledgeQuery(db_path=db_path, query="刘培强 情绪稳定", top_k=8, include_fact_properties=False)
    )
    bound = query_reference_knowledge(
        ReferenceKnowledgeQuery(db_path=db_path, query="刘培强 情绪稳定", top_k=8, entity_top_k=1, fact_binding_top_k=2)
    )

    assert disabled["atomic_facts"] == []
    assert disabled["entity_properties"] == []
    bound_lpq_facts = [
        fact for fact in bound["atomic_facts"] if fact.get("source_local_entity_id") == source_local_entity_id
    ]
    assert 1 <= len(bound_lpq_facts) <= 2
    assert any("情绪稳定" in fact["statement"] for fact in bound_lpq_facts)
    assert all((fact.get("binding") or {}).get("strategy") == "entity_fact_top_k_by_query_similarity" for fact in bound_lpq_facts)
    assert any(prop["property_name"] == "训练状态" for prop in bound["entity_properties"])


def test_reference_facts_properties_batches_large_entity_cluster(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text(
        "\n\n".join(
            [
                "## Part 1\n刘培强第一次接触550A训练。张鹏指导刘培强。",
                "## Part 2\n刘培强在模拟训练中保持稳定。550A记录训练状态。",
                "## Part 3\n张鹏提醒刘培强不要被情绪带偏。550A参与训练评估。",
            ]
        ),
        encoding="utf-8",
    )
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    facts_dir = tmp_path / "facts"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir, max_chunk_chars=80))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    kg_dir.mkdir()
    write_jsonl(
        kg_dir / "entities.jsonl",
        [
            {
                "extraction_id": f"ent_{index}",
                "entity_name": "刘培强",
                "entity_type": "character",
                "description": f"刘培强在第 {index} 个片段中参与训练。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
            }
            for index, chunk in enumerate(chunks, start=1)
        ]
        + [
            {
                "extraction_id": f"ent_550a_{index}",
                "entity_name": "550A",
                "entity_type": "object",
                "description": f"550A在第 {index} 个片段中参与训练。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
            }
            for index, chunk in enumerate(chunks, start=1)
        ]
        + [
            {
                "extraction_id": f"ent_zp_{index}",
                "entity_name": "张鹏",
                "entity_type": "character",
                "description": f"张鹏在第 {index} 个片段中参与训练指导。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
            }
            for index, chunk in enumerate(chunks, start=1)
        ],
    )
    write_jsonl(
        kg_dir / "relationships.jsonl",
        [
            {
                "extraction_id": f"rel_{index}",
                "src_id": "刘培强",
                "tgt_id": "550A",
                "keywords": "训练",
                "description": f"刘培强与550A在第 {index} 个片段中发生训练关系。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
                "relation_id": "rel_lpq_550a",
            }
            for index, chunk in enumerate(chunks, start=1)
        ]
        + [
            {
                "extraction_id": f"rel_lpq_zp_{index}",
                "src_id": "刘培强",
                "tgt_id": "张鹏",
                "keywords": "指导",
                "description": f"张鹏在第 {index} 个片段中指导刘培强。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
                "relation_id": "rel_lpq_zp",
            }
            for index, chunk in enumerate(chunks, start=1)
        ]
        + [
            {
                "extraction_id": f"rel_zp_550a_{index}",
                "src_id": "张鹏",
                "tgt_id": "550A",
                "keywords": "训练评估",
                "description": f"张鹏在第 {index} 个片段中围绕550A训练评估进行指导。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": index,
                "relation_id": "rel_zp_550a",
            }
            for index, chunk in enumerate(chunks, start=1)
        ],
    )
    write_jsonl(kg_dir / "trace.jsonl", [{"chunk_id": chunk["chunk_id"], "doc_id": chunk["doc_id"], "status": "manual"} for chunk in chunks])

    facts_summary = extract_reference_facts_properties(
        ReferenceFactPropertyExtractionConfig(
            library_dir=library_dir,
            kg_dir=kg_dir,
            output_dir=facts_dir,
            dry_run=True,
            max_evidence_chunks_per_job=1,
        ),
        llm_client=None,
    )
    trace = _read_jsonl(facts_dir / "trace.jsonl")

    assert facts_summary["eligible_entity_cluster_count"] == 3
    assert facts_summary["batched_fact_property_job_count"] == facts_summary["selection"]["total_jobs"]
    assert facts_summary["selection"]["total_jobs"] > facts_summary["eligible_entity_cluster_count"]
    assert {row["evidence_chunk_count"] for row in trace} == {1}
    assert all(row["batch_count"] >= 2 for row in trace)


def test_reference_kg_extraction_retries_failed_chunk(tmp_path: Path) -> None:
    class FlakyReferenceKGClient:
        provider = "fake"
        model = "flaky-reference-kg"

        def __init__(self) -> None:
            self.calls = 0

        def complete(self, prompt: str) -> LLMResult:
            self.calls += 1
            if self.calls == 1:
                return LLMResult(
                    text="",
                    provider=self.provider,
                    model=self.model,
                    raw_response={"fake": True, "empty": True},
                    usage={},
                )
            text = "\n".join(
                [
                    "entity<|#|>刘培强<|#|>character<|#|>刘培强 appears in the external reference material.",
                    "entity<|#|>550A<|#|>object<|#|>550A appears in the external reference material.",
                    "relation<|#|>刘培强<|#|>550A<|#|>training<|#|>刘培强 trains with 550A.",
                    "<|COMPLETE|>",
                ]
            )
            return LLMResult(
                text=text,
                provider=self.provider,
                model=self.model,
                raw_response={"fake": True, "text": text},
                usage={},
            )

    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("刘培强第一次接触550A训练。", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    client = FlakyReferenceKGClient()
    summary = extract_reference_kg(
        ReferenceKGExtractionConfig(library_dir=library_dir, output_dir=kg_dir, dry_run=False, max_retries=1),
        llm_client=client,
    )
    trace = _read_jsonl(kg_dir / "trace.jsonl")

    assert client.calls == 2
    assert summary["failed_count"] == 0
    assert summary["parsed_output_count"] == 1
    assert trace[0]["attempt"] == 2


def test_reference_kg_parser_tolerates_partial_tuple_delimiter_typo(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("韩临与雪于2030年结婚。", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    chunk = chunks[0]
    kg_dir.mkdir()
    write_jsonl(
        kg_dir / "entities.jsonl",
        [
            {
                "extraction_id": "ent_han",
                "entity_name": "韩临",
                "entity_type": "character",
                "description": "韩临是人物资料中的角色。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
            {
                "extraction_id": "ent_xue",
                "entity_name": "雪",
                "entity_type": "character",
                "description": "雪是韩临的配偶。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
        ],
    )
    raw_text = "\n".join(
        [
            "entity<|#|>韩临<|#|>character<|#|>韩临是人物资料中的角色。",
            "entity<|#|>雪<|#|>character<|#|>雪是韩临的配偶。",
            "relation<|#|>韩临<|#|>雪<|#|>婚姻|#|>韩临与雪于2030年结婚。",
            "<|COMPLETE|>",
        ]
    )
    parsed_entities, parsed_relationships = __import__(
        "dms.reference_library", fromlist=["_parse_lightrag_kg_output"]
    )._parse_lightrag_kg_output(raw_text, chunk)
    write_jsonl(kg_dir / "relationships.jsonl", parsed_relationships)
    write_jsonl(kg_dir / "trace.jsonl", [{"chunk_id": chunk["chunk_id"], "doc_id": chunk["doc_id"], "status": "manual"}])

    import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, db_path=db_path, reset=True)
    )

    assert parsed_entities
    assert parsed_relationships
    assert get_reference_asset_counts(db_path)["reference_source_local_relationships"] == 1


def test_reference_kg_entity_type_audit_flags_multiple_and_unsupported_types(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("韩临加入回响科技。", encoding="utf-8")
    library_dir = tmp_path / "library"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunk = _read_jsonl(library_dir / "reference_chunks.jsonl")[0]
    raw_text = "\n".join(
        [
            "entity<|#|>韩临<|#|>character, organization<|#|>韩临是人物资料中的角色。",
            "entity<|#|>回响科技<|#|>company<|#|>回响科技是外部资料中的公司。",
            "relation<|#|>韩临<|#|>回响科技<|#|>任职<|#|>韩临加入回响科技。",
            "<|COMPLETE|>",
        ]
    )

    parsed_entities, parsed_relationships = __import__(
        "dms.reference_library", fromlist=["_parse_lightrag_kg_output"]
    )._parse_lightrag_kg_output(raw_text, chunk)
    audit = __import__(
        "dms.reference_library", fromlist=["_reference_entity_type_audit"]
    )._reference_entity_type_audit(parsed_entities)

    han = next(entity for entity in parsed_entities if entity["entity_name"] == "韩临")
    echo = next(entity for entity in parsed_entities if entity["entity_name"] == "回响科技")
    assert han["entity_type"] == "character"
    assert han["entity_type_candidates"] == ["character", "organization"]
    assert {issue["issue"] for issue in han["entity_type_issues"]} == {"multiple_entity_types"}
    assert echo["entity_type"] == "concept"
    assert {issue["issue"] for issue in echo["entity_type_issues"]} == {"unsupported_entity_type"}
    assert audit["multiple_entity_type_count"] == 1
    assert audit["unsupported_entity_type_count"] == 1
    assert parsed_relationships


def test_source_local_entities_preserve_entity_type_conflict_audit(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("回响科技既被描述为组织，也被误标为地点。", encoding="utf-8")
    library_dir = tmp_path / "library"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunk = _read_jsonl(library_dir / "reference_chunks.jsonl")[0]
    rows = [
        {
            "extraction_id": "ent_org",
            "entity_name": "回响科技",
            "entity_type": "organization",
            "original_entity_type": "organization",
            "description": "回响科技是组织。",
            "source_id": chunk["chunk_id"],
            "source_doc_id": chunk["doc_id"],
            "source_scope_id": chunk["source_scope_id"],
            "file_path": chunk["source_path"],
            "timestamp": 1,
        },
        {
            "extraction_id": "ent_loc",
            "entity_name": "回响科技",
            "entity_type": "location",
            "original_entity_type": "location",
            "description": "回响科技被误标为地点。",
            "source_id": chunk["chunk_id"],
            "source_doc_id": chunk["doc_id"],
            "source_scope_id": chunk["source_scope_id"],
            "file_path": chunk["source_path"],
            "timestamp": 2,
        },
    ]

    reference_library = __import__(
        "dms.reference_library", fromlist=["_resolve_reference_source_local_entity_types", "_source_local_entities"]
    )
    source_local_entities = reference_library._resolve_reference_source_local_entity_types(
        reference_library._source_local_entities(rows)
    )

    assert len(source_local_entities) == 1
    entity = source_local_entities[0]
    assert entity["entity_type"] == "organization"
    assert entity["entity_types"] == ["organization", "location"]
    assert entity["entity_type_conflict"] is True
    assert any(issue["issue"] == "source_local_entity_type_conflict" for issue in entity["entity_type_issues"])


def test_entity_type_resolution_prefers_specific_type_before_disambiguation(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("M5既是系统对象，也被泛称为概念。M5节点是地点。", encoding="utf-8")
    library_dir = tmp_path / "library"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunk = _read_jsonl(library_dir / "reference_chunks.jsonl")[0]
    rows = [
        {
            "extraction_id": "ent_m5_object",
            "entity_name": "M5",
            "entity_type": "object",
            "original_entity_type": "object",
            "description": "M5是一个系统对象。",
            "source_id": chunk["chunk_id"],
            "source_doc_id": chunk["doc_id"],
            "source_scope_id": chunk["source_scope_id"],
            "file_path": chunk["source_path"],
            "timestamp": 1,
        },
        {
            "extraction_id": "ent_m5_concept",
            "entity_name": "M5",
            "entity_type": "concept",
            "original_entity_type": "concept",
            "description": "M5也被泛称为概念。",
            "source_id": chunk["chunk_id"],
            "source_doc_id": chunk["doc_id"],
            "source_scope_id": chunk["source_scope_id"],
            "file_path": chunk["source_path"],
            "timestamp": 2,
        },
        {
            "extraction_id": "ent_m5_node",
            "entity_name": "M5节点",
            "entity_type": "location",
            "original_entity_type": "location",
            "description": "M5节点是地点。",
            "source_id": chunk["chunk_id"],
            "source_doc_id": chunk["doc_id"],
            "source_scope_id": chunk["source_scope_id"],
            "file_path": chunk["source_path"],
            "timestamp": 3,
        },
    ]
    reference_library = __import__(
        "dms.reference_library",
        fromlist=[
            "_canonicalize_reference_source_local_graph",
            "_resolve_reference_source_local_entity_types",
            "_source_local_entities",
        ],
    )
    source_local_entities = reference_library._resolve_reference_source_local_entity_types(
        reference_library._source_local_entities(rows)
    )
    canonicalized, _relations = reference_library._canonicalize_reference_source_local_graph(
        source_local_entities,
        [],
        lexical_threshold=0.5,
    )

    by_name = {entity["entity_name"]: entity for entity in canonicalized}
    assert by_name["M5"]["entity_type"] == "object"
    assert by_name["M5"]["entity_type_conflict"] is True
    assert by_name["M5"]["entity_type_resolution"]["reason"] == "resolved_type_conflict"
    assert by_name["M5"]["cluster_id"] != by_name["M5节点"]["cluster_id"]


def test_entity_disambiguation_runs_within_resolved_entity_type_buckets(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "characters.md").write_text("Sentinel Alpha is a character also called Alpha.", encoding="utf-8")
    (refs_dir / "orgs.md").write_text("Alpha is an organization in another source.", encoding="utf-8")
    library_dir = tmp_path / "library"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    character_chunk, organization_chunk = chunks[0], chunks[1]
    rows = [
        {
            "extraction_id": "ent_character_long",
            "entity_name": "Sentinel Alpha",
            "entity_type": "character",
            "original_entity_type": "character",
            "description": "Sentinel Alpha is a character.",
            "source_id": character_chunk["chunk_id"],
            "source_doc_id": character_chunk["doc_id"],
            "source_scope_id": character_chunk["source_scope_id"],
            "file_path": character_chunk["source_path"],
            "timestamp": 1,
        },
        {
            "extraction_id": "ent_character_alias",
            "entity_name": "Alpha",
            "entity_type": "character",
            "original_entity_type": "character",
            "description": "Alpha is an alias for Sentinel Alpha.",
            "source_id": character_chunk["chunk_id"],
            "source_doc_id": character_chunk["doc_id"],
            "source_scope_id": character_chunk["source_scope_id"],
            "file_path": character_chunk["source_path"],
            "timestamp": 2,
        },
        {
            "extraction_id": "ent_org_alpha",
            "entity_name": "Alpha",
            "entity_type": "organization",
            "original_entity_type": "organization",
            "description": "Alpha is an organization.",
            "source_id": organization_chunk["chunk_id"],
            "source_doc_id": organization_chunk["doc_id"],
            "source_scope_id": organization_chunk["source_scope_id"],
            "file_path": organization_chunk["source_path"],
            "timestamp": 1,
        },
    ]
    reference_library = __import__(
        "dms.reference_library",
        fromlist=[
            "_canonicalize_reference_source_local_graph",
            "_resolve_reference_source_local_entity_types",
            "_source_local_entities",
        ],
    )
    source_local_entities = reference_library._resolve_reference_source_local_entity_types(
        reference_library._source_local_entities(rows)
    )
    canonicalized, _relations = reference_library._canonicalize_reference_source_local_graph(
        source_local_entities,
        [],
        lexical_threshold=0.5,
    )

    by_scope_name = {
        (entity["source_scope_id"], entity["entity_name"]): entity
        for entity in canonicalized
    }
    character_long = by_scope_name[(character_chunk["source_scope_id"], "Sentinel Alpha")]
    character_alias = by_scope_name[(character_chunk["source_scope_id"], "Alpha")]
    organization_alpha = by_scope_name[(organization_chunk["source_scope_id"], "Alpha")]

    assert character_long["cluster_id"] == character_alias["cluster_id"]
    assert character_alias["canonical_entity_name"] == "Sentinel Alpha"
    assert organization_alpha["entity_type"] == "organization"
    assert organization_alpha["cluster_id"] != character_alias["cluster_id"]


def test_cli_lightrag_style_reference_knowledge_flow(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("# 人物\n刘培强第一次接触550A训练。张鹏指导刘培强。\n", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    facts_dir = tmp_path / "facts"
    db_path = tmp_path / "reference_kg.sqlite"

    assert main(["ingest-reference-library", str(refs_dir), "--output-dir", str(library_dir)]) == 0
    assert main(["extract-reference-kg", str(library_dir), "--output-dir", str(kg_dir), "--no-dry-run", "--provider", "fake"]) == 0
    assert main(["extract-reference-facts-properties", str(library_dir), "--kg-dir", str(kg_dir), "--output-dir", str(facts_dir), "--no-dry-run", "--provider", "fake"]) == 0
    assert main(["import-reference-knowledge", "--library-dir", str(library_dir), "--kg-dir", str(kg_dir), "--facts-dir", str(facts_dir), "--output-db", str(db_path), "--overwrite"]) == 0
    assert main(["query-reference-knowledge", str(db_path), "--query", "刘培强 550A 张鹏", "--top-k", "6"]) == 0

    counts = get_reference_asset_counts(db_path)
    assert counts["reference_full_docs"] == 1
    assert counts["reference_llm_response_cache"] == 1
    assert counts["reference_source_local_entities"] >= 3
    assert counts["reference_entity_properties"] >= 1


def test_import_reference_knowledge_tolerates_trace_without_artifact_paths(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("# 人物\n刘培强第一次接触550A训练。\n", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    chunk = chunks[0]
    kg_dir.mkdir()
    write_jsonl(
        kg_dir / "entities.jsonl",
        [
            {
                "extraction_id": "ent_1",
                "entity_name": "刘培强",
                "entity_type": "character",
                "description": "刘培强第一次接触550A训练。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            }
        ],
    )
    write_jsonl(kg_dir / "relationships.jsonl", [])
    write_jsonl(kg_dir / "trace.jsonl", [{"chunk_id": chunk["chunk_id"], "doc_id": chunk["doc_id"], "status": "manual_demo"}])

    summary = import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, db_path=db_path, reset=True)
    )

    assert summary["asset_model"] == "source_local_external_reference_v1"
    assert summary["llm_response_cache"] == 1
    assert get_reference_asset_counts(db_path)["reference_source_local_entities"] == 1


def test_import_reference_knowledge_filters_negative_relationships(tmp_path: Path) -> None:
    refs_dir = tmp_path / "refs"
    refs_dir.mkdir()
    (refs_dir / "profiles.md").write_text("# 人物\n韩临与雪结婚。文本没有提及雪与回响科技存在直接关系。\n", encoding="utf-8")
    library_dir = tmp_path / "library"
    kg_dir = tmp_path / "kg"
    db_path = tmp_path / "reference_kg.sqlite"

    ingest_reference_library(ReferenceLibraryIngestConfig(input_path=refs_dir, output_dir=library_dir))
    chunks = _read_jsonl(library_dir / "reference_chunks.jsonl")
    chunk = chunks[0]
    kg_dir.mkdir()
    write_jsonl(
        kg_dir / "entities.jsonl",
        [
            {
                "extraction_id": "ent_han",
                "entity_name": "韩临",
                "entity_type": "character",
                "description": "韩临是人物资料中的角色。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
            {
                "extraction_id": "ent_xue",
                "entity_name": "雪",
                "entity_type": "character",
                "description": "雪是韩临的配偶。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
            {
                "extraction_id": "ent_echo",
                "entity_name": "回响科技",
                "entity_type": "organization",
                "description": "回响科技是外部资料中的组织。",
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
        ],
    )
    write_jsonl(
        kg_dir / "relationships.jsonl",
        [
            {
                "relation_id": "rel_han_xue",
                "extraction_id": "rel_1",
                "src_id": "韩临",
                "tgt_id": "雪",
                "keywords": "婚姻, 配偶",
                "description": "韩临与雪结婚，二人是配偶关系。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
            {
                "relation_id": "rel_xue_echo",
                "extraction_id": "rel_2",
                "src_id": "雪",
                "tgt_id": "回响科技",
                "keywords": "无直接关系",
                "description": "文本中未提及雪与回响科技之间存在任何直接关系。",
                "weight": 1.0,
                "source_id": chunk["chunk_id"],
                "source_doc_id": chunk["doc_id"],
                "source_scope_id": chunk["source_scope_id"],
                "file_path": chunk["source_path"],
                "timestamp": 1,
            },
        ],
    )
    write_jsonl(kg_dir / "trace.jsonl", [{"chunk_id": chunk["chunk_id"], "doc_id": chunk["doc_id"], "status": "manual_demo"}])

    summary = import_reference_knowledge(
        ReferenceKnowledgeImportConfig(library_dir=library_dir, kg_dir=kg_dir, db_path=db_path, reset=True)
    )
    counts = get_reference_asset_counts(db_path)
    result = query_reference_knowledge(ReferenceKnowledgeQuery(db_path=db_path, query="韩临 雪 回响科技", top_k=8))

    assert summary["raw_extracted_relationships"] == 2
    assert summary["extracted_relationships"] == 1
    assert summary["filtered_negative_relationships"] == 1
    assert counts["reference_extracted_relationships"] == 1
    assert counts["reference_source_local_relationships"] == 1
    assert {(item["src_id"], item["tgt_id"]) for item in result["relationships"]} == {("雪", "韩临")}
