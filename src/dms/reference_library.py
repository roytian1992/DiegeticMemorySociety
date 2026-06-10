from __future__ import annotations

import hashlib
import json
import re
import shutil
import sqlite3
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from dms.entity_types import ALLOWED_ENTITY_TYPES, is_supported_entity_type_label, normalize_entity_type
from dms.llm import LLMClient
from dms.progress import print_progress
from dms.storage.chroma_index import build_embedding_function


REFERENCE_SCHEMA_VERSION = 4
GRAPH_FIELD_SEP = "<SEP>"
VECTOR_NAMESPACE_CHUNKS = "chunks"
VECTOR_NAMESPACE_SOURCE_LOCAL_ENTITIES = "source_local_entities"
VECTOR_NAMESPACE_SOURCE_LOCAL_RELATIONSHIPS = "source_local_relationships"
VECTOR_NAMESPACE_ENTITY_CLUSTERS = "entity_clusters"
VECTOR_NAMESPACE_PSEUDO_RELATIONSHIPS = "pseudo_relationships"
VECTOR_NAMESPACE_ATOMIC_FACTS = "atomic_facts"
VECTOR_NAMESPACE_ENTITY_PROPERTIES = "entity_properties"
DOC_STATUS_PROCESSED = "processed"
REFERENCE_ASSET_MODEL = "source_local_external_reference_v1"

SUPPORTED_REFERENCE_EXTENSIONS = {".docx", ".md", ".markdown", ".txt", ".json", ".jsonl"}
NEGATIVE_RELATIONSHIP_PATTERNS = (
    "未提及",
    "没有提及",
    "未说明",
    "没有说明",
    "无直接关系",
    "不存在直接关系",
    "没有直接关系",
    "无明确关系",
    "没有明确关系",
    "无关系",
    "not mentioned",
    "not state",
    "not stated",
    "no direct relation",
    "no direct relationship",
    "no relationship",
)


@dataclass(frozen=True)
class ReferenceLibraryIngestConfig:
    input_path: Path
    output_dir: Path
    max_chunk_chars: int = 2400
    overwrite: bool = False
    workers: int = 4


@dataclass(frozen=True)
class ReferenceKGExtractionConfig:
    library_dir: Path
    output_dir: Path
    start: int = 1
    limit: int | None = None
    dry_run: bool = True
    overwrite: bool = False
    max_retries: int = 1
    workers: int = 4
    entity_types: tuple[str, ...] = (
        "character",
        "group",
        "organization",
        "location",
        "object",
        "concept",
        "occasion",
    )


@dataclass(frozen=True)
class ReferenceFactPropertyExtractionConfig:
    library_dir: Path
    kg_dir: Path
    output_dir: Path
    start: int = 1
    limit: int | None = None
    dry_run: bool = True
    overwrite: bool = False
    max_retries: int = 1
    workers: int = 4
    min_entity_degree: int = 2
    max_evidence_chunks_per_job: int = 12
    entity_disambiguation: bool = True
    disambiguation_lexical_threshold: float = 0.88


@dataclass(frozen=True)
class ReferenceKnowledgeImportConfig:
    library_dir: Path
    kg_dir: Path
    db_path: Path
    facts_dir: Path | None = None
    reset: bool = False
    entity_disambiguation: bool = True
    disambiguation_lexical_threshold: float = 0.88


@dataclass(frozen=True)
class ReferenceKnowledgeQuery:
    db_path: Path
    query: str
    source_doc_ids: tuple[str, ...] = ()
    source_paths: tuple[str, ...] = ()
    source_scope_ids: tuple[str, ...] = ()
    chroma_dir: Path | None = None
    collection_name: str = "dms_reference_knowledge"
    top_k: int = 8
    chunk_top_k: int = 8
    entity_top_k: int = 8
    relationship_top_k: int = 8
    include_fact_properties: bool = True
    fact_binding_top_k: int = 4
    property_binding_top_k: int = 3
    embedding_dim: int = 384
    embedding_provider: str = "hash"
    embedding_model: str | None = None
    embedding_base_url: str | None = None
    embedding_api_key: str | None = None
    embedding_max_tokens: int = 8192
    embedding_timeout: int = 60


@dataclass(frozen=True)
class ChromaReferenceIndexConfig:
    db_path: Path
    persist_dir: Path
    collection_name: str = "dms_reference_knowledge"
    reset: bool = False
    upsert_batch_size: int = 1000
    embedding_dim: int = 384
    embedding_provider: str = "hash"
    embedding_model: str | None = None
    embedding_base_url: str | None = None
    embedding_api_key: str | None = None
    embedding_max_tokens: int = 8192
    embedding_timeout: int = 60


@dataclass(frozen=True)
class ReferenceContextQuery:
    db_path: Path
    query: str
    matched_entities: tuple[dict[str, Any], ...] = ()
    source_doc_ids: tuple[str, ...] = ()
    source_paths: tuple[str, ...] = ()
    source_scope_ids: tuple[str, ...] = ()
    before_scene_id: str | None = None
    before_scene_order: int | None = None
    chroma_dir: Path | None = None
    collection_name: str = "dms_reference_knowledge"
    top_k: int = 6
    author_top_k: int = 6
    character_top_k: int = 6
    style_top_k: int = 4
    timeline_top_k: int = 4
    include_fact_properties: bool = True
    fact_binding_top_k: int = 4
    property_binding_top_k: int = 3
    embedding_dim: int = 384
    embedding_provider: str = "hash"
    embedding_model: str | None = None
    embedding_base_url: str | None = None
    embedding_api_key: str | None = None
    embedding_max_tokens: int = 8192
    embedding_timeout: int = 60


def ingest_reference_library(config: ReferenceLibraryIngestConfig) -> dict[str, Any]:
    input_path = Path(config.input_path)
    output_dir = Path(config.output_dir)
    if output_dir.exists() and any(output_dir.iterdir()):
        if not config.overwrite:
            raise FileExistsError(f"Output directory exists and is not empty: {output_dir}")
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    files = _reference_input_files(input_path)
    root = input_path if input_path.is_dir() else input_path.parent
    max_chunk_chars = max(int(config.max_chunk_chars or 1), 1)
    worker_count = min(max(int(config.workers or 1), 1), max(len(files), 1))
    file_args = [(file_index, path, root, max_chunk_chars) for file_index, path in enumerate(files, start=1)]
    if worker_count == 1 or len(file_args) <= 1:
        file_results = [_load_and_chunk_reference_file(*args) for args in file_args]
    else:
        file_results = []
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            futures = [executor.submit(_load_and_chunk_reference_file, *args) for args in file_args]
            for future in as_completed(futures):
                file_results.append(future.result())
    file_results = sorted(file_results, key=lambda row: int(row.get("file_index") or 0))
    raw_documents = [raw_doc for result in file_results for raw_doc in (result.get("raw_documents") or [])]
    chunks = [chunk for result in file_results for chunk in (result.get("chunks") or [])]

    raw_path = output_dir / "raw_documents.jsonl"
    chunks_path = output_dir / "reference_chunks.jsonl"
    summary_path = output_dir / "summary.json"
    _write_jsonl(raw_path, raw_documents)
    _write_jsonl(chunks_path, chunks)
    summary = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "input_path": str(input_path),
        "output_dir": str(output_dir),
        "file_count": len(files),
        "raw_document_count": len(raw_documents),
        "reference_chunk_count": len(chunks),
        "max_chunk_chars": config.max_chunk_chars,
        "workers": worker_count,
        "artifacts": {
            "raw_documents": str(raw_path),
            "reference_chunks": str(chunks_path),
            "summary": str(summary_path),
        },
    }
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return summary


def extract_reference_kg(
    config: ReferenceKGExtractionConfig,
    *,
    llm_client: LLMClient | None = None,
) -> dict[str, Any]:
    """Extract LightRAG-style entity/relation records from reference chunks."""
    if not config.dry_run and llm_client is None:
        raise ValueError("llm_client is required when dry_run is false")
    if config.start < 1:
        raise ValueError("start must be >= 1")

    output_dir = Path(config.output_dir)
    if output_dir.exists() and any(output_dir.iterdir()):
        if not config.overwrite:
            raise FileExistsError(f"Output directory exists and is not empty: {output_dir}")
        shutil.rmtree(output_dir)
    for directory in (
        output_dir,
        output_dir / "inputs",
        output_dir / "prompts",
        output_dir / "raw_outputs",
        output_dir / "parsed",
    ):
        directory.mkdir(parents=True, exist_ok=True)

    chunks = _read_jsonl(Path(config.library_dir) / "reference_chunks.jsonl")
    selected = _select_records(chunks, start=config.start, limit=config.limit)
    worker_count = min(max(int(config.workers or 1), 1), max(len(selected), 1))
    job_args = [
        (ordinal, chunk, len(selected), output_dir, config, llm_client)
        for ordinal, chunk in enumerate(selected, start=1)
    ]
    if worker_count == 1 or len(job_args) <= 1:
        results = [_extract_reference_kg_chunk(*args) for args in job_args]
    else:
        results = []
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            futures = [executor.submit(_extract_reference_kg_chunk, *args) for args in job_args]
            for future in as_completed(futures):
                results.append(future.result())
    results = sorted(results, key=lambda row: int(row.get("ordinal") or 0))
    chunk_extractions = [row["chunk_extraction"] for row in results if row.get("chunk_extraction")]
    entities = [entity for row in results for entity in (row.get("entities") or [])]
    relationships = [relation for row in results for relation in (row.get("relationships") or [])]
    trace_records = [row["trace"] for row in results]
    entity_type_audit = _reference_entity_type_audit(entities)
    completed_count = sum(int(row.get("llm_completed_count") or 0) for row in results)
    parsed_count = sum(1 for row in results if row.get("status") == "completed")
    failed_count = sum(1 for row in results if row.get("status") == "llm_failed")

    chunk_extractions_path = output_dir / "chunk_extractions.jsonl"
    entities_path = output_dir / "entities.jsonl"
    relationships_path = output_dir / "relationships.jsonl"
    trace_path = output_dir / "trace.jsonl"
    summary_path = output_dir / "summary.json"
    _write_jsonl(chunk_extractions_path, chunk_extractions)
    _write_jsonl(entities_path, entities)
    _write_jsonl(relationships_path, relationships)
    _write_jsonl(trace_path, trace_records)
    summary = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "dry_run_complete" if config.dry_run else "complete",
        "library_dir": str(config.library_dir),
        "output_dir": str(output_dir),
        "selection": {
            "start": config.start,
            "limit": config.limit,
            "selected_count": len(selected),
        },
        "entity_types": list(config.entity_types),
        "llm": {
            "provider": getattr(llm_client, "provider", None) if llm_client else None,
            "model": getattr(llm_client, "model", None) if llm_client else None,
        },
        "llm_completed_count": completed_count,
        "parsed_output_count": parsed_count,
        "failed_count": failed_count,
        "max_retries": config.max_retries,
        "workers": worker_count,
        "entity_extraction_count": len(entities),
        "relationship_extraction_count": len(relationships),
        "entity_type_audit": entity_type_audit,
        "artifacts": {
            "chunk_extractions": str(chunk_extractions_path),
            "entities": str(entities_path),
            "relationships": str(relationships_path),
            "trace": str(trace_path),
            "summary": str(summary_path),
        },
    }
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return summary


def _load_and_chunk_reference_file(
    file_index: int,
    path: Path,
    root: Path,
    max_chunk_chars: int,
) -> dict[str, Any]:
    file_records = _load_reference_file(path, root=root, file_index=file_index)
    chunks = [
        chunk
        for raw_doc in file_records
        for chunk in _chunk_reference_document(raw_doc, max_chunk_chars=max_chunk_chars)
    ]
    return {
        "file_index": file_index,
        "path": str(path),
        "raw_documents": file_records,
        "chunks": chunks,
    }


def _extract_reference_kg_chunk(
    ordinal: int,
    chunk: dict[str, Any],
    total: int,
    output_dir: Path,
    config: ReferenceKGExtractionConfig,
    llm_client: LLMClient | None,
) -> dict[str, Any]:
    print_progress(
        "reference_kg:chunk",
        ordinal - 1,
        total,
        detail=f"chunk={chunk.get('chunk_id')} status=start",
    )
    chunk_id = str(chunk.get("chunk_id") or f"chunk_{ordinal:06d}")
    prompt_context = _reference_chunk_prompt_context(chunk)
    prompt_text = _render_reference_kg_prompt(prompt_context, entity_types=config.entity_types)
    input_path = output_dir / "inputs" / f"{_safe_path_id(chunk_id)}.json"
    prompt_path = output_dir / "prompts" / f"{_safe_path_id(chunk_id)}.txt"
    raw_path = output_dir / "raw_outputs" / f"{_safe_path_id(chunk_id)}.txt"
    parsed_path = output_dir / "parsed" / f"{_safe_path_id(chunk_id)}.json"
    input_path.write_text(json.dumps(prompt_context, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    prompt_path.write_text(prompt_text.rstrip() + "\n", encoding="utf-8")

    raw_text = ""
    status = "dry_run_rendered" if config.dry_run else "llm_failed"
    error = None
    parsed_entities: list[dict[str, Any]] = []
    parsed_relationships: list[dict[str, Any]] = []
    llm_completed_count = 0
    parsed_payload = {
        "chunk_id": chunk_id,
        "status": "not_parsed",
        "reason": "dry_run" if config.dry_run else "llm_failed",
        "entities": [],
        "relationships": [],
    }

    if not config.dry_run:
        attempt_count = max(1, int(config.max_retries or 0) + 1)
        for attempt in range(1, attempt_count + 1):
            try:
                assert llm_client is not None
                result = llm_client.complete(prompt_text)
                llm_completed_count += 1
                raw_text = result.text
                if not raw_text.strip():
                    raise ValueError("empty LLM output")
                parsed_entities, parsed_relationships = _parse_lightrag_kg_output(raw_text, chunk)
                status = "completed"
                error = None
                parsed_payload = {
                    "chunk_id": chunk_id,
                    "status": "parsed",
                    "entities": parsed_entities,
                    "relationships": parsed_relationships,
                    "provider": result.provider,
                    "model": result.model,
                    "usage": result.usage,
                    "attempt": attempt,
                }
                break
            except Exception as exc:  # noqa: BLE001 - preserve per-chunk extraction failures.
                error = str(exc)
                parsed_payload = {
                    "chunk_id": chunk_id,
                    "status": "not_parsed",
                    "reason": "llm_failed",
                    "error": error,
                    "entities": [],
                    "relationships": [],
                    "attempt": attempt,
                }
                if attempt < attempt_count:
                    print_progress(
                        "reference_kg:chunk",
                        ordinal - 1,
                        total,
                        detail=f"chunk={chunk_id} status=retry attempt={attempt + 1}",
                    )
                    continue

    raw_path.write_text(raw_text.rstrip() + "\n", encoding="utf-8")
    parsed_path.write_text(json.dumps(parsed_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    chunk_extraction = None
    if status == "completed":
        chunk_extraction = {
            "chunk_id": chunk_id,
            "doc_id": chunk.get("doc_id"),
            "source_path": chunk.get("source_path"),
            "entity_count": len(parsed_entities),
            "relationship_count": len(parsed_relationships),
            "entities": parsed_entities,
            "relationships": parsed_relationships,
            "attempt": parsed_payload.get("attempt", 1),
        }
    trace = {
        "ordinal": ordinal,
        "chunk_id": chunk_id,
        "doc_id": chunk.get("doc_id"),
        "source_path": chunk.get("source_path"),
        "input_path": str(input_path),
        "prompt_path": str(prompt_path),
        "raw_output_path": str(raw_path),
        "parsed_path": str(parsed_path),
        "status": status,
        "error": error,
        "attempt": parsed_payload.get("attempt", 1),
        "prompt_char_count": len(prompt_text),
        "chunk_char_count": len(str(chunk.get("content") or "")),
        "entity_type_issue_count": _reference_entity_type_issue_count(parsed_entities),
        "multiple_entity_type_count": _reference_entity_type_issue_count(parsed_entities, issue="multiple_entity_types"),
        "unsupported_entity_type_count": _reference_entity_type_issue_count(parsed_entities, issue="unsupported_entity_type"),
    }
    print_progress("reference_kg:chunk", ordinal, total, detail=f"chunk={chunk_id} status={status}")
    return {
        "ordinal": ordinal,
        "status": status,
        "llm_completed_count": llm_completed_count,
        "entities": parsed_entities,
        "relationships": parsed_relationships,
        "chunk_extraction": chunk_extraction,
        "trace": trace,
    }


def extract_reference_facts_properties(
    config: ReferenceFactPropertyExtractionConfig,
    *,
    llm_client: LLMClient | None = None,
) -> dict[str, Any]:
    """Extract source-local atomic facts and entity properties from KG evidence."""
    if not config.dry_run and llm_client is None:
        raise ValueError("llm_client is required when dry_run is false")
    if config.start < 1:
        raise ValueError("start must be >= 1")

    output_dir = Path(config.output_dir)
    if output_dir.exists() and any(output_dir.iterdir()):
        if not config.overwrite:
            raise FileExistsError(f"Output directory exists and is not empty: {output_dir}")
        shutil.rmtree(output_dir)
    for directory in (
        output_dir,
        output_dir / "inputs",
        output_dir / "prompts",
        output_dir / "raw_outputs",
        output_dir / "parsed",
    ):
        directory.mkdir(parents=True, exist_ok=True)

    chunks = _read_jsonl(Path(config.library_dir) / "reference_chunks.jsonl")
    chunk_by_id = {str(chunk.get("chunk_id") or ""): chunk for chunk in chunks}
    extracted_entities = _read_jsonl(Path(config.kg_dir) / "entities.jsonl")
    entity_type_audit = _reference_entity_type_audit(extracted_entities)
    raw_extracted_relationships = _read_jsonl(Path(config.kg_dir) / "relationships.jsonl")
    extracted_relationships = [
        row
        for row in raw_extracted_relationships
        if not _is_negative_reference_relationship(row.get("description"), row.get("keywords"))
    ]
    source_local_entities = _resolve_reference_source_local_entity_types(_source_local_entities(extracted_entities))
    source_local_relationships = _source_local_relationships(extracted_relationships, source_local_entities)
    if config.entity_disambiguation:
        source_local_entities, source_local_relationships = _canonicalize_reference_source_local_graph(
            source_local_entities,
            source_local_relationships,
            lexical_threshold=config.disambiguation_lexical_threshold,
        )
    entity_clusters, entity_cluster_members = _reference_entity_clusters(
        source_local_entities,
        source_local_relationships,
    )
    pseudo_relationships = _reference_pseudo_relationships(source_local_relationships, entity_clusters)
    jobs = _reference_fact_property_jobs(
        source_local_entities,
        source_local_relationships,
        chunk_by_id,
        entity_clusters=entity_clusters,
        pseudo_relationships=pseudo_relationships,
        min_entity_degree=config.min_entity_degree,
        max_evidence_chunks_per_job=config.max_evidence_chunks_per_job,
    )
    selected = _select_records(jobs, start=config.start, limit=config.limit)
    worker_count = min(max(int(config.workers or 1), 1), max(len(selected), 1))
    job_args = [
        (ordinal, job, len(selected), output_dir, config, llm_client)
        for ordinal, job in enumerate(selected, start=1)
    ]
    if worker_count == 1 or len(job_args) <= 1:
        results = [_extract_reference_fact_property_job(*args) for args in job_args]
    else:
        results = []
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            futures = [executor.submit(_extract_reference_fact_property_job, *args) for args in job_args]
            for future in as_completed(futures):
                results.append(future.result())
    results = sorted(results, key=lambda row: int(row.get("ordinal") or 0))
    atomic_facts = _dedupe_reference_records(
        [fact for row in results for fact in (row.get("atomic_facts") or [])],
        key_field="fact_id",
    )
    entity_properties = _dedupe_reference_records(
        [prop for row in results for prop in (row.get("entity_properties") or [])],
        key_field="property_id",
    )
    extractions = [row["extraction"] for row in results if row.get("extraction")]
    trace_records = [row["trace"] for row in results]
    completed_count = sum(int(row.get("llm_completed_count") or 0) for row in results)
    parsed_count = sum(1 for row in results if row.get("status") == "completed")
    failed_count = sum(1 for row in results if row.get("status") == "llm_failed")

    atomic_facts_path = output_dir / "atomic_facts.jsonl"
    entity_properties_path = output_dir / "entity_properties.jsonl"
    extractions_path = output_dir / "fact_property_extractions.jsonl"
    trace_path = output_dir / "trace.jsonl"
    entity_clusters_path = output_dir / "entity_clusters.jsonl"
    entity_cluster_members_path = output_dir / "entity_cluster_members.jsonl"
    pseudo_relationships_path = output_dir / "pseudo_relationships.jsonl"
    summary_path = output_dir / "summary.json"
    _write_jsonl(atomic_facts_path, atomic_facts)
    _write_jsonl(entity_properties_path, entity_properties)
    _write_jsonl(extractions_path, extractions)
    _write_jsonl(trace_path, trace_records)
    _write_jsonl(entity_clusters_path, entity_clusters)
    _write_jsonl(entity_cluster_members_path, entity_cluster_members)
    _write_jsonl(pseudo_relationships_path, pseudo_relationships)
    summary = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "dry_run_complete" if config.dry_run else "complete",
        "library_dir": str(config.library_dir),
        "kg_dir": str(config.kg_dir),
        "output_dir": str(output_dir),
        "selection": {
            "start": config.start,
            "limit": config.limit,
            "selected_count": len(selected),
            "total_jobs": len(jobs),
            "min_entity_degree": config.min_entity_degree,
            "max_evidence_chunks_per_job": config.max_evidence_chunks_per_job,
            "entity_disambiguation": config.entity_disambiguation,
            "disambiguation_lexical_threshold": config.disambiguation_lexical_threshold,
        },
        "entity_cluster_count": len(entity_clusters),
        "entity_cluster_member_count": len(entity_cluster_members),
        "pseudo_relationship_count": len(pseudo_relationships),
        "entity_type_audit": entity_type_audit,
        "eligible_entity_cluster_count": sum(
            1 for cluster in entity_clusters if int(cluster.get("global_pseudo_degree") or 0) >= int(config.min_entity_degree or 0)
        ),
        "batched_fact_property_job_count": len(jobs),
        "llm": {
            "provider": getattr(llm_client, "provider", None) if llm_client else None,
            "model": getattr(llm_client, "model", None) if llm_client else None,
        },
        "llm_completed_count": completed_count,
        "parsed_output_count": parsed_count,
        "failed_count": failed_count,
        "max_retries": config.max_retries,
        "workers": worker_count,
        "atomic_fact_count": len(atomic_facts),
        "entity_property_count": len(entity_properties),
        "artifacts": {
            "atomic_facts": str(atomic_facts_path),
            "entity_properties": str(entity_properties_path),
            "fact_property_extractions": str(extractions_path),
            "trace": str(trace_path),
            "entity_clusters": str(entity_clusters_path),
            "entity_cluster_members": str(entity_cluster_members_path),
            "pseudo_relationships": str(pseudo_relationships_path),
            "summary": str(summary_path),
        },
    }
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return summary


def _extract_reference_fact_property_job(
    ordinal: int,
    job: dict[str, Any],
    total: int,
    output_dir: Path,
    config: ReferenceFactPropertyExtractionConfig,
    llm_client: LLMClient | None,
) -> dict[str, Any]:
    job_id = str(job.get("job_id") or f"job_{ordinal:06d}")
    print_progress(
        "reference_facts_properties:job",
        ordinal - 1,
        total,
        detail=f"job={job_id} status=start",
    )
    prompt_context = _reference_fact_property_prompt_context(job)
    prompt_text = _render_reference_fact_property_prompt(prompt_context)
    input_path = output_dir / "inputs" / f"{_safe_path_id(job_id)}.json"
    prompt_path = output_dir / "prompts" / f"{_safe_path_id(job_id)}.txt"
    raw_path = output_dir / "raw_outputs" / f"{_safe_path_id(job_id)}.txt"
    parsed_path = output_dir / "parsed" / f"{_safe_path_id(job_id)}.json"
    input_path.write_text(json.dumps(prompt_context, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    prompt_path.write_text(prompt_text.rstrip() + "\n", encoding="utf-8")

    raw_text = ""
    status = "dry_run_rendered" if config.dry_run else "llm_failed"
    error = None
    atomic_facts: list[dict[str, Any]] = []
    entity_properties: list[dict[str, Any]] = []
    llm_completed_count = 0
    parsed_payload: dict[str, Any] = {
        "job_id": job_id,
        "status": "not_parsed",
        "reason": "dry_run" if config.dry_run else "llm_failed",
        "atomic_facts": [],
        "entity_properties": [],
    }

    if not config.dry_run:
        attempt_count = max(1, int(config.max_retries or 0) + 1)
        for attempt in range(1, attempt_count + 1):
            try:
                assert llm_client is not None
                result = llm_client.complete(prompt_text)
                llm_completed_count += 1
                raw_text = result.text
                if not raw_text.strip():
                    raise ValueError("empty LLM output")
                atomic_facts, entity_properties = _parse_reference_fact_property_output(raw_text, job)
                status = "completed"
                error = None
                parsed_payload = {
                    "job_id": job_id,
                    "status": "parsed",
                    "atomic_facts": atomic_facts,
                    "entity_properties": entity_properties,
                    "provider": result.provider,
                    "model": result.model,
                    "usage": result.usage,
                    "attempt": attempt,
                }
                break
            except Exception as exc:  # noqa: BLE001 - preserve per-job extraction failures.
                error = str(exc)
                parsed_payload = {
                    "job_id": job_id,
                    "status": "not_parsed",
                    "reason": "llm_failed",
                    "error": error,
                    "atomic_facts": [],
                    "entity_properties": [],
                    "attempt": attempt,
                }
                if attempt < attempt_count:
                    print_progress(
                        "reference_facts_properties:job",
                        ordinal - 1,
                        total,
                        detail=f"job={job_id} status=retry attempt={attempt + 1}",
                    )
                    continue

    raw_path.write_text(raw_text.rstrip() + "\n", encoding="utf-8")
    parsed_path.write_text(json.dumps(parsed_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    extraction = None
    if status == "completed":
        extraction = {
            "job_id": job_id,
            "asset_type": job.get("asset_type"),
            "cluster_job_id": job.get("cluster_job_id"),
            "batch_index": job.get("batch_index"),
            "batch_count": job.get("batch_count"),
            "source_scope_id": job.get("source_scope_id"),
            "cluster_id": job.get("cluster_id"),
            "canonical_display_name": job.get("canonical_display_name"),
            "global_pseudo_degree": job.get("global_pseudo_degree"),
            "source_local_entity_id": job.get("source_local_entity_id"),
            "source_local_relation_id": job.get("source_local_relation_id"),
            "atomic_fact_count": len(atomic_facts),
            "entity_property_count": len(entity_properties),
            "atomic_facts": atomic_facts,
            "entity_properties": entity_properties,
            "attempt": parsed_payload.get("attempt", 1),
        }
    trace = {
        "ordinal": ordinal,
        "job_id": job_id,
        "asset_type": job.get("asset_type"),
        "cluster_job_id": job.get("cluster_job_id"),
        "batch_index": job.get("batch_index"),
        "batch_count": job.get("batch_count"),
        "source_scope_id": job.get("source_scope_id"),
        "cluster_id": job.get("cluster_id"),
        "canonical_display_name": job.get("canonical_display_name"),
        "global_pseudo_degree": job.get("global_pseudo_degree"),
        "input_path": str(input_path),
        "prompt_path": str(prompt_path),
        "raw_output_path": str(raw_path),
        "parsed_path": str(parsed_path),
        "status": status,
        "error": error,
        "attempt": parsed_payload.get("attempt", 1),
        "prompt_char_count": len(prompt_text),
        "evidence_chunk_count": len(job.get("evidence_chunks") or []),
    }
    print_progress("reference_facts_properties:job", ordinal, total, detail=f"job={job_id} status={status}")
    return {
        "ordinal": ordinal,
        "status": status,
        "llm_completed_count": llm_completed_count,
        "atomic_facts": atomic_facts,
        "entity_properties": entity_properties,
        "extraction": extraction,
        "trace": trace,
    }


def import_reference_knowledge(config: ReferenceKnowledgeImportConfig) -> dict[str, Any]:
    """Import LightRAG-style external reference knowledge assets into SQLite."""
    db_path = Path(config.db_path)
    db_path.parent.mkdir(parents=True, exist_ok=True)
    if config.reset and db_path.exists():
        db_path.unlink()

    raw_docs = _read_jsonl(Path(config.library_dir) / "raw_documents.jsonl")
    chunks = _read_jsonl(Path(config.library_dir) / "reference_chunks.jsonl")
    extracted_entities = _read_jsonl(Path(config.kg_dir) / "entities.jsonl")
    entity_type_audit = _reference_entity_type_audit(extracted_entities)
    raw_extracted_relationships = _read_jsonl(Path(config.kg_dir) / "relationships.jsonl")
    extracted_relationships = [
        row
        for row in raw_extracted_relationships
        if not _is_negative_reference_relationship(row.get("description"), row.get("keywords"))
    ]
    filtered_negative_relationships = len(raw_extracted_relationships) - len(extracted_relationships)
    llm_response_cache = _load_reference_llm_response_cache(Path(config.kg_dir))
    source_scopes = _reference_source_scopes(raw_docs)
    source_local_entities = _resolve_reference_source_local_entity_types(_source_local_entities(extracted_entities))
    source_local_relationships = _source_local_relationships(extracted_relationships, source_local_entities)
    if config.entity_disambiguation:
        source_local_entities, source_local_relationships = _canonicalize_reference_source_local_graph(
            source_local_entities,
            source_local_relationships,
            lexical_threshold=config.disambiguation_lexical_threshold,
        )
    extracted_fact_property_assets = _load_reference_fact_property_assets(config.facts_dir)
    source_local_entities, source_local_relationships = _attach_reference_fact_property_assets(
        source_local_entities,
        source_local_relationships,
        extracted_fact_property_assets,
    )
    entity_clusters, entity_cluster_members = _reference_entity_clusters(source_local_entities, source_local_relationships)
    pseudo_relationships = _reference_pseudo_relationships(source_local_relationships, entity_clusters)
    atomic_facts = _atomic_facts_from_source_local(source_local_entities, source_local_relationships)
    entity_properties = _entity_properties_from_source_local(source_local_entities)
    if extracted_fact_property_assets["atomic_facts"] or extracted_fact_property_assets["entity_properties"]:
        atomic_facts = _merge_reference_atomic_facts(atomic_facts, extracted_fact_property_assets["atomic_facts"])
        entity_properties = _merge_reference_entity_properties(entity_properties, extracted_fact_property_assets["entity_properties"])
    full_entities = _full_entities_by_doc(extracted_entities)
    full_relations = _full_relations_by_doc(extracted_relationships)
    entity_chunks = _chunks_by_entity(extracted_entities)
    relation_chunks = _chunks_by_relation(extracted_relationships)

    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        conn.execute(
            "INSERT OR REPLACE INTO reference_metadata(key, value) VALUES (?, ?)",
            ("schema_version", str(REFERENCE_SCHEMA_VERSION)),
        )
        conn.execute(
            "INSERT OR REPLACE INTO reference_metadata(key, value) VALUES (?, ?)",
            ("asset_model", REFERENCE_ASSET_MODEL),
        )
        for doc in raw_docs:
            _insert_reference_full_doc(conn, doc)
        for scope in source_scopes:
            _insert_reference_source_scope(conn, scope)
        for cache_record in llm_response_cache:
            _insert_reference_llm_response_cache(conn, cache_record)
        chunks_by_doc: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for chunk in chunks:
            chunks_by_doc[str(chunk.get("doc_id") or "")].append(chunk)
            _insert_reference_text_chunk(conn, chunk)
            _insert_reference_vector_document(conn, _chunk_vector_document(chunk))
        for doc in raw_docs:
            doc_id = str(doc.get("doc_id") or "")
            doc_chunks = chunks_by_doc.get(doc_id, [])
            _insert_reference_doc_status(conn, doc, doc_chunks)
        for entity in extracted_entities:
            _insert_reference_extracted_entity(conn, entity)
        for relationship in extracted_relationships:
            _insert_reference_extracted_relationship(conn, relationship)
        for entity in source_local_entities:
            _insert_reference_source_local_entity(conn, entity)
            _insert_reference_vector_document(conn, _source_local_entity_vector_document(entity))
        for relation in source_local_relationships:
            _insert_reference_source_local_relationship(conn, relation)
            _insert_reference_vector_document(conn, _source_local_relationship_vector_document(relation))
        for cluster in entity_clusters:
            _insert_reference_entity_cluster(conn, cluster)
            _insert_reference_vector_document(conn, _entity_cluster_vector_document(cluster))
        for member in entity_cluster_members:
            _insert_reference_entity_cluster_member(conn, member)
        for pseudo_relation in pseudo_relationships:
            _insert_reference_pseudo_relationship(conn, pseudo_relation)
            _insert_reference_vector_document(conn, _pseudo_relationship_vector_document(pseudo_relation))
        for fact in atomic_facts:
            _insert_reference_atomic_fact(conn, fact)
            _insert_reference_vector_document(conn, _atomic_fact_vector_document(fact))
        for prop in entity_properties:
            _insert_reference_entity_property(conn, prop)
            _insert_reference_vector_document(conn, _entity_property_vector_document(prop))
        for doc_id, entity_names in full_entities.items():
            _insert_reference_full_entities(conn, doc_id, entity_names)
        for doc_id, relation_pairs in full_relations.items():
            _insert_reference_full_relations(conn, doc_id, relation_pairs)
        for entity_name, chunk_ids in entity_chunks.items():
            _insert_reference_entity_chunks(conn, entity_name, chunk_ids)
        for relation_id, chunk_ids in relation_chunks.items():
            _insert_reference_relation_chunks(conn, relation_id, chunk_ids)
        conn.commit()

    return {
        "db_path": str(db_path),
        "library_dir": str(config.library_dir),
        "kg_dir": str(config.kg_dir),
        "schema_version": REFERENCE_SCHEMA_VERSION,
        "asset_model": REFERENCE_ASSET_MODEL,
        "full_docs": len(raw_docs),
        "source_scopes": len(source_scopes),
        "llm_response_cache": len(llm_response_cache),
        "text_chunks": len(chunks),
        "doc_status": len(raw_docs),
        "chunk_vectors": len(chunks),
        "extracted_entities": len(extracted_entities),
        "raw_extracted_relationships": len(raw_extracted_relationships),
        "extracted_relationships": len(extracted_relationships),
        "filtered_negative_relationships": filtered_negative_relationships,
        "entity_type_audit": entity_type_audit,
        "source_local_entities": len(source_local_entities),
        "source_local_relationships": len(source_local_relationships),
        "entity_clusters": len(entity_clusters),
        "entity_cluster_members": len(entity_cluster_members),
        "pseudo_relationships": len(pseudo_relationships),
        "atomic_facts": len(atomic_facts),
        "entity_properties": len(entity_properties),
        "extracted_atomic_facts": len(extracted_fact_property_assets["atomic_facts"]),
        "extracted_entity_properties": len(extracted_fact_property_assets["entity_properties"]),
        "facts_dir": str(config.facts_dir) if config.facts_dir else None,
        "entity_disambiguation": config.entity_disambiguation,
        "disambiguation_lexical_threshold": config.disambiguation_lexical_threshold,
        "entity_vectors": len(source_local_entities) + len(entity_clusters),
        "relationship_vectors": len(source_local_relationships) + len(pseudo_relationships),
        "claim_vectors": len(atomic_facts) + len(entity_properties),
        "full_entities": len(full_entities),
        "full_relations": len(full_relations),
        "entity_chunks": len(entity_chunks),
        "relation_chunks": len(relation_chunks),
    }


def build_chroma_reference_index(config: ChromaReferenceIndexConfig) -> dict[str, Any]:
    chromadb = _import_chromadb()
    persist_dir = Path(config.persist_dir)
    if config.reset and persist_dir.exists():
        shutil.rmtree(persist_dir)
    persist_dir.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(path=str(persist_dir))
    if config.reset:
        try:
            client.delete_collection(config.collection_name)
        except Exception:
            pass
    embedding_function = build_embedding_function(
        provider=config.embedding_provider,
        embedding_dim=config.embedding_dim,
        model_name=config.embedding_model,
        base_url=config.embedding_base_url,
        api_key=config.embedding_api_key,
        max_tokens=config.embedding_max_tokens,
        timeout=config.embedding_timeout,
    )
    collection = client.get_or_create_collection(
        name=config.collection_name,
        embedding_function=embedding_function,
        metadata={"hnsw:space": "cosine"},
    )
    records = list_reference_vector_documents(config.db_path)
    batches = _batches(records, max(int(config.upsert_batch_size or 1), 1))
    print_progress(
        "reference_chroma_index:start",
        0,
        len(batches),
        detail=f"documents={len(records)} collection={config.collection_name} persist_dir={persist_dir}",
    )
    batch_count = 0
    for batch in batches:
        batch_count += 1
        collection.upsert(
            ids=[str(record["vector_id"]) for record in batch],
            documents=[str(record["content"]) for record in batch],
            metadatas=[_reference_chroma_metadata(record) for record in batch],
        )
        print_progress(
            "reference_chroma_index:batch",
            batch_count,
            len(batches),
            detail=f"documents={len(batch)}",
        )
    return {
        "db_path": str(config.db_path),
        "persist_dir": str(persist_dir),
        "collection_name": config.collection_name,
        "document_count": len(records),
        "upsert_batch_size": max(int(config.upsert_batch_size or 1), 1),
        "upsert_batch_count": batch_count,
        "embedding": embedding_function.config_summary(),
    }


def build_reference_context(query: ReferenceContextQuery) -> tuple[dict[str, Any], dict[str, Any]]:
    top_k = max(int(query.top_k or 0), 0)
    if top_k <= 0:
        return _empty_reference_context(), {
            "enabled": True,
            "strategy": "disabled_by_top_k",
            "returned_counts": _reference_context_counts(_empty_reference_context()),
        }

    if _has_source_local_reference_assets(query.db_path):
        return _build_lightrag_reference_context(query)

    grouped = _empty_reference_context()
    return grouped, {
        "enabled": True,
        "strategy": "source_local_assets_missing",
        "asset_model": REFERENCE_ASSET_MODEL,
        "db_path": str(query.db_path),
        "query": query.query,
        "candidate_count": 0,
        "returned_counts": _reference_context_counts(grouped),
        "warning": "External reference context requires import-reference-knowledge assets.",
    }


def _build_lightrag_reference_context(query: ReferenceContextQuery) -> tuple[dict[str, Any], dict[str, Any]]:
    kg_result = query_reference_knowledge(
        ReferenceKnowledgeQuery(
            db_path=query.db_path,
            query=query.query,
            source_doc_ids=query.source_doc_ids,
            source_paths=query.source_paths,
            source_scope_ids=query.source_scope_ids,
            chroma_dir=query.chroma_dir,
            collection_name=query.collection_name,
            top_k=query.top_k,
            chunk_top_k=query.top_k,
            entity_top_k=query.author_top_k,
            relationship_top_k=query.author_top_k,
            include_fact_properties=query.include_fact_properties,
            fact_binding_top_k=query.fact_binding_top_k,
            property_binding_top_k=query.property_binding_top_k,
            embedding_dim=query.embedding_dim,
            embedding_provider=query.embedding_provider,
            embedding_model=query.embedding_model,
            embedding_base_url=query.embedding_base_url,
            embedding_api_key=query.embedding_api_key,
            embedding_max_tokens=query.embedding_max_tokens,
            embedding_timeout=query.embedding_timeout,
        )
    )
    matched_names = _matched_entity_names(query.matched_entities)
    author_items = []
    character_items = []
    timeline_items = []
    for entity in kg_result.get("entities") or []:
        item = _packet_lightrag_entity_item(entity)
        author_items.append(item)
        if _entity_item_visible_to_character(item, matched_names):
            character_items.append(item)
        if _looks_temporal_text(item.get("statement")):
            timeline_items.append(item)
    for relation in kg_result.get("relationships") or []:
        item = _packet_lightrag_relationship_item(relation)
        author_items.append(item)
        if _relation_item_visible_to_character(item, matched_names):
            character_items.append(item)
        if _looks_temporal_text(item.get("statement")):
            timeline_items.append(item)
    for fact in kg_result.get("relationship_facts") or []:
        item = _packet_lightrag_fact_item(fact)
        author_items.append(item)
        if _fact_item_visible_to_character(item, matched_names):
            character_items.append(item)
        if _looks_temporal_text(item.get("statement")):
            timeline_items.append(item)
    for fact in kg_result.get("atomic_facts") or []:
        item = _packet_lightrag_fact_item(fact)
        author_items.append(item)
        if _fact_item_visible_to_character(item, matched_names):
            character_items.append(item)
        if _looks_temporal_text(item.get("statement")):
            timeline_items.append(item)
    for prop in kg_result.get("entity_properties") or []:
        item = _packet_lightrag_property_item(prop)
        author_items.append(item)
        if _entity_item_visible_to_character(item, matched_names):
            character_items.append(item)
    for chunk in kg_result.get("chunks") or []:
        item = _packet_lightrag_chunk_item(chunk)
        author_items.append(item)
        if _looks_temporal_text(item.get("statement")):
            timeline_items.append(item)
    grouped = {
        "author_reference_context": _dedupe_packet_items(author_items)[: max(query.author_top_k or query.top_k, 0)],
        "character_reference_knowledge": _dedupe_packet_items(character_items)[: max(query.character_top_k or query.top_k, 0)],
        "style_reference_context": [],
        "timeline_reference_claims": _dedupe_packet_items(timeline_items)[: max(query.timeline_top_k or query.top_k, 0)],
    }
    trace = {
        "enabled": True,
        "strategy": kg_result.get("retrieval_strategy"),
        "asset_model": REFERENCE_ASSET_MODEL,
        "db_path": str(query.db_path),
        "chroma_dir": str(query.chroma_dir) if query.chroma_dir else None,
        "collection_name": query.collection_name if query.chroma_dir else None,
        "query": query.query,
        "source_filter": kg_result.get("source_filter"),
        "matched_entities": [entity.get("canonical_name") or entity.get("entity_id") for entity in query.matched_entities],
        "candidate_count": kg_result.get("count"),
        "returned_counts": _reference_context_counts(grouped),
        "visibility_policy": {
            "author_reference_context": "LightRAG-style external KG context is author-facing reference by default, not canon.",
            "character_reference_knowledge": "Only entity/relation context matching current character names is exposed as character knowledge.",
            "style_reference_context": "Style guidance requires a source-local style classifier in the external KG pipeline.",
            "timeline_reference_claims": "Temporal-looking entity/relation/chunk context is surfaced as author/system timeline reference.",
        },
    }
    return grouped, trace


def query_reference_knowledge(query: ReferenceKnowledgeQuery) -> dict[str, Any]:
    """Source-local retrieval over external reference chunks, claims, and pseudo graph assets."""
    top_k = max(int(query.top_k or 0), 0)
    include_fact_properties = bool(query.include_fact_properties)
    fact_binding_top_k = max(int(query.fact_binding_top_k or 0), 0)
    property_binding_top_k = max(int(query.property_binding_top_k or 0), 0)
    if top_k <= 0:
        return {
            "query": query.query,
            "mode": "source_local",
            "count": 0,
            "entities": [],
            "relationships": [],
            "relationship_facts": [],
            "atomic_facts": [],
            "entity_properties": [],
            "matched_clusters": [],
            "pseudo_relationships": [],
            "chunks": [],
            "source_filter": _reference_source_filter_summary(query),
            "evidence_board": _empty_reference_evidence_board(),
        }
    source_doc_ids = _resolve_reference_source_doc_ids(
        query.db_path,
        source_doc_ids=query.source_doc_ids,
        source_paths=query.source_paths,
    )
    source_scope_ids = _resolve_reference_source_scope_ids(
        query.db_path,
        source_scope_ids=query.source_scope_ids,
        source_doc_ids=source_doc_ids,
        source_paths=query.source_paths,
    )
    vector_top_k = max(query.chunk_top_k, query.entity_top_k, query.relationship_top_k, top_k)
    if include_fact_properties:
        vector_top_k = max(vector_top_k, top_k * max(fact_binding_top_k, property_binding_top_k, 1))
    vector_docs = _search_reference_vector_documents(
        query.db_path,
        query=query.query,
        chroma_dir=query.chroma_dir,
        collection_name=query.collection_name,
        source_doc_ids=source_doc_ids,
        source_scope_ids=source_scope_ids,
        top_k=vector_top_k,
        embedding_dim=query.embedding_dim,
        embedding_provider=query.embedding_provider,
        embedding_model=query.embedding_model,
        embedding_base_url=query.embedding_base_url,
        embedding_api_key=query.embedding_api_key,
        embedding_max_tokens=query.embedding_max_tokens,
        embedding_timeout=query.embedding_timeout,
    )
    chunk_hits = [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_CHUNKS][: max(query.chunk_top_k, 0)]
    entity_hits = _dedupe_vector_hits(
        [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_SOURCE_LOCAL_ENTITIES],
        key_field="source_local_entity_id",
    )[: max(query.entity_top_k, 0)]
    relation_hits = _dedupe_vector_hits(
        [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_SOURCE_LOCAL_RELATIONSHIPS],
        key_field="source_local_relation_id",
    )[: max(query.relationship_top_k, 0)]
    fact_hits = (
        _dedupe_vector_hits(
            [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_ATOMIC_FACTS],
            key_field="fact_id",
        )[:top_k]
        if include_fact_properties
        else []
    )
    property_hits = (
        _dedupe_vector_hits(
            [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_ENTITY_PROPERTIES],
            key_field="property_id",
        )[:top_k]
        if include_fact_properties
        else []
    )
    cluster_hits = _dedupe_vector_hits(
        [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_ENTITY_CLUSTERS],
        key_field="cluster_id",
    )[:top_k]
    pseudo_relation_hits = _dedupe_vector_hits(
        [row for row in vector_docs if row.get("vector_namespace") == VECTOR_NAMESPACE_PSEUDO_RELATIONSHIPS],
        key_field="pseudo_relation_id",
    )[:top_k]

    entities = [
        item
        for item in (
            _source_local_entity_payload_from_vector_hit(
                query.db_path,
                hit,
                source_doc_ids=source_doc_ids,
                source_scope_ids=source_scope_ids,
            )
            for hit in entity_hits
        )
        if item
    ]
    relationships = [
        item
        for item in (
            _source_local_relationship_payload_from_vector_hit(
                query.db_path,
                hit,
                source_doc_ids=source_doc_ids,
                source_scope_ids=source_scope_ids,
            )
            for hit in relation_hits
        )
        if item
    ]
    atomic_facts = [
        item
        for item in (
            _atomic_fact_payload_from_vector_hit(
                query.db_path,
                hit,
                source_doc_ids=source_doc_ids,
                source_scope_ids=source_scope_ids,
            )
            for hit in fact_hits
        )
        if item
    ]
    entity_properties = [
        item
        for item in (
            _entity_property_payload_from_vector_hit(
                query.db_path,
                hit,
                source_doc_ids=source_doc_ids,
                source_scope_ids=source_scope_ids,
            )
            for hit in property_hits
        )
        if item and _is_displayable_entity_property(item)
    ]
    if include_fact_properties:
        direct_fact_scores = {
            str(hit.get("fact_id") or ""): float(hit.get("score") or 0.0)
            for hit in fact_hits
            if str(hit.get("fact_id") or "").strip()
        }
        direct_property_scores = {
            str(hit.get("property_id") or ""): float(hit.get("score") or 0.0)
            for hit in property_hits
            if str(hit.get("property_id") or "").strip()
        }
        bound_facts = _bound_reference_atomic_facts_for_entities(
            query.db_path,
            query_text=query.query,
            entities=entities,
            source_doc_ids=source_doc_ids,
            source_scope_ids=source_scope_ids,
            top_k=fact_binding_top_k,
            direct_scores=direct_fact_scores,
        )
        bound_properties = _bound_reference_properties_for_entities(
            query.db_path,
            query_text=query.query,
            entities=entities,
            source_doc_ids=source_doc_ids,
            source_scope_ids=source_scope_ids,
            top_k=property_binding_top_k,
            direct_scores=direct_property_scores,
        )
        atomic_facts = _limit_reference_items_per_binding_group(
            _merge_reference_query_items(atomic_facts, bound_facts, key_field="fact_id"),
            top_k=fact_binding_top_k,
            fallback_limit=max(top_k, max(len(entities), 1) * max(fact_binding_top_k, 1)),
        )
        entity_properties = _limit_reference_items_per_binding_group(
            _merge_reference_query_items(entity_properties, bound_properties, key_field="property_id"),
            top_k=property_binding_top_k,
            fallback_limit=max(top_k, max(len(entities), 1) * max(property_binding_top_k, 1)),
        )
    else:
        atomic_facts = []
        entity_properties = []
    matched_clusters = [
        item
        for item in (
            _entity_cluster_payload_from_vector_hit(query.db_path, hit, source_scope_ids=source_scope_ids, source_doc_ids=source_doc_ids)
            for hit in cluster_hits
        )
        if item
    ]
    pseudo_relationships = [
        item
        for item in (
            _pseudo_relationship_payload_from_vector_hit(query.db_path, hit, source_scope_ids=source_scope_ids, source_doc_ids=source_doc_ids)
            for hit in pseudo_relation_hits
        )
        if item
    ]
    direct_chunks = [
        item
        for item in (
            _chunk_payload_from_vector_hit(query.db_path, hit, source_doc_ids=source_doc_ids, source_scope_ids=source_scope_ids)
            for hit in chunk_hits
        )
        if item
    ]
    chunks = _supporting_reference_chunks(
        query.db_path,
        direct_chunks=direct_chunks,
        entities=entities,
        relationships=relationships,
        atomic_facts=atomic_facts,
        entity_properties=entity_properties,
        top_k=max(query.chunk_top_k, top_k),
        source_doc_ids=source_doc_ids,
        source_scope_ids=source_scope_ids,
    )
    evidence_chunk_index = _reference_chunk_index_by_ids(
        query.db_path,
        _reference_evidence_chunk_ids(
            direct_chunks=chunks,
            entities=entities,
            relationships=relationships,
            atomic_facts=atomic_facts,
            entity_properties=entity_properties,
        ),
        source_doc_ids=source_doc_ids,
        source_scope_ids=source_scope_ids,
    )
    _attach_reference_evidence_chunks(entities, evidence_chunk_index, chunk_id_getter=_entity_evidence_chunk_ids)
    _attach_reference_evidence_chunks(relationships, evidence_chunk_index, chunk_id_getter=_relationship_evidence_chunk_ids)
    _attach_reference_evidence_chunks(atomic_facts, evidence_chunk_index, chunk_id_getter=_fact_evidence_chunk_ids)
    _attach_reference_evidence_chunks(entity_properties, evidence_chunk_index, chunk_id_getter=_property_evidence_chunk_ids)
    relationship_facts = _relationship_facts_for_query(
        relationships=relationships,
        atomic_facts=atomic_facts,
        evidence_chunk_index=evidence_chunk_index,
        top_k=top_k,
    )
    evidence_board = _reference_evidence_board(
        source_filter=_reference_source_filter_summary(
            query,
            resolved_source_doc_ids=source_doc_ids,
            resolved_source_scope_ids=source_scope_ids,
        ),
        matched_clusters=matched_clusters,
        entities=entities,
        relationships=relationships,
        relationship_facts=relationship_facts,
        atomic_facts=atomic_facts,
        entity_properties=entity_properties,
        chunks=chunks,
        pseudo_relationships=pseudo_relationships,
    )
    return {
        "query": query.query,
        "mode": "source_local",
        "retrieval_strategy": "chroma_sql_filtered" if query.chroma_dir else "sql_ranked",
        "asset_model": REFERENCE_ASSET_MODEL,
        "source_filter": evidence_board["source_filter"],
        "count": len(entities) + len(relationships) + len(atomic_facts) + len(entity_properties) + len(chunks),
        "matched_clusters": matched_clusters,
        "entities": entities,
        "relationships": relationships,
        "relationship_facts": relationship_facts,
        "atomic_facts": atomic_facts,
        "entity_properties": entity_properties,
        "pseudo_relationships": pseudo_relationships,
        "chunks": chunks,
        "evidence_board": evidence_board,
    }


def list_reference_context_records(
    db_path: str | Path,
    *,
    item_types: set[str] | None = None,
    knowledge_scopes: set[str] | None = None,
    before_scene_order: int | None = None,
    before_scene_id: str | None = None,
) -> list[dict[str, Any]]:
    max_order = before_scene_order
    if max_order is None and before_scene_id:
        max_order = _max_scene_order_before(before_scene_id)

    if not _has_source_local_reference_assets(db_path):
        return []
    return _filter_reference_context_records(
        _list_lightrag_reference_context_records(db_path),
        item_types=item_types,
        knowledge_scopes=knowledge_scopes,
        max_order=max_order,
    )


def _list_lightrag_reference_context_records(db_path: str | Path) -> list[dict[str, Any]]:
    """Expose LightRAG-style KG assets as external context records for the unified kernel."""
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        chunk_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_text_chunks ORDER BY full_doc_id, chunk_order_index")]
        doc_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_full_docs ORDER BY doc_id")]
        entity_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_source_local_entities ORDER BY source_scope_id, entity_name")]
        relation_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_source_local_relationships ORDER BY source_scope_id, src_id, tgt_id")]
        fact_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_atomic_facts ORDER BY source_scope_id, subject, fact_id")]
        property_rows = [dict(row) for row in conn.execute("SELECT * FROM reference_entity_properties ORDER BY source_scope_id, entity_name, property_id")]

    chunk_by_id = {str(row.get("chunk_id") or ""): row for row in chunk_rows}
    doc_by_id = {str(row.get("doc_id") or ""): row for row in doc_rows}
    records: list[dict[str, Any]] = []

    for raw_row in entity_rows:
        row = _source_local_entity_payload(raw_row)
        entity_name = str(row.get("entity_name") or "").strip()
        if not entity_name:
            continue
        source_chunk_ids = _split_graph_field(row.get("source_id"))
        source_doc_ids = row.get("source_doc_ids") or _source_doc_ids_for_context_chunks(source_chunk_ids, chunk_by_id)
        records.append(
            {
                "item_id": f"external_source_local_entity:{row.get('source_local_entity_id')}",
                "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
                "source_doc_ids": source_doc_ids,
                "chunk_id": source_chunk_ids[0] if source_chunk_ids else "",
                "source_chunk_ids": source_chunk_ids,
                "item_type": _lightrag_entity_context_item_type(row.get("entity_type")),
                "subject": entity_name,
                "statement": str(row.get("description") or "").strip(),
                "evidence": str(row.get("description") or "").strip(),
                "authority": 0.7,
                "knowledge_scope": "author_only",
                "known_to": [],
                "available_from": "unknown",
                "available_from_order": None,
                "timeline_hint": "",
                "confidence": 0.75,
                "source_role": "external_reference_entity",
                "metadata": {
                    "asset_model": REFERENCE_ASSET_MODEL,
                    "source_role": "external_reference_entity",
                    "source_scope_id": row.get("source_scope_id"),
                    "source_local_entity_id": row.get("source_local_entity_id"),
                    "cluster_id": row.get("cluster_id"),
                    "entity_name": entity_name,
                    "entity_type": row.get("entity_type"),
                    "facts": row.get("facts") or [],
                    "attributes": row.get("attributes") or [],
                    "source_doc_ids": source_doc_ids,
                    "source_chunk_ids": source_chunk_ids,
                    "source_docs": _context_source_docs(source_doc_ids, doc_by_id),
                    "raw": _json_loads(row.get("raw_json"), default={}),
                },
            }
        )

    for raw_row in relation_rows:
        row = _source_local_relationship_payload(raw_row)
        source_chunk_ids = _split_graph_field(row.get("source_id"))
        source_doc_ids = row.get("source_doc_ids") or _source_doc_ids_for_context_chunks(source_chunk_ids, chunk_by_id)
        src_id = str(row.get("src_id") or "").strip()
        tgt_id = str(row.get("tgt_id") or "").strip()
        relation_id = str(row.get("relation_id") or _reference_relation_id(src_id, tgt_id))
        records.append(
            {
                "item_id": f"external_source_local_relation:{row.get('source_local_relation_id')}",
                "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
                "source_doc_ids": source_doc_ids,
                "chunk_id": source_chunk_ids[0] if source_chunk_ids else "",
                "source_chunk_ids": source_chunk_ids,
                "item_type": "relationship_fact",
                "subject": f"{src_id} - {tgt_id}".strip(" -"),
                "statement": str(row.get("description") or "").strip(),
                "evidence": str(row.get("description") or "").strip(),
                "authority": 0.7,
                "knowledge_scope": "author_only",
                "known_to": [],
                "available_from": "unknown",
                "available_from_order": None,
                "timeline_hint": "",
                "confidence": 0.75,
                "source_role": "external_reference_relation",
                "metadata": {
                    "asset_model": REFERENCE_ASSET_MODEL,
                    "source_role": "external_reference_relation",
                    "source_scope_id": row.get("source_scope_id"),
                    "source_local_relation_id": row.get("source_local_relation_id"),
                    "pseudo_relation_id": row.get("pseudo_relation_id"),
                    "relation_id": relation_id,
                    "src_id": src_id,
                    "tgt_id": tgt_id,
                    "keywords": row.get("keywords"),
                    "facts": row.get("facts") or [],
                    "source_doc_ids": source_doc_ids,
                    "source_chunk_ids": source_chunk_ids,
                    "source_docs": _context_source_docs(source_doc_ids, doc_by_id),
                    "raw": _json_loads(row.get("raw_json"), default={}),
                },
            }
        )

    for raw_row in fact_rows:
        row = _atomic_fact_payload(raw_row)
        source_doc_id = str(row.get("source_doc_id") or "")
        records.append(
            {
                "item_id": f"external_fact:{row.get('fact_id')}",
                "source_doc_id": source_doc_id,
                "source_doc_ids": [source_doc_id] if source_doc_id else [],
                "chunk_id": row.get("source_chunk_id"),
                "source_chunk_ids": [row.get("source_chunk_id")] if row.get("source_chunk_id") else [],
                "item_type": "relationship_fact" if row.get("source_local_relation_id") else "world_bible",
                "subject": row.get("subject"),
                "statement": row.get("statement"),
                "evidence": row.get("statement"),
                "authority": 0.72,
                "knowledge_scope": "author_only",
                "known_to": [],
                "available_from": "unknown",
                "available_from_order": None,
                "timeline_hint": row.get("statement") if _looks_temporal_text(row.get("statement")) else "",
                "confidence": row.get("confidence"),
                "source_role": "external_reference_fact",
                "metadata": {
                    "asset_model": REFERENCE_ASSET_MODEL,
                    "source_role": "external_reference_fact",
                    "source_scope_id": row.get("source_scope_id"),
                    "fact_id": row.get("fact_id"),
                    "source_local_entity_id": row.get("source_local_entity_id"),
                    "source_local_relation_id": row.get("source_local_relation_id"),
                    "cluster_id": row.get("cluster_id"),
                "source_doc_ids": [source_doc_id] if source_doc_id else [],
                "source_chunk_ids": [row.get("source_chunk_id")] if row.get("source_chunk_id") else [],
                "source_docs": _context_source_docs([source_doc_id] if source_doc_id else [], doc_by_id),
                "raw": row.get("raw") or {},
            },
        }
        )

    for raw_row in property_rows:
        row = _entity_property_payload(raw_row)
        source_doc_id = str(row.get("source_doc_id") or "")
        records.append(
            {
                "item_id": f"external_property:{row.get('property_id')}",
                "source_doc_id": source_doc_id,
                "source_doc_ids": [source_doc_id] if source_doc_id else [],
                "chunk_id": row.get("source_chunk_id"),
                "source_chunk_ids": [row.get("source_chunk_id")] if row.get("source_chunk_id") else [],
                "item_type": "character_profile" if row.get("property_name") == "entity_type" and row.get("property_value") == "character" else "world_bible",
                "subject": row.get("entity_name"),
                "statement": row.get("statement"),
                "evidence": row.get("statement"),
                "authority": 0.7,
                "knowledge_scope": "author_only",
                "known_to": [],
                "available_from": "unknown",
                "available_from_order": None,
                "timeline_hint": "",
                "confidence": row.get("confidence"),
                "source_role": "external_reference_property",
                "metadata": {
                    "asset_model": REFERENCE_ASSET_MODEL,
                    "source_role": "external_reference_property",
                    "source_scope_id": row.get("source_scope_id"),
                    "property_id": row.get("property_id"),
                    "source_local_entity_id": row.get("source_local_entity_id"),
                    "cluster_id": row.get("cluster_id"),
                    "property_name": row.get("property_name"),
                    "property_value": row.get("property_value"),
                    "source_doc_ids": [source_doc_id] if source_doc_id else [],
                    "source_chunk_ids": [row.get("source_chunk_id")] if row.get("source_chunk_id") else [],
                    "source_docs": _context_source_docs([source_doc_id] if source_doc_id else [], doc_by_id),
                    "raw": row.get("raw") or {},
                },
            }
        )

    for row in chunk_rows:
        content = str(row.get("content") or "").strip()
        if not content:
            continue
        full_doc_id = str(row.get("full_doc_id") or "").strip()
        records.append(
            {
                "item_id": f"external_chunk:{row.get('chunk_id')}",
                "source_doc_id": full_doc_id,
                "source_doc_ids": [full_doc_id] if full_doc_id else [],
                "chunk_id": row.get("chunk_id"),
                "source_chunk_ids": [row.get("chunk_id")] if row.get("chunk_id") else [],
                "item_type": "author_note",
                "subject": str(row.get("heading") or row.get("title") or full_doc_id or "external reference").strip(),
                "statement": content,
                "evidence": content,
                "authority": 0.6,
                "knowledge_scope": "author_only",
                "known_to": [],
                "available_from": "unknown",
                "available_from_order": None,
                "timeline_hint": "",
                "confidence": 0.7,
                "source_role": "external_reference_chunk",
                "metadata": {
                    "asset_model": REFERENCE_ASSET_MODEL,
                    "source_role": "external_reference_chunk",
                    "source_scope_id": row.get("source_scope_id"),
                    "source_doc_ids": [full_doc_id] if full_doc_id else [],
                    "source_chunk_ids": [row.get("chunk_id")] if row.get("chunk_id") else [],
                    "source_docs": _context_source_docs([full_doc_id] if full_doc_id else [], doc_by_id),
                    "raw": _json_loads(row.get("raw_json"), default={}),
                },
            }
        )

    return sorted(records, key=lambda item: (-float(item.get("authority") or 0), -float(item.get("confidence") or 0), str(item.get("item_id") or "")))


def _filter_reference_context_records(
    records: list[dict[str, Any]],
    *,
    item_types: set[str] | None,
    knowledge_scopes: set[str] | None,
    max_order: int | None,
) -> list[dict[str, Any]]:
    filtered = []
    for record in records:
        if item_types and record.get("item_type") not in item_types:
            continue
        if knowledge_scopes and record.get("knowledge_scope") not in knowledge_scopes:
            continue
        available_order = record.get("available_from_order")
        if max_order is not None and available_order is not None and int(available_order) > max_order:
            continue
        filtered.append(record)
    return filtered


def list_reference_vector_documents(
    db_path: str | Path,
    *,
    vector_namespaces: set[str] | None = None,
    source_doc_ids: set[str] | None = None,
    source_scope_ids: set[str] | None = None,
) -> list[dict[str, Any]]:
    sql = "SELECT * FROM reference_vector_documents WHERE 1 = 1"
    params: list[Any] = []
    if vector_namespaces:
        placeholders = ", ".join("?" for _ in vector_namespaces)
        sql += f" AND vector_namespace IN ({placeholders})"
        params.extend(sorted(vector_namespaces))
    sql += " ORDER BY vector_namespace, source_doc_id, vector_id"
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        rows = [_reference_vector_document_payload(row) for row in conn.execute(sql, params)]
    if source_doc_ids:
        rows = [
            row
            for row in rows
            if _vector_document_matches_sources(row, source_doc_ids=source_doc_ids, source_scope_ids=None)
        ]
    if source_scope_ids:
        rows = [
            row
            for row in rows
            if _vector_document_matches_sources(row, source_doc_ids=None, source_scope_ids=source_scope_ids)
        ]
    return rows


def get_reference_asset_counts(db_path: str | Path) -> dict[str, int]:
    tables = [
        "reference_full_docs",
        "reference_source_scopes",
        "reference_doc_status",
        "reference_llm_response_cache",
        "reference_text_chunks",
        "reference_extracted_entities",
        "reference_extracted_relationships",
        "reference_source_local_entities",
        "reference_source_local_relationships",
        "reference_entity_clusters",
        "reference_entity_cluster_members",
        "reference_pseudo_relationships",
        "reference_atomic_facts",
        "reference_entity_properties",
        "reference_vector_documents",
        "reference_full_entities",
        "reference_full_relations",
        "reference_entity_chunks",
        "reference_relation_chunks",
    ]
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        return {
            table: int(conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0])
            for table in tables
        }


def _search_reference_vector_documents(
    db_path: str | Path,
    *,
    query: str,
    chroma_dir: Path | None,
    collection_name: str,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
    top_k: int,
    embedding_dim: int,
    embedding_provider: str,
    embedding_model: str | None,
    embedding_base_url: str | None,
    embedding_api_key: str | None,
    embedding_max_tokens: int,
    embedding_timeout: int,
) -> list[dict[str, Any]]:
    sql_docs = list_reference_vector_documents(
        db_path,
        source_doc_ids=source_doc_ids or None,
        source_scope_ids=source_scope_ids or None,
    )
    if not sql_docs:
        return []
    limit = max(int(top_k or 1), 1)
    if chroma_dir is None:
        return _rank_reference_vector_sql_docs(sql_docs, query)[: max(limit * 3, limit)]

    chromadb = _import_chromadb()
    docs_by_id = {str(row["vector_id"]): row for row in sql_docs}
    allowed_doc_ids = set(docs_by_id)
    client = chromadb.PersistentClient(path=str(chroma_dir))
    embedding_function = build_embedding_function(
        provider=embedding_provider,
        embedding_dim=embedding_dim,
        model_name=embedding_model,
        base_url=embedding_base_url,
        api_key=embedding_api_key,
        max_tokens=embedding_max_tokens,
        timeout=embedding_timeout,
    )
    collection = client.get_collection(
        name=collection_name,
        embedding_function=embedding_function,
    )
    collection_count = int(collection.count())
    candidate_count = max(min(max(limit * 30, len(sql_docs), 1), max(collection_count, 1)), 1)
    result = collection.query(
        query_texts=[query],
        n_results=candidate_count,
        include=["documents", "metadatas", "distances"],
    )
    hits: list[dict[str, Any]] = []
    ids = result.get("ids", [[]])[0]
    distances = result.get("distances", [[]])[0]
    for vector_id, distance in zip(ids, distances):
        vector_id = str(vector_id)
        if vector_id not in allowed_doc_ids:
            continue
        row = dict(docs_by_id[vector_id])
        row["score"] = 1.0 - float(distance) if distance is not None else None
        hits.append(row)
        if len(hits) >= max(limit * 3, limit):
            break
    return hits


def _rank_reference_vector_sql_docs(docs: list[dict[str, Any]], query: str) -> list[dict[str, Any]]:
    query_tokens = set(_tokens(query))
    ranked = []
    for row in docs:
        tokens = set(_tokens(row.get("content") or ""))
        overlap = len(tokens.intersection(query_tokens))
        namespace_weight = {
            VECTOR_NAMESPACE_SOURCE_LOCAL_ENTITIES: 0.35,
            VECTOR_NAMESPACE_ATOMIC_FACTS: 0.34,
            VECTOR_NAMESPACE_SOURCE_LOCAL_RELATIONSHIPS: 0.28,
            VECTOR_NAMESPACE_ENTITY_PROPERTIES: 0.24,
            VECTOR_NAMESPACE_ENTITY_CLUSTERS: 0.18,
            VECTOR_NAMESPACE_PSEUDO_RELATIONSHIPS: 0.16,
            VECTOR_NAMESPACE_CHUNKS: 0.1,
        }.get(str(row.get("vector_namespace") or ""), 0.0)
        payload = dict(row)
        payload["score"] = round(overlap + namespace_weight, 4)
        ranked.append(payload)
    return sorted(
        ranked,
        key=lambda item: (
            float(item.get("score") or 0.0),
            str(item.get("vector_namespace") or ""),
            str(item.get("vector_id") or ""),
        ),
        reverse=True,
    )


def _reference_text_similarity(query: str, text: Any) -> float:
    query_tokens = set(_tokens(query))
    text_tokens = set(_tokens(text))
    if not query_tokens or not text_tokens:
        return 0.0
    overlap = len(query_tokens.intersection(text_tokens))
    coverage = overlap / max(len(query_tokens), 1)
    density = overlap / max(len(text_tokens), 1)
    return round(overlap + coverage + density, 6)


def _bound_reference_atomic_facts_for_entities(
    db_path: str | Path,
    *,
    query_text: str,
    entities: list[dict[str, Any]],
    source_doc_ids: set[str],
    source_scope_ids: set[str],
    top_k: int,
    direct_scores: dict[str, float],
) -> list[dict[str, Any]]:
    if top_k <= 0 or not entities:
        return []
    entity_ids = _dedupe_ordered(entity.get("source_local_entity_id") for entity in entities)
    cluster_ids = _dedupe_ordered(entity.get("cluster_id") for entity in entities)
    clauses = []
    params: list[str] = []
    if entity_ids:
        clauses.append(f"source_local_entity_id IN ({', '.join('?' for _ in entity_ids)})")
        params.extend(entity_ids)
    if cluster_ids:
        clauses.append(f"cluster_id IN ({', '.join('?' for _ in cluster_ids)})")
        params.extend(cluster_ids)
    if not clauses:
        return []
    with _connect(db_path) as conn:
        rows = [
            _atomic_fact_payload(row)
            for row in conn.execute(
                f"SELECT * FROM reference_atomic_facts WHERE {' OR '.join(f'({clause})' for clause in clauses)}",
                params,
            )
        ]
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    entity_ids_set = set(entity_ids)
    for fact in rows:
        if source_doc_ids and str(fact.get("source_doc_id") or "") not in source_doc_ids:
            continue
        if source_scope_ids and str(fact.get("source_scope_id") or "") not in source_scope_ids:
            continue
        group_key = _reference_fact_binding_group(fact, preferred_entity_ids=entity_ids_set)
        if not group_key:
            continue
        scored = dict(fact)
        fact_id = str(scored.get("fact_id") or "")
        scored["score"] = round(max(direct_scores.get(fact_id, 0.0), _reference_text_similarity(query_text, scored.get("statement"))), 6)
        scored["binding"] = {
            "strategy": "entity_fact_top_k_by_query_similarity",
            "bound_to": group_key,
            "top_k": top_k,
        }
        grouped[group_key].append(scored)
    return [
        item
        for group_items in grouped.values()
        for item in sorted(
            group_items,
            key=lambda row: (float(row.get("score") or 0.0), str(row.get("fact_id") or "")),
            reverse=True,
        )[:top_k]
    ]


def _bound_reference_properties_for_entities(
    db_path: str | Path,
    *,
    query_text: str,
    entities: list[dict[str, Any]],
    source_doc_ids: set[str],
    source_scope_ids: set[str],
    top_k: int,
    direct_scores: dict[str, float],
) -> list[dict[str, Any]]:
    if top_k <= 0 or not entities:
        return []
    entity_ids = _dedupe_ordered(entity.get("source_local_entity_id") for entity in entities)
    cluster_ids = _dedupe_ordered(entity.get("cluster_id") for entity in entities)
    clauses = []
    params: list[str] = []
    if entity_ids:
        clauses.append(f"source_local_entity_id IN ({', '.join('?' for _ in entity_ids)})")
        params.extend(entity_ids)
    if cluster_ids:
        clauses.append(f"cluster_id IN ({', '.join('?' for _ in cluster_ids)})")
        params.extend(cluster_ids)
    if not clauses:
        return []
    with _connect(db_path) as conn:
        rows = [
            _entity_property_payload(row)
            for row in conn.execute(
                f"SELECT * FROM reference_entity_properties WHERE {' OR '.join(f'({clause})' for clause in clauses)}",
                params,
            )
        ]
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    entity_ids_set = set(entity_ids)
    for prop in rows:
        if not _is_displayable_entity_property(prop):
            continue
        if source_doc_ids and str(prop.get("source_doc_id") or "") not in source_doc_ids:
            continue
        if source_scope_ids and str(prop.get("source_scope_id") or "") not in source_scope_ids:
            continue
        group_key = _reference_fact_binding_group(prop, preferred_entity_ids=entity_ids_set)
        if not group_key:
            continue
        scored = dict(prop)
        property_id = str(scored.get("property_id") or "")
        text = "\n".join(
            str(value or "")
            for value in (scored.get("property_name"), scored.get("property_value"), scored.get("statement"))
            if str(value or "").strip()
        )
        scored["score"] = round(max(direct_scores.get(property_id, 0.0), _reference_text_similarity(query_text, text)), 6)
        scored["binding"] = {
            "strategy": "entity_property_top_k_by_query_similarity",
            "bound_to": group_key,
            "top_k": top_k,
        }
        grouped[group_key].append(scored)
    return [
        item
        for group_items in grouped.values()
        for item in sorted(
            group_items,
            key=lambda row: (float(row.get("score") or 0.0), str(row.get("property_id") or "")),
            reverse=True,
        )[:top_k]
    ]


def _reference_fact_binding_group(item: dict[str, Any], *, preferred_entity_ids: set[str]) -> str:
    source_local_entity_id = str(item.get("source_local_entity_id") or "").strip()
    if source_local_entity_id and (not preferred_entity_ids or source_local_entity_id in preferred_entity_ids):
        return f"entity:{source_local_entity_id}"
    cluster_id = str(item.get("cluster_id") or "").strip()
    if cluster_id:
        return f"cluster:{cluster_id}"
    return f"entity:{source_local_entity_id}" if source_local_entity_id else ""


def _merge_reference_query_items(
    left: list[dict[str, Any]],
    right: list[dict[str, Any]],
    *,
    key_field: str,
) -> list[dict[str, Any]]:
    selected: dict[str, dict[str, Any]] = {}
    for item in left + right:
        key = str(item.get(key_field) or "")
        if not key:
            continue
        current = selected.get(key)
        if current is None or float(item.get("score") or 0.0) > float(current.get("score") or 0.0):
            selected[key] = item
    return sorted(
        selected.values(),
        key=lambda row: (float(row.get("score") or 0.0), str(row.get(key_field) or "")),
        reverse=True,
    )


def _limit_reference_items_per_binding_group(
    items: list[dict[str, Any]],
    *,
    top_k: int,
    fallback_limit: int,
) -> list[dict[str, Any]]:
    if top_k <= 0:
        return []
    grouped_counts: dict[str, int] = defaultdict(int)
    selected = []
    for item in sorted(
        items,
        key=lambda row: (float(row.get("score") or 0.0), str(row.get("fact_id") or row.get("property_id") or "")),
        reverse=True,
    ):
        binding = item.get("binding") or {}
        group_key = str(binding.get("bound_to") or _reference_fact_binding_group(item, preferred_entity_ids=set()))
        if group_key:
            if grouped_counts[group_key] >= top_k:
                continue
            grouped_counts[group_key] += 1
        selected.append(item)
        if len(selected) >= max(int(fallback_limit or 0), 0):
            break
    return selected


def _dedupe_vector_hits(rows: list[dict[str, Any]], *, key_field: str) -> list[dict[str, Any]]:
    selected: dict[str, dict[str, Any]] = {}
    for row in rows:
        key = str(row.get(key_field) or row.get("vector_id") or "")
        if not key:
            continue
        current = selected.get(key)
        if current is None or float(row.get("score") or 0.0) > float(current.get("score") or 0.0):
            selected[key] = row
    return sorted(
        selected.values(),
        key=lambda item: (float(item.get("score") or 0.0), str(item.get("vector_id") or "")),
        reverse=True,
    )


def _vector_document_matches_sources(
    row: dict[str, Any],
    *,
    source_doc_ids: set[str] | None,
    source_scope_ids: set[str] | None,
) -> bool:
    if source_doc_ids:
        row_doc_ids = set(_vector_document_source_doc_ids(row))
        if not row_doc_ids.intersection(source_doc_ids):
            return False
    if source_scope_ids:
        row_scope_ids = set(_vector_document_source_scope_ids(row))
        if not row_scope_ids.intersection(source_scope_ids):
            return False
    return True


def _vector_document_source_doc_ids(row: dict[str, Any]) -> list[str]:
    metadata = row.get("metadata") if isinstance(row.get("metadata"), dict) else {}
    values = []
    if row.get("source_doc_id"):
        values.append(str(row.get("source_doc_id")))
    for key in ("source_doc_ids", "source_docs"):
        metadata_value = metadata.get(key)
        if isinstance(metadata_value, list):
            for item in metadata_value:
                if isinstance(item, dict):
                    values.append(str(item.get("doc_id") or item.get("source_doc_id") or ""))
                else:
                    values.append(str(item or ""))
    return _dedupe_ordered(value for value in values if value)


def _vector_document_source_scope_ids(row: dict[str, Any]) -> list[str]:
    metadata = row.get("metadata") if isinstance(row.get("metadata"), dict) else {}
    values = []
    if row.get("source_scope_id"):
        values.append(str(row.get("source_scope_id")))
    metadata_scope_id = metadata.get("source_scope_id")
    if metadata_scope_id:
        values.append(str(metadata_scope_id))
    metadata_scope_ids = metadata.get("source_scope_ids")
    if isinstance(metadata_scope_ids, list):
        values.extend(str(item or "") for item in metadata_scope_ids)
    return _dedupe_ordered(value for value in values if value)


def _empty_reference_context() -> dict[str, Any]:
    return {
        "author_reference_context": [],
        "character_reference_knowledge": [],
        "style_reference_context": [],
        "timeline_reference_claims": [],
    }


def _reference_context_counts(context: dict[str, Any]) -> dict[str, int]:
    return {
        "author_reference_context": len(context.get("author_reference_context") or []),
        "character_reference_knowledge": len(context.get("character_reference_knowledge") or []),
        "style_reference_context": len(context.get("style_reference_context") or []),
        "timeline_reference_claims": len(context.get("timeline_reference_claims") or []),
    }


def _apply_reference_schema(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        PRAGMA journal_mode=WAL;

        CREATE TABLE IF NOT EXISTS reference_metadata (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_full_docs (
            doc_id TEXT PRIMARY KEY,
            source_scope_id TEXT,
            source_scope_name TEXT,
            source_path TEXT,
            relative_path TEXT,
            file_name TEXT,
            format TEXT,
            record_index INTEGER,
            title TEXT,
            content TEXT NOT NULL,
            content_sha256 TEXT,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_doc_status (
            doc_id TEXT PRIMARY KEY,
            status TEXT NOT NULL,
            chunks_count INTEGER NOT NULL DEFAULT 0,
            chunks_list_json TEXT NOT NULL,
            content_summary TEXT,
            content_length INTEGER NOT NULL DEFAULT 0,
            file_path TEXT,
            created_at TEXT,
            updated_at TEXT,
            metadata_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_llm_response_cache (
            cache_id TEXT PRIMARY KEY,
            chunk_id TEXT NOT NULL,
            doc_id TEXT,
            status TEXT,
            provider TEXT,
            model TEXT,
            raw_text TEXT NOT NULL,
            parsed_json TEXT NOT NULL,
            trace_json TEXT NOT NULL,
            created_at TEXT
        );

        CREATE TABLE IF NOT EXISTS reference_text_chunks (
            chunk_id TEXT PRIMARY KEY,
            full_doc_id TEXT NOT NULL,
            source_scope_id TEXT,
            chunk_order_index INTEGER NOT NULL,
            tokens INTEGER,
            content TEXT NOT NULL,
            file_path TEXT,
            source_path TEXT,
            relative_path TEXT,
            title TEXT,
            heading TEXT,
            content_sha256 TEXT,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_extracted_entities (
            extraction_id TEXT PRIMARY KEY,
            entity_name TEXT NOT NULL,
            entity_type TEXT NOT NULL,
            description TEXT NOT NULL,
            source_id TEXT NOT NULL,
            source_doc_id TEXT,
            source_scope_id TEXT,
            file_path TEXT,
            timestamp INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_extracted_relationships (
            extraction_id TEXT PRIMARY KEY,
            src_id TEXT NOT NULL,
            tgt_id TEXT NOT NULL,
            keywords TEXT,
            description TEXT NOT NULL,
            weight REAL NOT NULL DEFAULT 1.0,
            source_id TEXT NOT NULL,
            source_doc_id TEXT,
            source_scope_id TEXT,
            file_path TEXT,
            timestamp INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_source_scopes (
            source_scope_id TEXT PRIMARY KEY,
            source_scope_name TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            source_paths_json TEXT NOT NULL,
            created_at INTEGER,
            metadata_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_source_local_entities (
            source_local_entity_id TEXT PRIMARY KEY,
            source_scope_id TEXT NOT NULL,
            entity_name TEXT NOT NULL,
            canonical_entity_name TEXT NOT NULL,
            cluster_id TEXT NOT NULL,
            entity_type TEXT NOT NULL,
            description TEXT NOT NULL,
            facts_json TEXT NOT NULL,
            attributes_json TEXT NOT NULL,
            timeline_claims_json TEXT NOT NULL,
            source_id TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            file_path TEXT,
            created_at INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_source_local_relationships (
            source_local_relation_id TEXT PRIMARY KEY,
            source_scope_id TEXT NOT NULL,
            relation_id TEXT NOT NULL,
            pseudo_relation_id TEXT NOT NULL,
            src_id TEXT NOT NULL,
            tgt_id TEXT NOT NULL,
            src_cluster_id TEXT NOT NULL,
            tgt_cluster_id TEXT NOT NULL,
            keywords TEXT,
            description TEXT NOT NULL,
            facts_json TEXT NOT NULL,
            weight REAL NOT NULL DEFAULT 1.0,
            source_id TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            file_path TEXT,
            created_at INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_entity_clusters (
            cluster_id TEXT PRIMARY KEY,
            canonical_display_name TEXT NOT NULL,
            aliases_json TEXT NOT NULL,
            member_source_local_entity_ids_json TEXT NOT NULL,
            source_scope_ids_json TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            global_pseudo_degree INTEGER NOT NULL DEFAULT 0,
            evidence_degree INTEGER NOT NULL DEFAULT 0,
            source_coverage INTEGER NOT NULL DEFAULT 0,
            source_separation_note TEXT,
            summary TEXT,
            created_at INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_entity_cluster_members (
            cluster_id TEXT NOT NULL,
            source_local_entity_id TEXT NOT NULL,
            source_scope_id TEXT NOT NULL,
            entity_name TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            PRIMARY KEY(cluster_id, source_local_entity_id)
        );

        CREATE TABLE IF NOT EXISTS reference_pseudo_relationships (
            pseudo_relation_id TEXT PRIMARY KEY,
            src_cluster_id TEXT NOT NULL,
            tgt_cluster_id TEXT NOT NULL,
            src_display_name TEXT NOT NULL,
            tgt_display_name TEXT NOT NULL,
            member_source_local_relation_ids_json TEXT NOT NULL,
            source_scope_ids_json TEXT NOT NULL,
            source_doc_ids_json TEXT NOT NULL,
            weight REAL NOT NULL DEFAULT 1.0,
            description TEXT,
            keywords TEXT,
            created_at INTEGER,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_atomic_facts (
            fact_id TEXT PRIMARY KEY,
            source_scope_id TEXT NOT NULL,
            source_local_entity_id TEXT,
            source_local_relation_id TEXT,
            cluster_id TEXT,
            subject TEXT NOT NULL,
            predicate TEXT NOT NULL,
            object TEXT,
            statement TEXT NOT NULL,
            fact_type TEXT NOT NULL,
            confidence REAL NOT NULL DEFAULT 0.7,
            source_chunk_id TEXT,
            source_doc_id TEXT,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_entity_properties (
            property_id TEXT PRIMARY KEY,
            source_scope_id TEXT NOT NULL,
            source_local_entity_id TEXT NOT NULL,
            cluster_id TEXT NOT NULL,
            entity_name TEXT NOT NULL,
            property_name TEXT NOT NULL,
            property_value TEXT NOT NULL,
            statement TEXT NOT NULL,
            confidence REAL NOT NULL DEFAULT 0.7,
            source_chunk_id TEXT,
            source_doc_id TEXT,
            raw_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_vector_documents (
            vector_id TEXT PRIMARY KEY,
            vector_namespace TEXT NOT NULL,
            source_id TEXT NOT NULL,
            source_doc_id TEXT,
            source_scope_id TEXT,
            chunk_id TEXT,
            entity_name TEXT,
            source_local_entity_id TEXT,
            cluster_id TEXT,
            relation_id TEXT,
            source_local_relation_id TEXT,
            pseudo_relation_id TEXT,
            fact_id TEXT,
            property_id TEXT,
            src_id TEXT,
            tgt_id TEXT,
            content TEXT NOT NULL,
            metadata_json TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS reference_full_entities (
            doc_id TEXT PRIMARY KEY,
            entity_names_json TEXT NOT NULL,
            count INTEGER NOT NULL DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS reference_full_relations (
            doc_id TEXT PRIMARY KEY,
            relation_pairs_json TEXT NOT NULL,
            count INTEGER NOT NULL DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS reference_entity_chunks (
            entity_name TEXT PRIMARY KEY,
            chunk_ids_json TEXT NOT NULL,
            count INTEGER NOT NULL DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS reference_relation_chunks (
            relation_id TEXT PRIMARY KEY,
            chunk_ids_json TEXT NOT NULL,
            count INTEGER NOT NULL DEFAULT 0
        );

        CREATE INDEX IF NOT EXISTS idx_reference_text_chunks_doc
            ON reference_text_chunks(full_doc_id, chunk_order_index);
        CREATE INDEX IF NOT EXISTS idx_reference_text_chunks_scope
            ON reference_text_chunks(source_scope_id, full_doc_id);
        CREATE INDEX IF NOT EXISTS idx_reference_llm_response_cache_chunk
            ON reference_llm_response_cache(chunk_id, doc_id);
        CREATE INDEX IF NOT EXISTS idx_reference_extracted_entities_source
            ON reference_extracted_entities(source_doc_id, source_id, entity_name);
        CREATE INDEX IF NOT EXISTS idx_reference_extracted_entities_scope
            ON reference_extracted_entities(source_scope_id, entity_name);
        CREATE INDEX IF NOT EXISTS idx_reference_extracted_relationships_source
            ON reference_extracted_relationships(source_doc_id, source_id, src_id, tgt_id);
        CREATE INDEX IF NOT EXISTS idx_reference_extracted_relationships_scope
            ON reference_extracted_relationships(source_scope_id, src_id, tgt_id);
        CREATE INDEX IF NOT EXISTS idx_reference_vector_documents_namespace
            ON reference_vector_documents(vector_namespace, source_doc_id);
        CREATE INDEX IF NOT EXISTS idx_reference_vector_documents_scope
            ON reference_vector_documents(vector_namespace, source_scope_id);
        CREATE INDEX IF NOT EXISTS idx_reference_vector_documents_entity
            ON reference_vector_documents(entity_name);
        CREATE INDEX IF NOT EXISTS idx_reference_vector_documents_cluster
            ON reference_vector_documents(cluster_id);
        CREATE INDEX IF NOT EXISTS idx_reference_vector_documents_relation
            ON reference_vector_documents(relation_id);
        CREATE INDEX IF NOT EXISTS idx_reference_source_local_entities_scope
            ON reference_source_local_entities(source_scope_id, canonical_entity_name);
        CREATE INDEX IF NOT EXISTS idx_reference_source_local_entities_cluster
            ON reference_source_local_entities(cluster_id);
        CREATE INDEX IF NOT EXISTS idx_reference_source_local_relationships_scope
            ON reference_source_local_relationships(source_scope_id, src_id, tgt_id);
        CREATE INDEX IF NOT EXISTS idx_reference_source_local_relationships_pseudo
            ON reference_source_local_relationships(pseudo_relation_id);
        CREATE INDEX IF NOT EXISTS idx_reference_atomic_facts_subject
            ON reference_atomic_facts(subject, source_scope_id);
        CREATE INDEX IF NOT EXISTS idx_reference_entity_properties_entity
            ON reference_entity_properties(entity_name, source_scope_id);
        """
    )
    _ensure_reference_schema_columns(conn)


def _ensure_reference_schema_columns(conn: sqlite3.Connection) -> None:
    columns = {
        "reference_full_docs": {
            "source_scope_id": "TEXT",
            "source_scope_name": "TEXT",
        },
        "reference_text_chunks": {
            "source_scope_id": "TEXT",
        },
        "reference_extracted_entities": {
            "source_scope_id": "TEXT",
        },
        "reference_extracted_relationships": {
            "source_scope_id": "TEXT",
        },
        "reference_vector_documents": {
            "source_scope_id": "TEXT",
            "source_local_entity_id": "TEXT",
            "cluster_id": "TEXT",
            "source_local_relation_id": "TEXT",
            "pseudo_relation_id": "TEXT",
            "fact_id": "TEXT",
            "property_id": "TEXT",
        },
        "reference_entity_clusters": {
            "source_separation_note": "TEXT",
        },
    }
    for table, table_columns in columns.items():
        existing = {str(row["name"]) for row in conn.execute(f"PRAGMA table_info({table})")}
        for column, column_type in table_columns.items():
            if column not in existing:
                conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {column_type}")


def _insert_reference_full_doc(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_full_docs
        (doc_id, source_scope_id, source_scope_name, source_path, relative_path, file_name, format, record_index, title,
         content, content_sha256, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["doc_id"],
            row.get("source_scope_id") or _default_source_scope_id(row["doc_id"]),
            row.get("source_scope_name") or row.get("title") or row.get("file_name") or row["doc_id"],
            row.get("source_path"),
            row.get("relative_path"),
            row.get("file_name"),
            row.get("format"),
            int(row.get("record_index") or 0),
            row.get("title"),
            str(row.get("content") or ""),
            row.get("content_sha256"),
            json.dumps(row.get("raw") or row, ensure_ascii=False),
        ),
    )


def _insert_reference_doc_status(conn: sqlite3.Connection, doc: dict[str, Any], chunks: list[dict[str, Any]]) -> None:
    now = datetime.now(timezone.utc).isoformat()
    content = str(doc.get("content") or "")
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_doc_status
        (doc_id, status, chunks_count, chunks_list_json, content_summary,
         content_length, file_path, created_at, updated_at, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            doc["doc_id"],
            DOC_STATUS_PROCESSED,
            len(chunks),
            json.dumps([chunk.get("chunk_id") for chunk in chunks], ensure_ascii=False),
            content,
            len(content),
            doc.get("source_path"),
            now,
            now,
            json.dumps(
                {
                    "relative_path": doc.get("relative_path"),
                    "file_name": doc.get("file_name"),
                    "format": doc.get("format"),
                    "title": doc.get("title"),
                },
                ensure_ascii=False,
            ),
        ),
    )


def _insert_reference_llm_response_cache(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_llm_response_cache
        (cache_id, chunk_id, doc_id, status, provider, model, raw_text,
         parsed_json, trace_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["cache_id"],
            row["chunk_id"],
            row.get("doc_id"),
            row.get("status"),
            row.get("provider"),
            row.get("model"),
            row.get("raw_text") or "",
            json.dumps(row.get("parsed") or {}, ensure_ascii=False),
            json.dumps(row.get("trace") or {}, ensure_ascii=False),
            row.get("created_at") or datetime.now(timezone.utc).isoformat(),
        ),
    )


def _insert_reference_text_chunk(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    content = str(row.get("content") or "")
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_text_chunks
        (chunk_id, full_doc_id, source_scope_id, chunk_order_index, tokens, content, file_path,
         source_path, relative_path, title, heading, content_sha256, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["chunk_id"],
            row["doc_id"],
            row.get("source_scope_id") or _default_source_scope_id(row["doc_id"]),
            int(row.get("chunk_index") or row.get("chunk_order_index") or 0),
            len(_tokens(content)),
            content,
            row.get("source_path"),
            row.get("source_path"),
            row.get("relative_path"),
            row.get("title"),
            row.get("heading"),
            row.get("content_sha256"),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_extracted_entity(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_extracted_entities
        (extraction_id, entity_name, entity_type, description, source_id,
         source_doc_id, source_scope_id, file_path, timestamp, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["extraction_id"],
            row["entity_name"],
            row["entity_type"],
            row["description"],
            row["source_id"],
            row.get("source_doc_id"),
            row.get("source_scope_id") or _default_source_scope_id(row.get("source_doc_id")),
            row.get("file_path"),
            int(row.get("timestamp") or 0),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_extracted_relationship(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_extracted_relationships
        (extraction_id, src_id, tgt_id, keywords, description, weight,
         source_id, source_doc_id, source_scope_id, file_path, timestamp, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["extraction_id"],
            row["src_id"],
            row["tgt_id"],
            row.get("keywords") or "",
            row["description"],
            float(row.get("weight") or 1.0),
            row["source_id"],
            row.get("source_doc_id"),
            row.get("source_scope_id") or _default_source_scope_id(row.get("source_doc_id")),
            row.get("file_path"),
            int(row.get("timestamp") or 0),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_vector_document(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_vector_documents
        (vector_id, vector_namespace, source_id, source_doc_id, source_scope_id, chunk_id,
         entity_name, source_local_entity_id, cluster_id, relation_id, source_local_relation_id,
         pseudo_relation_id, fact_id, property_id, src_id, tgt_id, content, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["vector_id"],
            row["vector_namespace"],
            row["source_id"],
            row.get("source_doc_id"),
            row.get("source_scope_id"),
            row.get("chunk_id"),
            row.get("entity_name"),
            row.get("source_local_entity_id"),
            row.get("cluster_id"),
            row.get("relation_id"),
            row.get("source_local_relation_id"),
            row.get("pseudo_relation_id"),
            row.get("fact_id"),
            row.get("property_id"),
            row.get("src_id"),
            row.get("tgt_id"),
            row["content"],
            json.dumps(row.get("metadata") or {}, ensure_ascii=False),
        ),
    )


def _insert_reference_full_entities(conn: sqlite3.Connection, doc_id: str, entity_names: list[str]) -> None:
    values = _dedupe_ordered(entity_names)
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_full_entities
        (doc_id, entity_names_json, count)
        VALUES (?, ?, ?)
        """,
        (doc_id, json.dumps(values, ensure_ascii=False), len(values)),
    )


def _insert_reference_full_relations(conn: sqlite3.Connection, doc_id: str, relation_pairs: list[list[str]]) -> None:
    values = _dedupe_relation_pairs(relation_pairs)
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_full_relations
        (doc_id, relation_pairs_json, count)
        VALUES (?, ?, ?)
        """,
        (doc_id, json.dumps(values, ensure_ascii=False), len(values)),
    )


def _insert_reference_entity_chunks(conn: sqlite3.Connection, entity_name: str, chunk_ids: list[str]) -> None:
    values = _dedupe_ordered(chunk_ids)
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_entity_chunks
        (entity_name, chunk_ids_json, count)
        VALUES (?, ?, ?)
        """,
        (entity_name, json.dumps(values, ensure_ascii=False), len(values)),
    )


def _insert_reference_relation_chunks(conn: sqlite3.Connection, relation_id: str, chunk_ids: list[str]) -> None:
    values = _dedupe_ordered(chunk_ids)
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_relation_chunks
        (relation_id, chunk_ids_json, count)
        VALUES (?, ?, ?)
        """,
        (relation_id, json.dumps(values, ensure_ascii=False), len(values)),
    )


def _insert_reference_source_scope(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_source_scopes
        (source_scope_id, source_scope_name, source_doc_ids_json, source_paths_json,
         created_at, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            row["source_scope_id"],
            row["source_scope_name"],
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
            json.dumps(row.get("source_paths") or [], ensure_ascii=False),
            int(row.get("created_at") or int(time.time())),
            json.dumps(row.get("metadata") or {}, ensure_ascii=False),
        ),
    )


def _insert_reference_source_local_entity(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_source_local_entities
        (source_local_entity_id, source_scope_id, entity_name, canonical_entity_name,
         cluster_id, entity_type, description, facts_json, attributes_json,
         timeline_claims_json, source_id, source_doc_ids_json, file_path, created_at,
         raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["source_local_entity_id"],
            row["source_scope_id"],
            row["entity_name"],
            row["canonical_entity_name"],
            row["cluster_id"],
            row["entity_type"],
            row["description"],
            json.dumps(row.get("facts") or [], ensure_ascii=False),
            json.dumps(row.get("attributes") or [], ensure_ascii=False),
            json.dumps(row.get("timeline_claims") or [], ensure_ascii=False),
            row["source_id"],
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
            row.get("file_path") or "",
            int(row.get("created_at") or int(time.time())),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_source_local_relationship(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_source_local_relationships
        (source_local_relation_id, source_scope_id, relation_id, pseudo_relation_id,
         src_id, tgt_id, src_cluster_id, tgt_cluster_id, keywords, description,
         facts_json, weight, source_id, source_doc_ids_json, file_path, created_at,
         raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["source_local_relation_id"],
            row["source_scope_id"],
            row["relation_id"],
            row["pseudo_relation_id"],
            row["src_id"],
            row["tgt_id"],
            row["src_cluster_id"],
            row["tgt_cluster_id"],
            row.get("keywords") or "",
            row["description"],
            json.dumps(row.get("facts") or [], ensure_ascii=False),
            float(row.get("weight") or 1.0),
            row["source_id"],
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
            row.get("file_path") or "",
            int(row.get("created_at") or int(time.time())),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_entity_cluster(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_entity_clusters
        (cluster_id, canonical_display_name, aliases_json,
         member_source_local_entity_ids_json, source_scope_ids_json,
         source_doc_ids_json, global_pseudo_degree, evidence_degree,
         source_coverage, source_separation_note, summary, created_at, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["cluster_id"],
            row["canonical_display_name"],
            json.dumps(row.get("aliases") or [], ensure_ascii=False),
            json.dumps(row.get("member_source_local_entity_ids") or [], ensure_ascii=False),
            json.dumps(row.get("source_scope_ids") or [], ensure_ascii=False),
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
            int(row.get("global_pseudo_degree") or 0),
            int(row.get("evidence_degree") or 0),
            int(row.get("source_coverage") or 0),
            row.get("source_separation_note") or "",
            row.get("summary") or "",
            int(row.get("created_at") or int(time.time())),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_entity_cluster_member(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_entity_cluster_members
        (cluster_id, source_local_entity_id, source_scope_id, entity_name,
         source_doc_ids_json)
        VALUES (?, ?, ?, ?, ?)
        """,
        (
            row["cluster_id"],
            row["source_local_entity_id"],
            row["source_scope_id"],
            row["entity_name"],
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
        ),
    )


def _insert_reference_pseudo_relationship(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_pseudo_relationships
        (pseudo_relation_id, src_cluster_id, tgt_cluster_id, src_display_name,
         tgt_display_name, member_source_local_relation_ids_json,
         source_scope_ids_json, source_doc_ids_json, weight, description,
         keywords, created_at, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["pseudo_relation_id"],
            row["src_cluster_id"],
            row["tgt_cluster_id"],
            row["src_display_name"],
            row["tgt_display_name"],
            json.dumps(row.get("member_source_local_relation_ids") or [], ensure_ascii=False),
            json.dumps(row.get("source_scope_ids") or [], ensure_ascii=False),
            json.dumps(row.get("source_doc_ids") or [], ensure_ascii=False),
            float(row.get("weight") or 1.0),
            row.get("description") or "",
            row.get("keywords") or "",
            int(row.get("created_at") or int(time.time())),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_atomic_fact(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_atomic_facts
        (fact_id, source_scope_id, source_local_entity_id, source_local_relation_id,
         cluster_id, subject, predicate, object, statement, fact_type, confidence,
         source_chunk_id, source_doc_id, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["fact_id"],
            row["source_scope_id"],
            row.get("source_local_entity_id"),
            row.get("source_local_relation_id"),
            row.get("cluster_id"),
            row["subject"],
            row["predicate"],
            row.get("object") or "",
            row["statement"],
            row["fact_type"],
            float(row.get("confidence") or 0.7),
            row.get("source_chunk_id"),
            row.get("source_doc_id"),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _insert_reference_entity_property(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO reference_entity_properties
        (property_id, source_scope_id, source_local_entity_id, cluster_id,
         entity_name, property_name, property_value, statement, confidence,
         source_chunk_id, source_doc_id, raw_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            row["property_id"],
            row["source_scope_id"],
            row["source_local_entity_id"],
            row["cluster_id"],
            row["entity_name"],
            row["property_name"],
            row["property_value"],
            row["statement"],
            float(row.get("confidence") or 0.7),
            row.get("source_chunk_id"),
            row.get("source_doc_id"),
            json.dumps(row, ensure_ascii=False),
        ),
    )


def _reference_input_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        if input_path.suffix.lower() not in SUPPORTED_REFERENCE_EXTENSIONS:
            raise ValueError(f"Unsupported reference file extension: {input_path.suffix}")
        return [input_path]
    if not input_path.is_dir():
        raise FileNotFoundError(f"Reference input path not found: {input_path}")
    files = [
        path
        for path in sorted(input_path.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_REFERENCE_EXTENSIONS
    ]
    if not files:
        raise ValueError(f"No supported reference files found under: {input_path}")
    return files


def _load_reference_file(path: Path, *, root: Path, file_index: int) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        text = path.read_text(encoding="utf-8")
        return [_raw_reference_document(path, root=root, file_index=file_index, record_index=1, title=path.stem, content=text)]
    if suffix == ".docx":
        payload = _load_docx_reference_payload(path)
        return [
            _raw_reference_document(
                path,
                root=root,
                file_index=file_index,
                record_index=1,
                title=str(payload.get("title") or path.stem),
                content=str(payload.get("content") or ""),
                raw_payload=payload,
            )
        ]
    if suffix == ".jsonl":
        records = []
        with path.open("r", encoding="utf-8") as handle:
            for index, line in enumerate(handle, start=1):
                line = line.strip()
                if not line:
                    continue
                payload = json.loads(line)
                records.append(
                    _raw_reference_document(
                        path,
                        root=root,
                        file_index=file_index,
                        record_index=index,
                        title=str(payload.get("heading") or payload.get("title") or payload.get("record_id") or path.stem)
                        if isinstance(payload, dict)
                        else f"{path.stem}_{index}",
                        content=_json_record_to_text(payload),
                        raw_payload=payload,
                    )
                )
        return records
    if suffix == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        docs = _flatten_json_reference_payload(payload, path=path, root=root, file_index=file_index)
        return docs or [
            _raw_reference_document(
                path,
                root=root,
                file_index=file_index,
                record_index=1,
                title=path.stem,
                content=_json_record_to_text(payload),
                raw_payload=payload,
            )
        ]
    raise ValueError(f"Unsupported reference file extension: {suffix}")


def _raw_reference_document(
    path: Path,
    *,
    root: Path,
    file_index: int,
    record_index: int,
    title: str,
    content: str,
    raw_payload: Any | None = None,
) -> dict[str, Any]:
    rel = path.relative_to(root) if path.is_relative_to(root) else path
    doc_key = f"{rel.as_posix()}:{record_index}"
    digest = hashlib.sha1(doc_key.encode("utf-8")).hexdigest()
    content_text = str(content or "")
    doc_id = f"ref_doc_{file_index:04d}_{record_index:04d}_{digest}"
    source_scope_id = _default_source_scope_id(doc_id)
    source_scope_name = str(title or path.stem).strip()
    return {
        "doc_id": doc_id,
        "source_scope_id": source_scope_id,
        "source_scope_name": source_scope_name,
        "source_path": str(path),
        "relative_path": rel.as_posix(),
        "file_name": path.name,
        "format": path.suffix.lower().lstrip("."),
        "record_index": record_index,
        "title": str(title or path.stem).strip(),
        "content": content_text,
        "content_sha256": hashlib.sha256(content_text.encode("utf-8")).hexdigest(),
        "raw": raw_payload if raw_payload is not None else {},
    }


def _load_docx_reference_payload(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared, this is an environment guard.
        raise RuntimeError("python-docx is required to ingest .docx reference files") from exc

    document = Document(str(path))
    lines: list[str] = []
    paragraph_count = 0
    non_empty_paragraph_count = 0
    for paragraph in document.paragraphs:
        text = _clean_docx_text(paragraph.text)
        if not text:
            continue
        paragraph_count += 1
        non_empty_paragraph_count += 1
        style_name = str(getattr(paragraph.style, "name", "") or "")
        rendered = _render_docx_paragraph(text, style_name)
        if rendered:
            lines.append(rendered)
    table_count = 0
    table_row_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        table_count += 1
        lines.append(f"\n## Table {table_index}")
        for row in table.rows:
            cells = [_clean_docx_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            table_row_count += 1
            lines.append("- " + " | ".join(cells))
    content = "\n".join(lines).strip()
    title = _docx_title(document, path=path, fallback_content=content)
    return {
        "title": title,
        "content": content,
        "paragraph_count": paragraph_count,
        "non_empty_paragraph_count": non_empty_paragraph_count,
        "table_count": table_count,
        "table_row_count": table_row_count,
        "parser": "python-docx",
    }


def _render_docx_paragraph(text: str, style_name: str) -> str:
    normalized_style = style_name.strip().lower()
    if normalized_style == "title":
        return f"# {text}"
    heading_match = re.match(r"heading\s+([1-6])", normalized_style)
    if heading_match:
        level = min(int(heading_match.group(1)), 6)
        return f"{'#' * level} {text}"
    if "list" in normalized_style or normalized_style.startswith("bullet"):
        return f"- {text}"
    return text


def _docx_title(document: Any, *, path: Path, fallback_content: str) -> str:
    core_title = str(getattr(document.core_properties, "title", "") or "").strip()
    if core_title:
        return core_title
    for line in fallback_content.splitlines():
        text = line.strip()
        if not text:
            continue
        text = re.sub(r"^\s{0,3}#{1,6}\s+", "", text).strip()
        text = re.sub(r"^-\s+", "", text).strip()
        if text:
            return text
    return path.stem


def _clean_docx_text(value: Any) -> str:
    text = str(value or "").replace("\xa0", " ").strip()
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _flatten_json_reference_payload(payload: Any, *, path: Path, root: Path, file_index: int) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    if isinstance(payload, dict) and isinstance(payload.get("documents"), list):
        record_index = 0
        for doc in payload["documents"]:
            if not isinstance(doc, dict):
                continue
            doc_title = str(doc.get("doc_id") or doc.get("title") or path.stem)
            chunks = doc.get("chunks")
            if isinstance(chunks, list):
                for chunk in chunks:
                    if not isinstance(chunk, dict):
                        continue
                    record_index += 1
                    title = str(chunk.get("heading") or chunk.get("chunk_id") or doc_title)
                    records.append(
                        _raw_reference_document(
                            path,
                            root=root,
                            file_index=file_index,
                            record_index=record_index,
                            title=title,
                            content=_json_record_to_text({**doc, **chunk, "parent_doc": doc_title}),
                            raw_payload={"document": doc, "chunk": chunk},
                        )
                    )
            else:
                record_index += 1
                records.append(
                    _raw_reference_document(
                        path,
                        root=root,
                        file_index=file_index,
                        record_index=record_index,
                        title=doc_title,
                        content=_json_record_to_text(doc),
                        raw_payload=doc,
                    )
                )
        return records
    if isinstance(payload, list):
        return [
            _raw_reference_document(
                path,
                root=root,
                file_index=file_index,
                record_index=index,
                title=str(item.get("title") or item.get("heading") or item.get("record_id") or f"{path.stem}_{index}")
                if isinstance(item, dict)
                else f"{path.stem}_{index}",
                content=_json_record_to_text(item),
                raw_payload=item,
            )
            for index, item in enumerate(payload, start=1)
        ]
    return []


def _json_record_to_text(payload: Any) -> str:
    if isinstance(payload, dict):
        parts = []
        for key, value in payload.items():
            if key in {"raw", "documents", "chunks"}:
                continue
            if isinstance(value, (dict, list)):
                rendered = json.dumps(value, ensure_ascii=False)
            else:
                rendered = str(value)
            if rendered and rendered not in {"None", ""}:
                parts.append(f"{key}: {rendered}")
        return "\n".join(parts)
    if isinstance(payload, list):
        return "\n".join(_json_record_to_text(item) for item in payload)
    return str(payload or "")


def _chunk_reference_document(raw_doc: dict[str, Any], *, max_chunk_chars: int) -> list[dict[str, Any]]:
    content = str(raw_doc.get("content") or "")
    sections = _markdown_like_sections(content) if raw_doc.get("format") in {"docx", "md", "markdown", "txt"} else []
    if not sections:
        sections = [{"heading": raw_doc.get("title") or "", "content": content, "section_index": 1}]

    chunks: list[dict[str, Any]] = []
    chunk_index = 0
    for section in sections:
        for part_index, text in enumerate(_split_reference_text_by_boundaries(section["content"], target_chars=max_chunk_chars), start=1):
            if not str(text or "").strip():
                continue
            chunk_index += 1
            chunk_id = f"{raw_doc['doc_id']}_chunk_{chunk_index:04d}"
            chunks.append(
                {
                    "chunk_id": chunk_id,
                    "doc_id": raw_doc["doc_id"],
                    "source_scope_id": raw_doc.get("source_scope_id") or _default_source_scope_id(raw_doc["doc_id"]),
                    "source_scope_name": raw_doc.get("source_scope_name") or raw_doc.get("title") or "",
                    "source_path": raw_doc.get("source_path"),
                    "relative_path": raw_doc.get("relative_path"),
                    "format": raw_doc.get("format"),
                    "title": raw_doc.get("title"),
                    "heading": section.get("heading") or raw_doc.get("title") or "",
                    "section_index": section.get("section_index"),
                    "chunk_index": chunk_index,
                    "section_part_index": part_index,
                    "content": text,
                    "content_sha256": hashlib.sha256(str(text).encode("utf-8")).hexdigest(),
                    "parent_content_sha256": raw_doc.get("content_sha256"),
                    "char_count": len(str(text)),
                    "raw_document": {
                        "doc_id": raw_doc.get("doc_id"),
                        "source_scope_id": raw_doc.get("source_scope_id"),
                        "source_scope_name": raw_doc.get("source_scope_name"),
                        "source_path": raw_doc.get("source_path"),
                        "relative_path": raw_doc.get("relative_path"),
                        "title": raw_doc.get("title"),
                    },
                }
            )
    return chunks


def _markdown_like_sections(content: str) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current_heading = ""
    current_heading_line = ""
    current_lines: list[str] = []
    heading_stack: list[str] = []
    section_index = 0
    for line in str(content or "").splitlines():
        heading_match = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$", line)
        if heading_match:
            if _section_has_body(current_lines, current_heading_line):
                section_index += 1
                sections.append(
                    {
                        "heading": current_heading,
                        "content": "\n".join(current_lines).strip(),
                        "section_index": section_index,
                    }
                )
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            heading_stack = heading_stack[: max(level - 1, 0)]
            heading_stack.append(heading_text)
            current_heading = " / ".join(heading_stack)
            current_heading_line = line
            current_lines = [line]
        else:
            current_lines.append(line)
    if _section_has_body(current_lines, current_heading_line):
        section_index += 1
        sections.append({"heading": current_heading, "content": "\n".join(current_lines).strip(), "section_index": section_index})
    return [section for section in sections if str(section.get("content") or "").strip()]


def _section_has_body(lines: list[str], heading_line: str) -> bool:
    body_lines = [line.strip() for line in lines if line.strip()]
    if not body_lines:
        return False
    if heading_line and len(body_lines) == 1 and body_lines[0] == heading_line.strip():
        return False
    return True


def _split_reference_text_by_boundaries(text: str, *, target_chars: int) -> list[str]:
    value = str(text or "").strip()
    if len(value) <= target_chars:
        return [value] if value else []
    paragraphs = re.split(r"\n\s*\n", value)
    units: list[str] = []
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if len(paragraph) <= target_chars:
            units.append(paragraph)
            continue
        sentence_units = [
            sentence.strip()
            for sentence in re.split(r"(?<=[。！？!?；;])\s*|(?<=[.!?])\s+|\n+", paragraph)
            if sentence.strip()
        ]
        units.extend(sentence_units or [paragraph])

    parts: list[str] = []
    current = ""
    for unit in units:
        candidate = f"{current}\n\n{unit}".strip() if current else unit
        if len(candidate) <= target_chars or not current:
            current = candidate
        else:
            parts.append(current.strip())
            current = unit
    if current:
        parts.append(current.strip())
    return parts


def _reference_chunk_prompt_context(chunk: dict[str, Any]) -> dict[str, Any]:
    return {
        "chunk_id": chunk.get("chunk_id"),
        "doc_id": chunk.get("doc_id"),
        "source_scope_id": chunk.get("source_scope_id"),
        "source_scope_name": chunk.get("source_scope_name"),
        "source_path": chunk.get("source_path"),
        "relative_path": chunk.get("relative_path"),
        "format": chunk.get("format"),
        "title": chunk.get("title"),
        "heading": chunk.get("heading"),
        "content": chunk.get("content"),
    }


def _render_reference_kg_prompt(chunk: dict[str, Any], *, entity_types: tuple[str, ...]) -> str:
    tuple_delimiter = "<|#|>"
    completion_delimiter = "<|COMPLETE|>"
    entity_type_text = ", ".join(entity_types)
    return f"""---Role---
You are a Knowledge Graph Specialist responsible for extracting entities and relationships from external reference text.

---Instructions---
1. Identify clearly defined and meaningful entities in the input text.
2. Entity output has exactly 4 fields:
entity{tuple_delimiter}entity_name{tuple_delimiter}entity_type{tuple_delimiter}entity_description
3. Entity type must be one of: {entity_type_text}. If none apply, use concept.
4. Identify direct, meaningful relationships between extracted entities.
5. Relationship output has exactly 5 fields:
relation{tuple_delimiter}source_entity{tuple_delimiter}target_entity{tuple_delimiter}relationship_keywords{tuple_delimiter}relationship_description
6. Only output positive relationships explicitly supported by the current chunk. Never output "no relationship", "not mentioned", "unknown", or speculative negative relationships.
7. Treat relationships as undirected unless the text explicitly says otherwise. Do not output duplicates.
8. Base descriptions only on the current chunk. Do not import facts from other files or general knowledge.
9. Output all entities first, then relationships. Output only records, one per line.
10. End with {completion_delimiter}.

---Data to be Processed---
<Entity_types>
[{entity_type_text}]

<Reference Chunk>
{json.dumps(chunk, ensure_ascii=False, indent=2)}

<Output>
"""


def _reference_fact_property_prompt_context(job: dict[str, Any]) -> dict[str, Any]:
    return {
        "job_id": job.get("job_id"),
        "asset_type": job.get("asset_type"),
        "cluster_job_id": job.get("cluster_job_id"),
        "batch_index": job.get("batch_index"),
        "batch_count": job.get("batch_count"),
        "source_scope_id": job.get("source_scope_id"),
        "source_local_entity_id": job.get("source_local_entity_id"),
        "source_local_relation_id": job.get("source_local_relation_id"),
        "cluster_id": job.get("cluster_id"),
        "canonical_display_name": job.get("canonical_display_name"),
        "global_pseudo_degree": job.get("global_pseudo_degree"),
        "source_scope_ids": job.get("source_scope_ids") or [],
        "source_doc_ids": job.get("source_doc_ids") or [],
        "entity": job.get("entity"),
        "relation": job.get("relation"),
        "members": job.get("members") or [],
        "relationships": job.get("relationships") or [],
        "evidence_chunks": job.get("evidence_chunks") or [],
    }


def _render_reference_fact_property_prompt(job: dict[str, Any]) -> str:
    return f"""---Role---
You extract source-grounded entity properties and atomic facts from external reference evidence for a creative writing memory system.

---Instructions---
1. Use only the given evidence chunks. Do not import outside knowledge.
2. The job is centered on one pseudo-canonical entity cluster selected from the reference graph. The cluster is for retrieval focus only; keep every output grounded in the cited source chunk.
3. Extract compact atomic facts that are useful for later retrieval and reasoning.
4. Extract entity properties only when the evidence states a stable attribute, profile detail, role, preference, capability, trait, status, affiliation, location, object ownership/use, or temporal attribute.
5. Do not output trivial type labels such as "entity_type", "type", or "category".
6. Keep each fact/property source-local. If different source chunks or files conflict, output separate source-grounded rows rather than reconciling them.
7. Prefer facts/properties about the cluster entity and its graph-neighbor relationships. Do not extract unrelated background facts from evidence chunks.
8. Prefer concise natural-language Chinese statements when the evidence is Chinese; preserve proper names.
9. Output JSON only with this schema:
{{
  "atomic_facts": [
    {{
      "subject": "entity or relation source",
      "predicate": "short predicate",
      "object": "optional object",
      "fact": "one source-grounded sentence",
      "evidence": "short supporting quote or paraphrase",
      "source_chunk_id": "chunk id from evidence_chunks",
      "confidence": 0.0
    }}
  ],
  "entity_properties": [
    {{
      "entity": "entity name",
      "property": "property name",
      "value": "property value",
      "statement": "one source-grounded property sentence",
      "evidence": "short supporting quote or paraphrase",
      "source_chunk_id": "chunk id from evidence_chunks",
      "confidence": 0.0
    }}
  ]
}}

---Data---
<Reference Fact/Property Job>
{json.dumps(job, ensure_ascii=False, indent=2)}

<Output>
"""


def _parse_lightrag_kg_output(raw_text: str, chunk: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    chunk_id = str(chunk.get("chunk_id") or "")
    doc_id = str(chunk.get("doc_id") or "")
    source_scope_id = str(chunk.get("source_scope_id") or _default_source_scope_id(doc_id))
    file_path = str(chunk.get("source_path") or "unknown_source")
    timestamp = int(time.time())
    entities: list[dict[str, Any]] = []
    relationships: list[dict[str, Any]] = []
    records = _split_lightrag_records(raw_text)
    entity_index = 0
    relationship_index = 0
    for record in records:
        fields = [field.strip() for field in _normalize_lightrag_delimiters(record).split("<|#|>")]
        if not fields:
            continue
        kind = fields[0].strip().lower()
        if kind == "entity" and len(fields) == 4:
            parsed = _parse_lightrag_entity_fields(
                fields,
                chunk_id=chunk_id,
                doc_id=doc_id,
                source_scope_id=source_scope_id,
                file_path=file_path,
                timestamp=timestamp,
            )
            if parsed:
                entity_index += 1
                parsed["extraction_id"] = _stable_reference_id("ref_ent_ext", chunk_id, parsed["entity_name"], parsed["description"], entity_index)
                entities.append(parsed)
        elif kind in {"relation", "relationship"} and len(fields) == 5:
            parsed = _parse_lightrag_relation_fields(
                fields,
                chunk_id=chunk_id,
                doc_id=doc_id,
                source_scope_id=source_scope_id,
                file_path=file_path,
                timestamp=timestamp,
            )
            if parsed:
                relationship_index += 1
                parsed["extraction_id"] = _stable_reference_id("ref_rel_ext", chunk_id, parsed["src_id"], parsed["tgt_id"], parsed["description"], relationship_index)
                relationships.append(parsed)
    return entities, relationships


def _reference_entity_type_issue_count(entities: list[dict[str, Any]], *, issue: str | None = None) -> int:
    count = 0
    for entity in entities:
        for item in _reference_entity_type_issue_list(entity):
            if not issue or item.get("issue") == issue:
                count += 1
    return count


def _reference_entity_type_audit(entities: list[dict[str, Any]]) -> dict[str, Any]:
    multiple_type_entities = []
    unsupported_type_entities = []
    for entity in entities:
        issues = _reference_entity_type_issue_list(entity)
        if any(issue.get("issue") == "multiple_entity_types" for issue in issues):
            multiple_type_entities.append(entity)
        if any(issue.get("issue") == "unsupported_entity_type" for issue in issues):
            unsupported_type_entities.append(entity)
    source_local_conflicts = _reference_source_local_entity_type_conflicts(entities)
    return {
        "multiple_entity_type_count": len(multiple_type_entities),
        "unsupported_entity_type_count": len(unsupported_type_entities),
        "source_local_type_conflict_count": len(source_local_conflicts),
        "multiple_entity_type_examples": _reference_entity_type_issue_examples(multiple_type_entities),
        "unsupported_entity_type_examples": _reference_entity_type_issue_examples(unsupported_type_entities),
        "source_local_type_conflict_examples": source_local_conflicts[:20],
    }


def _reference_entity_type_issue_examples(entities: list[dict[str, Any]], *, limit: int = 20) -> list[dict[str, Any]]:
    examples = []
    for entity in entities[:limit]:
        raw_entity_type = entity.get("original_entity_type") or entity.get("entity_type")
        examples.append(
            {
                "entity_name": entity.get("entity_name"),
                "entity_type": entity.get("entity_type"),
                "original_entity_type": raw_entity_type,
                "entity_type_candidates": entity.get("entity_type_candidates") or _reference_entity_type_candidates(raw_entity_type),
                "source_chunk_id": entity.get("source_id"),
                "source_doc_id": entity.get("source_doc_id"),
                "source_scope_id": entity.get("source_scope_id"),
                "issues": _reference_entity_type_issue_list(entity),
            }
        )
    return examples


def _reference_entity_type_issue_list(entity: dict[str, Any]) -> list[dict[str, Any]]:
    issues = entity.get("entity_type_issues")
    if isinstance(issues, list) and issues:
        return [issue for issue in issues if isinstance(issue, dict)]
    raw_entity_type = entity.get("original_entity_type") or entity.get("entity_type")
    candidates = entity.get("entity_type_candidates")
    if not isinstance(candidates, list) or not candidates:
        candidates = _reference_entity_type_candidates(raw_entity_type)
    return _reference_entity_type_issues(raw_entity_type, [str(candidate) for candidate in candidates])


def _reference_source_local_entity_type_conflicts(entities: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for entity in entities:
        entity_name = str(entity.get("entity_name") or "").strip()
        source_scope_id = str(entity.get("source_scope_id") or _default_source_scope_id(entity.get("source_doc_id")))
        if entity_name and source_scope_id:
            grouped[(source_scope_id, entity_name)].append(entity)
    conflicts = []
    for (source_scope_id, entity_name), values in sorted(grouped.items()):
        normalized_types = _dedupe_ordered(value.get("entity_type") for value in values)
        original_types = _dedupe_ordered(value.get("original_entity_type") or value.get("entity_type") for value in values)
        if len(normalized_types) <= 1:
            continue
        conflicts.append(
            {
                "source_scope_id": source_scope_id,
                "entity_name": entity_name,
                "normalized_entity_types": normalized_types,
                "original_entity_types": original_types,
                "source_chunk_ids": _dedupe_ordered(value.get("source_id") for value in values),
                "chosen_entity_type": Counter(str(value.get("entity_type") or "concept") for value in values).most_common(1)[0][0],
            }
        )
    return conflicts


def _parse_reference_fact_property_output(raw_text: str, job: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    payload = _load_reference_json_object(raw_text)
    facts_payload = payload.get("atomic_facts") if isinstance(payload, dict) else []
    properties_payload = payload.get("entity_properties") if isinstance(payload, dict) else []
    if not isinstance(facts_payload, list):
        facts_payload = []
    if not isinstance(properties_payload, list):
        properties_payload = []
    valid_chunk_ids = {str(chunk.get("chunk_id") or "") for chunk in (job.get("evidence_chunks") or []) if str(chunk.get("chunk_id") or "").strip()}
    fallback_chunk_id = next(iter(valid_chunk_ids), "")
    fallback_doc_id = _source_doc_id_for_fact_property_chunk(job, fallback_chunk_id)
    atomic_facts = [
        fact
        for row in facts_payload
        for fact in [_normalize_reference_atomic_fact(row, job, valid_chunk_ids=valid_chunk_ids, fallback_chunk_id=fallback_chunk_id, fallback_doc_id=fallback_doc_id)]
        if fact
    ]
    entity_properties = [
        prop
        for row in properties_payload
        for prop in [_normalize_reference_entity_property(row, job, valid_chunk_ids=valid_chunk_ids, fallback_chunk_id=fallback_chunk_id, fallback_doc_id=fallback_doc_id)]
        if prop
    ]
    return atomic_facts, entity_properties


def _normalize_reference_atomic_fact(
    row: Any,
    job: dict[str, Any],
    *,
    valid_chunk_ids: set[str],
    fallback_chunk_id: str,
    fallback_doc_id: str,
) -> dict[str, Any] | None:
    if not isinstance(row, dict):
        return None
    subject = _clean_lightrag_text(row.get("subject") or _reference_fact_property_default_subject(job), remove_inner_quotes=True)
    predicate = _clean_lightrag_text(row.get("predicate") or "states", remove_inner_quotes=True)
    obj = _clean_lightrag_text(row.get("object") or row.get("target") or "", remove_inner_quotes=True)
    statement = _clean_lightrag_text(row.get("fact") or row.get("statement") or row.get("description") or row.get("evidence"))
    if not subject or not statement:
        return None
    source_chunk_id = _valid_reference_source_chunk_id(row.get("source_chunk_id"), valid_chunk_ids, fallback_chunk_id)
    source_doc_id = _source_doc_id_for_fact_property_chunk(job, source_chunk_id) or fallback_doc_id
    source_scope_id = _source_scope_id_for_fact_property_chunk(job, source_chunk_id) or str(job.get("source_scope_id") or "")
    source_local_entity_id = (
        _source_local_entity_id_for_fact_property_chunk(job, source_chunk_id)
        if job.get("asset_type") in {"entity", "entity_cluster"}
        else ""
    )
    source_local_relation_id = str(job.get("source_local_relation_id") or "") if job.get("asset_type") == "relationship" else ""
    cluster_id = str(job.get("cluster_id") or "") if source_local_entity_id or job.get("asset_type") == "entity_cluster" else ""
    return {
        "fact_id": _stable_reference_id("reffact", job.get("job_id"), source_chunk_id, subject, predicate, obj, statement),
        "source_scope_id": source_scope_id,
        "source_local_entity_id": source_local_entity_id,
        "source_local_relation_id": source_local_relation_id,
        "cluster_id": cluster_id,
        "subject": subject,
        "predicate": predicate or "states",
        "object": obj,
        "statement": statement,
        "fact_type": f"{job.get('asset_type') or 'reference'}_atomic_fact",
        "confidence": _reference_confidence(row.get("confidence"), default=0.72),
        "source_chunk_id": source_chunk_id,
        "source_doc_id": source_doc_id,
        "evidence": _clean_lightrag_text(row.get("evidence")),
    }


def _normalize_reference_entity_property(
    row: Any,
    job: dict[str, Any],
    *,
    valid_chunk_ids: set[str],
    fallback_chunk_id: str,
    fallback_doc_id: str,
) -> dict[str, Any] | None:
    if not isinstance(row, dict):
        return None
    if job.get("asset_type") not in {"entity", "entity_cluster"}:
        return None
    entity = job.get("entity") or {}
    entity_name = _clean_lightrag_text(row.get("entity") or entity.get("entity_name"), remove_inner_quotes=True)
    property_name = _clean_lightrag_text(row.get("property") or row.get("property_name") or row.get("name"), remove_inner_quotes=True)
    property_value = _clean_lightrag_text(row.get("value") or row.get("property_value") or row.get("description"))
    statement = _clean_lightrag_text(row.get("statement") or row.get("fact") or "")
    if not statement and entity_name and property_name and property_value:
        statement = f"{entity_name}.{property_name}: {property_value}"
    if not entity_name or not property_name or not property_value or not statement:
        return None
    if property_name.strip().lower() in {"entity_type", "type", "category"}:
        return None
    source_chunk_id = _valid_reference_source_chunk_id(row.get("source_chunk_id"), valid_chunk_ids, fallback_chunk_id)
    source_doc_id = _source_doc_id_for_fact_property_chunk(job, source_chunk_id) or fallback_doc_id
    source_scope_id = _source_scope_id_for_fact_property_chunk(job, source_chunk_id) or str(job.get("source_scope_id") or "")
    source_local_entity_id = _source_local_entity_id_for_fact_property_chunk(job, source_chunk_id)
    cluster_id = str(job.get("cluster_id") or "")
    return {
        "property_id": _stable_reference_id("refprop", source_local_entity_id, source_chunk_id, property_name, property_value, statement),
        "source_scope_id": source_scope_id,
        "source_local_entity_id": source_local_entity_id,
        "cluster_id": cluster_id,
        "entity_name": entity_name,
        "property_name": property_name,
        "property_value": property_value,
        "statement": statement,
        "confidence": _reference_confidence(row.get("confidence"), default=0.72),
        "source_chunk_id": source_chunk_id,
        "source_doc_id": source_doc_id,
        "evidence": _clean_lightrag_text(row.get("evidence")),
    }


def _load_reference_json_object(raw_text: str) -> dict[str, Any]:
    text = str(raw_text or "").strip()
    if not text:
        return {}
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1).strip()
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            return {}
        try:
            payload = json.loads(match.group(0))
        except json.JSONDecodeError:
            return {}
    return payload if isinstance(payload, dict) else {}


def _reference_fact_property_default_subject(job: dict[str, Any]) -> str:
    if job.get("asset_type") == "relationship":
        relation = job.get("relation") or {}
        src = str(relation.get("src_id") or "").strip()
        tgt = str(relation.get("tgt_id") or "").strip()
        return f"{src} - {tgt}".strip(" -")
    entity = job.get("entity") or {}
    return str(entity.get("entity_name") or "").strip()


def _valid_reference_source_chunk_id(value: Any, valid_chunk_ids: set[str], fallback_chunk_id: str) -> str:
    chunk_id = str(value or "").strip()
    return chunk_id if chunk_id and chunk_id in valid_chunk_ids else fallback_chunk_id


def _source_doc_id_for_fact_property_chunk(job: dict[str, Any], chunk_id: str) -> str:
    for chunk in job.get("evidence_chunks") or []:
        if str(chunk.get("chunk_id") or "") == str(chunk_id or ""):
            return str(chunk.get("doc_id") or "")
    return ""


def _source_scope_id_for_fact_property_chunk(job: dict[str, Any], chunk_id: str) -> str:
    for chunk in job.get("evidence_chunks") or []:
        if str(chunk.get("chunk_id") or "") == str(chunk_id or ""):
            return str(chunk.get("source_scope_id") or "")
    return ""


def _source_local_entity_id_for_fact_property_chunk(job: dict[str, Any], chunk_id: str) -> str:
    for member in job.get("members") or []:
        source_ids = _split_graph_field(member.get("source_id"))
        if str(chunk_id or "") in source_ids and member.get("source_local_entity_id"):
            return str(member.get("source_local_entity_id") or "")
    return str(job.get("source_local_entity_id") or "")


def _reference_confidence(value: Any, *, default: float) -> float:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        parsed = default
    return min(max(parsed, 0.0), 1.0)


def _split_lightrag_records(raw_text: str) -> list[str]:
    text = _normalize_lightrag_delimiters(str(raw_text or "")).replace("<|COMPLETE|>", "\n")
    records = []
    for line in text.splitlines():
        value = line.strip().strip("()")
        if not value:
            continue
        if value.lower().startswith(("entity<|#|>", "relation<|#|>", "relationship<|#|>")):
            records.append(value)
    return records


def _normalize_lightrag_delimiters(text: str) -> str:
    return re.sub(r"(?<!<)\|#\|>", "<|#|>", text)


def _parse_lightrag_entity_fields(
    fields: list[str],
    *,
    chunk_id: str,
    doc_id: str,
    source_scope_id: str,
    file_path: str,
    timestamp: int,
) -> dict[str, Any] | None:
    entity_name = _clean_lightrag_text(fields[1], remove_inner_quotes=True)
    raw_entity_type = _clean_lightrag_text(fields[2], remove_inner_quotes=True)
    entity_type_candidates = _reference_entity_type_candidates(raw_entity_type)
    primary_entity_type = entity_type_candidates[0] if entity_type_candidates else "concept"
    normalized_entity_type = normalize_entity_type(primary_entity_type)
    type_issues = _reference_entity_type_issues(raw_entity_type, entity_type_candidates)
    description = _clean_lightrag_text(fields[3])
    if not entity_name or not description:
        return None
    return {
        "entity_name": entity_name,
        "entity_type": normalized_entity_type,
        "original_entity_type": raw_entity_type,
        "entity_type_candidates": entity_type_candidates,
        "entity_type_issues": type_issues,
        "description": description,
        "source_id": chunk_id,
        "source_doc_id": doc_id,
        "source_scope_id": source_scope_id,
        "file_path": file_path,
        "timestamp": timestamp,
    }


def _reference_entity_type_candidates(raw_entity_type: Any) -> list[str]:
    text = str(raw_entity_type or "").strip().lower()
    if not text:
        return []
    parts = [
        part.strip().replace(" ", "_").replace("-", "_").strip("_")
        for part in re.split(r"[,，/、|;；]+", text)
        if part.strip()
    ]
    return _dedupe_ordered(parts)


def _reference_entity_type_issues(raw_entity_type: Any, candidates: list[str]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    if len(candidates) > 1:
        issues.append(
            {
                "issue": "multiple_entity_types",
                "raw_entity_type": str(raw_entity_type or ""),
                "candidates": candidates,
                "normalized_choice": normalize_entity_type(candidates[0]),
            }
        )
    unsupported = [candidate for candidate in candidates if not is_supported_entity_type_label(candidate)]
    if unsupported:
        issues.append(
            {
                "issue": "unsupported_entity_type",
                "raw_entity_type": str(raw_entity_type or ""),
                "unsupported": unsupported,
                "allowed_entity_types": sorted(ALLOWED_ENTITY_TYPES),
                "normalized_choice": normalize_entity_type(candidates[0] if candidates else "concept"),
            }
        )
    return issues


def _parse_lightrag_relation_fields(
    fields: list[str],
    *,
    chunk_id: str,
    doc_id: str,
    source_scope_id: str,
    file_path: str,
    timestamp: int,
) -> dict[str, Any] | None:
    source = _clean_lightrag_text(fields[1], remove_inner_quotes=True)
    target = _clean_lightrag_text(fields[2], remove_inner_quotes=True)
    keywords = _clean_lightrag_text(fields[3], remove_inner_quotes=True).replace("，", ",")
    description = _clean_lightrag_text(fields[4])
    if not source or not target or source == target or not description:
        return None
    if _is_negative_reference_relationship(description, keywords):
        return None
    return {
        "src_id": source,
        "tgt_id": target,
        "weight": 1.0,
        "description": description,
        "keywords": keywords,
        "source_id": chunk_id,
        "source_doc_id": doc_id,
        "source_scope_id": source_scope_id,
        "file_path": file_path,
        "timestamp": timestamp,
        "relation_id": _reference_relation_id(source, target),
    }


def _clean_lightrag_text(value: Any, *, remove_inner_quotes: bool = False) -> str:
    text = str(value or "").strip().strip("\ufeff")
    text = text.strip().strip('"').strip("'").strip()
    text = re.sub(r"\s+", " ", text)
    if remove_inner_quotes:
        text = text.replace('"', "").replace("'", "")
    return text


def _is_negative_reference_relationship(description: Any, keywords: Any = "") -> bool:
    text = f"{description or ''} {keywords or ''}".strip().lower()
    if not text:
        return True
    compact_text = re.sub(r"\s+", " ", text)
    return any(pattern in compact_text for pattern in NEGATIVE_RELATIONSHIP_PATTERNS)


def _reference_source_scopes(raw_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for doc in raw_docs:
        doc_id = str(doc.get("doc_id") or "")
        scope_id = str(doc.get("source_scope_id") or _default_source_scope_id(doc_id))
        if scope_id:
            grouped[scope_id].append(doc)
    scopes = []
    for scope_id, docs in sorted(grouped.items()):
        first = docs[0]
        scope_name = str(first.get("source_scope_name") or first.get("title") or first.get("file_name") or scope_id)
        scopes.append(
            {
                "source_scope_id": scope_id,
                "source_scope_name": scope_name,
                "source_doc_ids": _dedupe_ordered(doc.get("doc_id") for doc in docs),
                "source_paths": _dedupe_ordered(doc.get("source_path") or doc.get("relative_path") for doc in docs),
                "created_at": int(time.time()),
                "metadata": {
                    "default_policy": "one raw reference document per source_scope unless a future manifest groups documents",
                    "doc_count": len(docs),
                },
            }
        )
    return scopes


def _source_local_entities(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        entity_name = str(row.get("entity_name") or "").strip()
        source_scope_id = str(row.get("source_scope_id") or _default_source_scope_id(row.get("source_doc_id")))
        if entity_name and source_scope_id:
            grouped[(source_scope_id, entity_name)].append(row)
    entities = []
    for (source_scope_id, entity_name), values in sorted(grouped.items()):
        sorted_values = sorted(values, key=lambda item: (int(item.get("timestamp") or 0), str(item.get("source_id") or "")))
        descriptions = _dedupe_ordered(row.get("description") for row in sorted_values)
        entity_type_evidence = _reference_entity_type_evidence(sorted_values)
        type_resolution = _resolve_reference_entity_type(entity_type_evidence)
        entity_type_issues = [
            issue
            for row in sorted_values
            for issue in _reference_entity_type_issue_list(row)
            if isinstance(issue, dict)
        ]
        entity_type_issues.extend(type_resolution.get("entity_type_issues") or [])
        source_ids = _dedupe_ordered(row.get("source_id") for row in sorted_values)
        source_doc_ids = _dedupe_ordered(row.get("source_doc_id") for row in sorted_values)
        file_paths = _dedupe_ordered(row.get("file_path") for row in sorted_values)
        canonical_name = _canonical_reference_entity_name(entity_name)
        cluster_id = _reference_entity_cluster_id(canonical_name, type_resolution["entity_type"])
        source_local_entity_id = _reference_source_local_entity_id(source_scope_id, canonical_name)
        description = _merge_descriptions(f"{source_scope_id}:{entity_name}", descriptions)
        facts = [
            _entity_fact_from_description(
                source_local_entity_id=source_local_entity_id,
                source_scope_id=source_scope_id,
                cluster_id=cluster_id,
                entity_name=entity_name,
                description=str(row.get("description") or ""),
                source_chunk_id=str(row.get("source_id") or ""),
                source_doc_id=str(row.get("source_doc_id") or ""),
            )
            for row in sorted_values
            if str(row.get("description") or "").strip()
        ]
        attributes: list[dict[str, Any]] = []
        entities.append(
            {
                "source_local_entity_id": source_local_entity_id,
                "source_scope_id": source_scope_id,
                "entity_name": entity_name,
                "canonical_entity_name": canonical_name,
                "cluster_id": cluster_id,
                "entity_type": type_resolution["entity_type"],
                "entity_types": type_resolution["entity_types"],
                "original_entity_types": type_resolution["original_entity_types"],
                "entity_type_conflict": type_resolution["entity_type_conflict"],
                "entity_type_resolution": type_resolution["entity_type_resolution"],
                "entity_type_evidence": entity_type_evidence,
                "entity_type_issues": entity_type_issues,
                "description": description,
                "facts": facts,
                "attributes": attributes,
                "timeline_claims": [fact for fact in facts if _looks_temporal_text(fact.get("statement"))],
                "source_id": GRAPH_FIELD_SEP.join(source_ids),
                "source_doc_ids": source_doc_ids,
                "file_path": GRAPH_FIELD_SEP.join(file_paths),
                "created_at": int(time.time()),
                "description_fragments": descriptions,
            }
        )
    return entities


def _resolve_reference_source_local_entity_types(entities: list[dict[str, Any]]) -> list[dict[str, Any]]:
    resolved = []
    for entity in entities:
        rewritten = dict(entity)
        resolution = _resolve_reference_entity_type(entity.get("entity_type_evidence") or [])
        rewritten["entity_type"] = resolution["entity_type"]
        rewritten["entity_types"] = resolution["entity_types"]
        rewritten["original_entity_types"] = resolution["original_entity_types"]
        rewritten["entity_type_conflict"] = resolution["entity_type_conflict"]
        rewritten["entity_type_resolution"] = resolution["entity_type_resolution"]
        base_issues = [
            issue
            for issue in (entity.get("entity_type_issues") or [])
            if isinstance(issue, dict) and issue.get("issue") != "source_local_entity_type_conflict"
        ]
        rewritten["entity_type_issues"] = base_issues + list(resolution.get("entity_type_issues") or [])
        resolved.append(rewritten)
    return resolved


def _reference_entity_type_evidence(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    evidence = []
    for row in rows:
        raw_entity_type = row.get("original_entity_type") or row.get("entity_type") or "concept"
        candidates = row.get("entity_type_candidates")
        if not isinstance(candidates, list) or not candidates:
            candidates = _reference_entity_type_candidates(raw_entity_type)
        normalized_candidates = _dedupe_ordered(normalize_entity_type(candidate) for candidate in candidates)
        evidence.append(
            {
                "source_chunk_id": row.get("source_id"),
                "source_doc_id": row.get("source_doc_id"),
                "source_scope_id": row.get("source_scope_id"),
                "raw_entity_type": raw_entity_type,
                "entity_type_candidates": [str(candidate) for candidate in candidates],
                "normalized_entity_types": normalized_candidates or [normalize_entity_type(row.get("entity_type") or "concept")],
                "chosen_entity_type": normalize_entity_type(row.get("entity_type") or raw_entity_type or "concept"),
            }
        )
    return evidence


def _resolve_reference_entity_type(type_evidence: list[dict[str, Any]]) -> dict[str, Any]:
    normalized_values = _dedupe_ordered(
        entity_type
        for evidence in type_evidence
        for entity_type in (evidence.get("normalized_entity_types") or [])
    )
    original_values = _dedupe_ordered(evidence.get("raw_entity_type") for evidence in type_evidence)
    if not normalized_values:
        normalized_values = ["concept"]
    counts = Counter(
        str(entity_type or "concept")
        for evidence in type_evidence
        for entity_type in (evidence.get("normalized_entity_types") or ["concept"])
    )
    chosen = _choose_reference_resolved_entity_type(counts)
    conflict = len(normalized_values) > 1
    issues: list[dict[str, Any]] = []
    reason = "single_type"
    if conflict:
        reason = "resolved_type_conflict"
        issues.append(
            {
                "issue": "source_local_entity_type_conflict",
                "entity_types": normalized_values,
                "original_entity_types": original_values,
                "chosen_entity_type": chosen,
                "resolution_policy": "frequency_then_specificity",
            }
        )
    return {
        "entity_type": chosen,
        "entity_types": normalized_values,
        "original_entity_types": original_values,
        "entity_type_conflict": conflict,
        "entity_type_issues": issues,
        "entity_type_resolution": {
            "chosen_entity_type": chosen,
            "reason": reason,
            "type_counts": dict(counts),
            "policy": "frequency_then_specificity",
        },
    }


def _choose_reference_resolved_entity_type(counts: Counter[str]) -> str:
    if not counts:
        return "concept"
    specificity = {
        "character": 6,
        "organization": 5,
        "group": 4,
        "location": 4,
        "object": 3,
        "occasion": 3,
        "concept": 1,
    }
    return sorted(
        counts,
        key=lambda entity_type: (-counts[entity_type], -specificity.get(entity_type, 0), entity_type),
    )[0]


def _source_local_relationships(rows: list[dict[str, Any]], entities: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entity_by_scope_name = {
        (str(entity.get("source_scope_id") or ""), _canonical_reference_entity_name(entity.get("entity_name"))): entity
        for entity in entities
    }
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        if _is_negative_reference_relationship(row.get("description"), row.get("keywords")):
            continue
        source_scope_id = str(row.get("source_scope_id") or _default_source_scope_id(row.get("source_doc_id")))
        src, tgt = _sorted_relation_pair(row.get("src_id"), row.get("tgt_id"))
        if source_scope_id and src and tgt:
            grouped[(source_scope_id, _reference_relation_id(src, tgt))].append(row)
    relationships = []
    for (source_scope_id, relation_id), values in sorted(grouped.items()):
        sorted_values = sorted(values, key=lambda item: (int(item.get("timestamp") or 0), str(item.get("source_id") or "")))
        first = sorted_values[0]
        src, tgt = _sorted_relation_pair(first.get("src_id"), first.get("tgt_id"))
        src_entity = entity_by_scope_name.get((source_scope_id, _canonical_reference_entity_name(src)))
        tgt_entity = entity_by_scope_name.get((source_scope_id, _canonical_reference_entity_name(tgt)))
        src_cluster_id = str((src_entity or {}).get("cluster_id") or _reference_entity_cluster_id(src))
        tgt_cluster_id = str((tgt_entity or {}).get("cluster_id") or _reference_entity_cluster_id(tgt))
        pseudo_relation_id = _reference_pseudo_relation_id(src_cluster_id, tgt_cluster_id)
        descriptions = _dedupe_ordered(row.get("description") for row in sorted_values)
        source_ids = _dedupe_ordered(row.get("source_id") for row in sorted_values)
        source_doc_ids = _dedupe_ordered(row.get("source_doc_id") for row in sorted_values)
        file_paths = _dedupe_ordered(row.get("file_path") for row in sorted_values)
        keywords = _merge_keywords(row.get("keywords") for row in sorted_values)
        source_local_relation_id = _reference_source_local_relation_id(source_scope_id, relation_id)
        facts = [
            _relationship_fact_from_description(
                source_local_relation_id=source_local_relation_id,
                source_scope_id=source_scope_id,
                src=src,
                tgt=tgt,
                description=str(row.get("description") or ""),
                source_chunk_id=str(row.get("source_id") or ""),
                source_doc_id=str(row.get("source_doc_id") or ""),
            )
            for row in sorted_values
            if str(row.get("description") or "").strip()
        ]
        relationships.append(
            {
                "source_local_relation_id": source_local_relation_id,
                "source_scope_id": source_scope_id,
                "relation_id": relation_id,
                "pseudo_relation_id": pseudo_relation_id,
                "src_id": src,
                "tgt_id": tgt,
                "src_cluster_id": src_cluster_id,
                "tgt_cluster_id": tgt_cluster_id,
                "keywords": keywords,
                "description": _merge_descriptions(f"{source_scope_id}:{src}-{tgt}", descriptions),
                "facts": facts,
                "weight": max((float(row.get("weight") or 1.0) for row in sorted_values), default=1.0),
                "source_id": GRAPH_FIELD_SEP.join(source_ids),
                "source_doc_ids": source_doc_ids,
                "file_path": GRAPH_FIELD_SEP.join(file_paths),
                "created_at": int(time.time()),
                "description_fragments": descriptions,
            }
        )
    return relationships


def _reference_fact_property_jobs(
    source_local_entities: list[dict[str, Any]],
    source_local_relationships: list[dict[str, Any]],
    chunk_by_id: dict[str, dict[str, Any]],
    *,
    entity_clusters: list[dict[str, Any]],
    pseudo_relationships: list[dict[str, Any]],
    min_entity_degree: int = 2,
    max_evidence_chunks_per_job: int = 12,
) -> list[dict[str, Any]]:
    jobs: list[dict[str, Any]] = []
    entities_by_cluster: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for entity in source_local_entities:
        cluster_id = str(entity.get("cluster_id") or "")
        if cluster_id:
            entities_by_cluster[cluster_id].append(entity)
    relations_by_cluster: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for relation in source_local_relationships:
        for cluster_id in (str(relation.get("src_cluster_id") or ""), str(relation.get("tgt_cluster_id") or "")):
            if cluster_id:
                relations_by_cluster[cluster_id].append(relation)
    pseudo_by_cluster: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for pseudo_relation in pseudo_relationships:
        for cluster_id in (str(pseudo_relation.get("src_cluster_id") or ""), str(pseudo_relation.get("tgt_cluster_id") or "")):
            if cluster_id:
                pseudo_by_cluster[cluster_id].append(pseudo_relation)

    for cluster in sorted(entity_clusters, key=lambda item: str(item.get("canonical_display_name") or item.get("cluster_id") or "")):
        degree = int(cluster.get("global_pseudo_degree") or 0)
        if int(min_entity_degree or 0) > 0 and degree < int(min_entity_degree or 0):
            continue
        cluster_id = str(cluster.get("cluster_id") or "")
        members = sorted(
            entities_by_cluster.get(cluster_id, []),
            key=lambda item: (str(item.get("source_scope_id") or ""), str(item.get("entity_name") or "")),
        )
        relations = sorted(
            _dedupe_reference_records(relations_by_cluster.get(cluster_id, []), key_field="source_local_relation_id"),
            key=lambda item: (str(item.get("source_scope_id") or ""), str(item.get("relation_id") or "")),
        )
        chunk_ids = _reference_fact_property_cluster_chunk_ids(members, relations)
        evidence_chunks = [_reference_fact_property_evidence_chunk(chunk_by_id[chunk_id]) for chunk_id in chunk_ids if chunk_id in chunk_by_id]
        if not members or not evidence_chunks:
            continue
        jobs.extend(
            _reference_fact_property_cluster_jobs(
                cluster,
                members=members,
                relations=relations,
                pseudo_relations=pseudo_by_cluster.get(cluster_id, []),
                evidence_chunks=evidence_chunks,
                max_evidence_chunks_per_job=max_evidence_chunks_per_job,
            )
        )
    return jobs


def _reference_fact_property_cluster_jobs(
    cluster: dict[str, Any],
    *,
    members: list[dict[str, Any]],
    relations: list[dict[str, Any]],
    pseudo_relations: list[dict[str, Any]],
    evidence_chunks: list[dict[str, Any]],
    max_evidence_chunks_per_job: int,
) -> list[dict[str, Any]]:
    chunk_batches = _reference_fact_property_chunk_batches(
        evidence_chunks,
        max_evidence_chunks_per_job=max_evidence_chunks_per_job,
    )
    cluster_id = str(cluster.get("cluster_id") or "")
    cluster_job_id = _stable_reference_id("reffpclusterjob", cluster_id, ",".join(str(chunk.get("chunk_id") or "") for chunk in evidence_chunks))
    jobs = []
    for batch_index, batch_chunks in enumerate(chunk_batches, start=1):
        batch_chunk_ids = {str(chunk.get("chunk_id") or "") for chunk in batch_chunks if str(chunk.get("chunk_id") or "").strip()}
        batch_members = _reference_fact_property_batch_members(members, batch_chunk_ids)
        batch_relations = _reference_fact_property_batch_relations(relations, batch_chunk_ids)
        if not batch_members and not batch_relations:
            continue
        representative = (batch_members or members)[0]
        batch_pseudo_relations = _reference_fact_property_batch_pseudo_relations(pseudo_relations, batch_relations)
        job_id = _stable_reference_id(
            "reffpjob",
            cluster_job_id,
            batch_index,
            ",".join(str(chunk.get("chunk_id") or "") for chunk in batch_chunks),
        )
        jobs.append(
            {
                "job_id": job_id,
                "cluster_job_id": cluster_job_id,
                "batch_index": batch_index,
                "batch_count": len(chunk_batches),
                "asset_type": "entity_cluster",
                "source_scope_id": representative.get("source_scope_id"),
                "source_scope_ids": cluster.get("source_scope_ids") or [],
                "source_doc_ids": _dedupe_ordered(chunk.get("doc_id") for chunk in batch_chunks),
                "source_local_entity_id": representative.get("source_local_entity_id"),
                "source_local_relation_id": "",
                "cluster_id": cluster_id,
                "canonical_display_name": cluster.get("canonical_display_name"),
                "global_pseudo_degree": int(cluster.get("global_pseudo_degree") or 0),
                "entity": {
                    "entity_name": cluster.get("canonical_display_name"),
                    "canonical_entity_name": cluster.get("canonical_display_name"),
                    "entity_type": _reference_cluster_entity_type(batch_members or members),
                    "description": _merge_descriptions(
                        str(cluster.get("canonical_display_name") or cluster_id),
                        _reference_batch_entity_descriptions(batch_members or members, batch_chunk_ids),
                    ),
                },
                "relation": None,
                "members": [
                    {
                        "source_local_entity_id": member.get("source_local_entity_id"),
                        "source_scope_id": member.get("source_scope_id"),
                        "entity_name": member.get("entity_name"),
                        "canonical_entity_name": member.get("canonical_entity_name"),
                        "entity_type": member.get("entity_type"),
                        "description": _reference_batch_graph_descriptions(member, batch_chunk_ids),
                        "source_id": GRAPH_FIELD_SEP.join(_reference_graph_source_ids_in_batch(member, batch_chunk_ids)),
                        "source_doc_ids": _dedupe_ordered(chunk.get("doc_id") for chunk in batch_chunks),
                    }
                    for member in batch_members
                ],
                "relationships": [
                    {
                        "source_local_relation_id": relation.get("source_local_relation_id"),
                        "source_scope_id": relation.get("source_scope_id"),
                        "src_id": relation.get("src_id"),
                        "tgt_id": relation.get("tgt_id"),
                        "src_cluster_id": relation.get("src_cluster_id"),
                        "tgt_cluster_id": relation.get("tgt_cluster_id"),
                        "keywords": relation.get("keywords"),
                        "description": _reference_batch_graph_descriptions(relation, batch_chunk_ids),
                        "source_doc_ids": _dedupe_ordered(chunk.get("doc_id") for chunk in batch_chunks),
                    }
                    for relation in batch_relations
                ],
                "pseudo_relationships": [
                    {
                        "pseudo_relation_id": relation.get("pseudo_relation_id"),
                        "src_cluster_id": relation.get("src_cluster_id"),
                        "tgt_cluster_id": relation.get("tgt_cluster_id"),
                        "src_display_name": relation.get("src_display_name"),
                        "tgt_display_name": relation.get("tgt_display_name"),
                        "source_scope_ids": relation.get("source_scope_ids") or [],
                        "source_doc_ids": relation.get("source_doc_ids") or [],
                        "description": relation.get("description"),
                        "keywords": relation.get("keywords"),
                    }
                    for relation in batch_pseudo_relations
                ],
                "evidence_chunks": batch_chunks,
            }
        )
    return jobs


def _reference_fact_property_chunk_batches(
    evidence_chunks: list[dict[str, Any]],
    *,
    max_evidence_chunks_per_job: int,
) -> list[list[dict[str, Any]]]:
    max_chunks = max(1, int(max_evidence_chunks_per_job or 1))
    by_scope: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for chunk in evidence_chunks:
        scope_id = str(chunk.get("source_scope_id") or "")
        by_scope[scope_id].append(chunk)
    batches: list[list[dict[str, Any]]] = []
    for _scope_id, scope_chunks in sorted(by_scope.items(), key=lambda item: item[0]):
        sorted_chunks = sorted(
            scope_chunks,
            key=lambda item: (
                str(item.get("doc_id") or ""),
                int(item.get("chunk_index") or 0),
                str(item.get("chunk_id") or ""),
            ),
        )
        batches.extend(_batches(sorted_chunks, max_chunks))
    return batches


def _reference_fact_property_batch_members(
    members: list[dict[str, Any]],
    batch_chunk_ids: set[str],
) -> list[dict[str, Any]]:
    return [member for member in members if batch_chunk_ids.intersection(_split_graph_field(member.get("source_id")))]


def _reference_fact_property_batch_relations(
    relations: list[dict[str, Any]],
    batch_chunk_ids: set[str],
) -> list[dict[str, Any]]:
    return [relation for relation in relations if batch_chunk_ids.intersection(_split_graph_field(relation.get("source_id")))]


def _reference_fact_property_batch_pseudo_relations(
    pseudo_relations: list[dict[str, Any]],
    batch_relations: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    pseudo_relation_ids = {str(relation.get("pseudo_relation_id") or "") for relation in batch_relations if str(relation.get("pseudo_relation_id") or "").strip()}
    return [relation for relation in pseudo_relations if str(relation.get("pseudo_relation_id") or "") in pseudo_relation_ids]


def _reference_batch_entity_descriptions(members: list[dict[str, Any]], batch_chunk_ids: set[str]) -> list[str]:
    values = []
    for member in members:
        values.extend(_reference_batch_graph_description_values(member, batch_chunk_ids))
    return values


def _reference_batch_graph_descriptions(row: dict[str, Any], batch_chunk_ids: set[str]) -> str:
    return _merge_descriptions(str(row.get("entity_name") or row.get("source_local_relation_id") or ""), _reference_batch_graph_description_values(row, batch_chunk_ids))


def _reference_batch_graph_description_values(row: dict[str, Any], batch_chunk_ids: set[str]) -> list[str]:
    facts = row.get("facts") if isinstance(row.get("facts"), list) else []
    fact_values = [
        str(fact.get("statement") or "")
        for fact in facts
        if str(fact.get("source_chunk_id") or "") in batch_chunk_ids and str(fact.get("statement") or "").strip()
    ]
    if fact_values:
        return _dedupe_ordered(fact_values)
    fragments = row.get("description_fragments") if isinstance(row.get("description_fragments"), list) else []
    source_ids = _split_graph_field(row.get("source_id"))
    if fragments and source_ids:
        return [str(fragment) for source_id, fragment in zip(source_ids, fragments) if source_id in batch_chunk_ids and str(fragment or "").strip()]
    description = str(row.get("description") or "").strip()
    return [description] if description else []


def _reference_graph_source_ids_in_batch(row: dict[str, Any], batch_chunk_ids: set[str]) -> list[str]:
    return [source_id for source_id in _split_graph_field(row.get("source_id")) if source_id in batch_chunk_ids]


def _reference_fact_property_cluster_chunk_ids(
    members: list[dict[str, Any]],
    relations: list[dict[str, Any]],
) -> list[str]:
    values: list[str] = []
    for member in members:
        values.extend(_split_graph_field(member.get("source_id")))
    for relation in relations:
        values.extend(_split_graph_field(relation.get("source_id")))
    return _dedupe_ordered(values)


def _reference_cluster_entity_type(members: list[dict[str, Any]]) -> str:
    values = [str(member.get("entity_type") or "concept") for member in members if str(member.get("entity_type") or "").strip()]
    if not values:
        return "concept"
    return normalize_entity_type(Counter(values).most_common(1)[0][0])


def _reference_fact_property_evidence_chunk(chunk: dict[str, Any]) -> dict[str, Any]:
    return {
        "chunk_id": chunk.get("chunk_id"),
        "doc_id": chunk.get("doc_id"),
        "source_scope_id": chunk.get("source_scope_id"),
        "source_scope_name": chunk.get("source_scope_name"),
        "source_path": chunk.get("source_path"),
        "relative_path": chunk.get("relative_path"),
        "title": chunk.get("title"),
        "heading": chunk.get("heading"),
        "section_index": chunk.get("section_index"),
        "chunk_index": chunk.get("chunk_index"),
        "section_part_index": chunk.get("section_part_index"),
        "char_count": chunk.get("char_count"),
        "content": chunk.get("content"),
    }


def _reference_entity_clusters(
    source_local_entities: list[dict[str, Any]],
    source_local_relationships: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for entity in source_local_entities:
        cluster_id = str(entity.get("cluster_id") or "")
        if cluster_id:
            grouped[cluster_id].append(entity)
    neighbor_clusters: dict[str, set[str]] = defaultdict(set)
    for relation in source_local_relationships:
        src_cluster_id = str(relation.get("src_cluster_id") or "")
        tgt_cluster_id = str(relation.get("tgt_cluster_id") or "")
        if src_cluster_id and tgt_cluster_id and src_cluster_id != tgt_cluster_id:
            neighbor_clusters[src_cluster_id].add(tgt_cluster_id)
            neighbor_clusters[tgt_cluster_id].add(src_cluster_id)
    clusters = []
    members = []
    for cluster_id, values in sorted(grouped.items()):
        aliases = _dedupe_ordered(entity.get("entity_name") for entity in values)
        source_scope_ids = _dedupe_ordered(entity.get("source_scope_id") for entity in values)
        source_doc_ids = _dedupe_ordered(doc_id for entity in values for doc_id in (entity.get("source_doc_ids") or []))
        member_ids = _dedupe_ordered(entity.get("source_local_entity_id") for entity in values)
        display_name = aliases[0] if aliases else cluster_id
        source_separation_note = ""
        if len(source_scope_ids) > 1:
            source_separation_note = (
                f"{display_name} appears in {len(source_scope_ids)} external source scopes. "
                "Keep source-local facts separated unless one source claim is explicitly adopted into Working Truth."
            )
        clusters.append(
            {
                "cluster_id": cluster_id,
                "canonical_display_name": display_name,
                "aliases": aliases,
                "member_source_local_entity_ids": member_ids,
                "source_scope_ids": source_scope_ids,
                "source_doc_ids": source_doc_ids,
                "global_pseudo_degree": len(neighbor_clusters.get(cluster_id, set())),
                "evidence_degree": len(source_doc_ids),
                "source_coverage": len(source_scope_ids),
                "source_separation_note": source_separation_note,
                "summary": _reference_cluster_summary(display_name, source_scope_ids),
                "created_at": int(time.time()),
            }
        )
        for entity in values:
            members.append(
                {
                    "cluster_id": cluster_id,
                    "source_local_entity_id": entity.get("source_local_entity_id"),
                    "source_scope_id": entity.get("source_scope_id"),
                    "entity_name": entity.get("entity_name"),
                    "source_doc_ids": entity.get("source_doc_ids") or [],
                }
            )
    return clusters, members


def _canonicalize_reference_source_local_graph(
    source_local_entities: list[dict[str, Any]],
    source_local_relationships: list[dict[str, Any]],
    *,
    lexical_threshold: float,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    canonical_map = _reference_entity_canonical_map(source_local_entities, lexical_threshold=lexical_threshold)
    if not canonical_map:
        return source_local_entities, source_local_relationships

    scoped_name_to_cluster: dict[tuple[str, str], str] = {}
    name_to_clusters: dict[str, set[str]] = defaultdict(set)
    rewritten_entities = []
    for entity in source_local_entities:
        entity_type = normalize_entity_type(entity.get("entity_type") or "concept")
        canonical_name = _reference_canonical_name_for_map(
            entity.get("canonical_entity_name") or entity.get("entity_name"),
            canonical_map,
            entity_type=entity_type,
        )
        cluster_id = _reference_entity_cluster_id(canonical_name, entity_type)
        rewritten = dict(entity)
        rewritten["canonical_entity_name"] = canonical_name
        rewritten["cluster_id"] = cluster_id
        rewritten_entities.append(rewritten)
        for name in (entity.get("entity_name"), entity.get("canonical_entity_name"), canonical_name):
            normalized = _reference_entity_alias_norm(name)
            if normalized:
                scoped_name_to_cluster[(str(entity.get("source_scope_id") or ""), normalized)] = cluster_id
                name_to_clusters[normalized].add(cluster_id)

    rewritten_relationships = []
    for relation in source_local_relationships:
        relation_scope_id = str(relation.get("source_scope_id") or "")
        src_norm = _reference_entity_alias_norm(relation.get("src_id"))
        tgt_norm = _reference_entity_alias_norm(relation.get("tgt_id"))
        src_cluster_id = (
            scoped_name_to_cluster.get((relation_scope_id, src_norm))
            or _single_reference_cluster_for_alias(name_to_clusters, src_norm)
            or relation.get("src_cluster_id")
        )
        tgt_cluster_id = (
            scoped_name_to_cluster.get((relation_scope_id, tgt_norm))
            or _single_reference_cluster_for_alias(name_to_clusters, tgt_norm)
            or relation.get("tgt_cluster_id")
        )
        rewritten = dict(relation)
        rewritten["src_cluster_id"] = str(src_cluster_id or "")
        rewritten["tgt_cluster_id"] = str(tgt_cluster_id or "")
        rewritten["pseudo_relation_id"] = _reference_pseudo_relation_id(rewritten["src_cluster_id"], rewritten["tgt_cluster_id"])
        rewritten_relationships.append(rewritten)
    return rewritten_entities, rewritten_relationships


def _reference_entity_canonical_map(
    source_local_entities: list[dict[str, Any]],
    *,
    lexical_threshold: float,
) -> dict[tuple[str, str], str]:
    names_by_type: dict[str, list[str]] = defaultdict(list)
    for entity in source_local_entities:
        entity_type = normalize_entity_type(entity.get("entity_type") or "concept")
        names_by_type[entity_type].extend(
            str(value)
            for value in (entity.get("canonical_entity_name"), entity.get("entity_name"))
            if str(value or "").strip()
        )

    canonical_map: dict[tuple[str, str], str] = {}
    for entity_type, raw_names in sorted(names_by_type.items()):
        canonical_map.update(
            _reference_entity_canonical_map_for_type(
                entity_type,
                _dedupe_ordered(raw_names),
                lexical_threshold=lexical_threshold,
            )
        )
    return canonical_map


def _reference_entity_canonical_map_for_type(
    entity_type: str,
    names: list[str],
    *,
    lexical_threshold: float,
) -> dict[tuple[str, str], str]:
    if len(names) < 2:
        return {}
    parent = {name: name for name in names}

    def find(name: str) -> str:
        while parent[name] != name:
            parent[name] = parent[parent[name]]
            name = parent[name]
        return name

    def union(left: str, right: str) -> None:
        root_left = find(left)
        root_right = find(right)
        if root_left == root_right:
            return
        canonical = _choose_reference_canonical_name([root_left, root_right])
        other = root_right if canonical == root_left else root_left
        parent[other] = canonical
        parent[canonical] = canonical

    by_norm: dict[str, list[str]] = defaultdict(list)
    by_last_token: dict[str, list[str]] = defaultdict(list)
    for name in names:
        norm = _reference_entity_alias_norm(name)
        if not norm:
            continue
        by_norm[norm].append(name)
        tokens = norm.split()
        if tokens:
            by_last_token[tokens[-1]].append(name)

    for values in by_norm.values():
        if len(values) > 1:
            first = values[0]
            for other in values[1:]:
                union(first, other)

    for values in by_last_token.values():
        if not (1 < len(values) <= 8):
            continue
        for index, left in enumerate(values):
            for right in values[index + 1 :]:
                left_norm = _reference_entity_alias_norm(left)
                right_norm = _reference_entity_alias_norm(right)
                left_tokens = set(left_norm.split())
                right_tokens = set(right_norm.split())
                overlap = len(left_tokens & right_tokens) / max(1, min(len(left_tokens), len(right_tokens)))
                ratio = SequenceMatcher(None, left_norm, right_norm).ratio()
                if overlap >= 0.75 or ratio >= float(lexical_threshold or 0.88):
                    union(left, right)

    groups_by_root: dict[str, list[str]] = defaultdict(list)
    for name in names:
        root = find(name)
        groups_by_root[root].append(name)
    canonical_map = {}
    for values in groups_by_root.values():
        canonical = _choose_reference_canonical_name(values)
        for name in values:
            if canonical and canonical != name:
                canonical_map[(entity_type, _reference_entity_alias_norm(name))] = canonical
    return canonical_map


def _single_reference_cluster_for_alias(name_to_clusters: dict[str, set[str]], normalized_alias: str) -> str | None:
    cluster_ids = name_to_clusters.get(normalized_alias) or set()
    if len(cluster_ids) == 1:
        return next(iter(cluster_ids))
    return None


def _reference_canonical_name_for_map(value: Any, canonical_map: dict[tuple[str, str], str], *, entity_type: Any) -> str:
    text = _canonical_reference_entity_name(value)
    type_key = normalize_entity_type(entity_type or "concept")
    return canonical_map.get((type_key, _reference_entity_alias_norm(text)), text)


def _choose_reference_canonical_name(values: list[str]) -> str:
    candidates = [_canonical_reference_entity_name(value) for value in values if _canonical_reference_entity_name(value)]
    if not candidates:
        return ""
    return sorted(candidates, key=lambda item: (-len(item), item))[0]


def _reference_entity_alias_norm(value: Any) -> str:
    text = _canonical_reference_entity_name(value).lower()
    text = re.sub(r"[\s_\-·•:：,，.。()（）\\[\\]【】\"'“”‘’]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _reference_pseudo_relationships(
    source_local_relationships: list[dict[str, Any]],
    clusters: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    cluster_names = {str(cluster.get("cluster_id") or ""): str(cluster.get("canonical_display_name") or "") for cluster in clusters}
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for relation in source_local_relationships:
        pseudo_relation_id = str(relation.get("pseudo_relation_id") or "")
        if pseudo_relation_id:
            grouped[pseudo_relation_id].append(relation)
    pseudo_relationships = []
    for pseudo_relation_id, values in sorted(grouped.items()):
        first = values[0]
        src_cluster_id = str(first.get("src_cluster_id") or "")
        tgt_cluster_id = str(first.get("tgt_cluster_id") or "")
        descriptions = _dedupe_ordered(relation.get("description") for relation in values)
        pseudo_relationships.append(
            {
                "pseudo_relation_id": pseudo_relation_id,
                "src_cluster_id": src_cluster_id,
                "tgt_cluster_id": tgt_cluster_id,
                "src_display_name": cluster_names.get(src_cluster_id) or src_cluster_id,
                "tgt_display_name": cluster_names.get(tgt_cluster_id) or tgt_cluster_id,
                "member_source_local_relation_ids": _dedupe_ordered(relation.get("source_local_relation_id") for relation in values),
                "source_scope_ids": _dedupe_ordered(relation.get("source_scope_id") for relation in values),
                "source_doc_ids": _dedupe_ordered(doc_id for relation in values for doc_id in (relation.get("source_doc_ids") or [])),
                "weight": sum(float(relation.get("weight") or 1.0) for relation in values),
                "description": _merge_descriptions(pseudo_relation_id, descriptions),
                "keywords": _merge_keywords(relation.get("keywords") for relation in values),
                "created_at": int(time.time()),
            }
        )
    return pseudo_relationships


def _atomic_facts_from_source_local(
    source_local_entities: list[dict[str, Any]],
    source_local_relationships: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    facts = []
    for entity in source_local_entities:
        facts.extend(entity.get("facts") or [])
    for relation in source_local_relationships:
        facts.extend(relation.get("facts") or [])
    selected: dict[str, dict[str, Any]] = {}
    for fact in facts:
        fact_id = str(fact.get("fact_id") or "")
        if fact_id and fact_id not in selected:
            selected[fact_id] = fact
    return list(selected.values())


def _load_reference_fact_property_assets(facts_dir: Path | None) -> dict[str, list[dict[str, Any]]]:
    if facts_dir is None:
        return {"atomic_facts": [], "entity_properties": []}
    path = Path(facts_dir)
    atomic_facts_path = path / "atomic_facts.jsonl"
    entity_properties_path = path / "entity_properties.jsonl"
    return {
        "atomic_facts": _read_jsonl(atomic_facts_path) if atomic_facts_path.exists() else [],
        "entity_properties": _read_jsonl(entity_properties_path) if entity_properties_path.exists() else [],
    }


def _attach_reference_fact_property_assets(
    source_local_entities: list[dict[str, Any]],
    source_local_relationships: list[dict[str, Any]],
    assets: dict[str, list[dict[str, Any]]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    facts_by_entity: dict[str, list[dict[str, Any]]] = defaultdict(list)
    facts_by_relation: dict[str, list[dict[str, Any]]] = defaultdict(list)
    properties_by_entity: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for fact in assets.get("atomic_facts") or []:
        entity_id = str(fact.get("source_local_entity_id") or "")
        relation_id = str(fact.get("source_local_relation_id") or "")
        if entity_id:
            facts_by_entity[entity_id].append(fact)
        if relation_id:
            facts_by_relation[relation_id].append(fact)
    for prop in assets.get("entity_properties") or []:
        entity_id = str(prop.get("source_local_entity_id") or "")
        if entity_id:
            properties_by_entity[entity_id].append(prop)

    enriched_entities = []
    for entity in source_local_entities:
        entity_id = str(entity.get("source_local_entity_id") or "")
        enriched = dict(entity)
        if facts_by_entity.get(entity_id):
            enriched["facts"] = _merge_reference_atomic_facts(list(enriched.get("facts") or []), facts_by_entity[entity_id])
            enriched["timeline_claims"] = [fact for fact in enriched["facts"] if _looks_temporal_text(fact.get("statement"))]
        if properties_by_entity.get(entity_id):
            enriched["attributes"] = _merge_reference_entity_properties(list(enriched.get("attributes") or []), properties_by_entity[entity_id])
        enriched_entities.append(enriched)

    enriched_relationships = []
    for relation in source_local_relationships:
        relation_id = str(relation.get("source_local_relation_id") or "")
        enriched = dict(relation)
        if facts_by_relation.get(relation_id):
            enriched["facts"] = _merge_reference_atomic_facts(list(enriched.get("facts") or []), facts_by_relation[relation_id])
        enriched_relationships.append(enriched)
    return enriched_entities, enriched_relationships


def _merge_reference_atomic_facts(base: list[dict[str, Any]], extracted: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _dedupe_reference_records([*base, *extracted], key_field="fact_id")


def _merge_reference_entity_properties(base: list[dict[str, Any]], extracted: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _dedupe_reference_records([*base, *extracted], key_field="property_id")


def _dedupe_reference_records(records: list[dict[str, Any]], *, key_field: str) -> list[dict[str, Any]]:
    selected: dict[str, dict[str, Any]] = {}
    anonymous_index = 0
    for record in records:
        if not isinstance(record, dict):
            continue
        key = str(record.get(key_field) or "").strip()
        if not key:
            anonymous_index += 1
            key = f"anonymous_{anonymous_index}_{_hash_id(json.dumps(record, ensure_ascii=False, sort_keys=True))}"
        if key not in selected:
            selected[key] = record
    return list(selected.values())


def _entity_properties_from_source_local(source_local_entities: list[dict[str, Any]]) -> list[dict[str, Any]]:
    properties = []
    for entity in source_local_entities:
        properties.extend(entity.get("attributes") or [])
    selected: dict[str, dict[str, Any]] = {}
    for prop in properties:
        property_id = str(prop.get("property_id") or "")
        if property_id and property_id not in selected:
            selected[property_id] = prop
    return list(selected.values())


def _is_displayable_entity_property(prop: dict[str, Any]) -> bool:
    property_name = str(prop.get("property_name") or "").strip().lower()
    if property_name in {"entity_type", "type", "category"}:
        return False
    return bool(str(prop.get("statement") or "").strip())


def _full_entities_by_doc(rows: list[dict[str, Any]]) -> dict[str, list[str]]:
    grouped: dict[str, list[str]] = defaultdict(list)
    for row in rows:
        doc_id = str(row.get("source_doc_id") or "")
        entity_name = str(row.get("entity_name") or "")
        if doc_id and entity_name:
            grouped[doc_id].append(entity_name)
    return {doc_id: _dedupe_ordered(values) for doc_id, values in grouped.items()}


def _full_relations_by_doc(rows: list[dict[str, Any]]) -> dict[str, list[list[str]]]:
    grouped: dict[str, list[list[str]]] = defaultdict(list)
    for row in rows:
        doc_id = str(row.get("source_doc_id") or "")
        if not doc_id:
            continue
        src, tgt = _sorted_relation_pair(row.get("src_id"), row.get("tgt_id"))
        if src and tgt:
            grouped[doc_id].append([src, tgt])
    return {doc_id: _dedupe_relation_pairs(values) for doc_id, values in grouped.items()}


def _chunks_by_entity(rows: list[dict[str, Any]]) -> dict[str, list[str]]:
    grouped: dict[str, list[str]] = defaultdict(list)
    for row in rows:
        entity_name = str(row.get("entity_name") or "")
        chunk_id = str(row.get("source_id") or "")
        if entity_name and chunk_id:
            grouped[entity_name].append(chunk_id)
    return {entity_name: _dedupe_ordered(values) for entity_name, values in grouped.items()}


def _chunks_by_relation(rows: list[dict[str, Any]]) -> dict[str, list[str]]:
    grouped: dict[str, list[str]] = defaultdict(list)
    for row in rows:
        relation_id = str(row.get("relation_id") or _reference_relation_id(row.get("src_id"), row.get("tgt_id")))
        chunk_id = str(row.get("source_id") or "")
        if relation_id and chunk_id:
            grouped[relation_id].append(chunk_id)
    return {relation_id: _dedupe_ordered(values) for relation_id, values in grouped.items()}


def _chunk_vector_document(chunk: dict[str, Any]) -> dict[str, Any]:
    chunk_id = str(chunk.get("chunk_id") or "")
    return {
        "vector_id": f"chunk-{_hash_id(chunk_id)}",
        "vector_namespace": VECTOR_NAMESPACE_CHUNKS,
        "source_id": chunk_id,
        "source_doc_id": chunk.get("doc_id"),
        "source_scope_id": chunk.get("source_scope_id") or _default_source_scope_id(chunk.get("doc_id")),
        "chunk_id": chunk_id,
        "content": str(chunk.get("content") or ""),
        "metadata": {
            "title": chunk.get("title"),
            "heading": chunk.get("heading"),
            "source_path": chunk.get("source_path"),
            "relative_path": chunk.get("relative_path"),
            "source_scope_id": chunk.get("source_scope_id"),
            "source_scope_name": chunk.get("source_scope_name"),
        },
    }


def _source_local_entity_vector_document(entity: dict[str, Any]) -> dict[str, Any]:
    entity_name = str(entity.get("entity_name") or "")
    source_local_entity_id = str(entity.get("source_local_entity_id") or "")
    facts_text = "\n".join(str(fact.get("statement") or "") for fact in (entity.get("facts") or []))
    attrs_text = "\n".join(
        f"{prop.get('property_name')}: {prop.get('property_value')}"
        for prop in (entity.get("attributes") or [])
    )
    content = "\n".join(
        part
        for part in (
            entity_name,
            str(entity.get("entity_type") or ""),
            str(entity.get("description") or ""),
            facts_text,
            attrs_text,
        )
        if part
    )
    source_doc_ids = entity.get("source_doc_ids") or []
    source_ids = _split_graph_field(entity.get("source_id"))
    return {
        "vector_id": f"slent-{_hash_id(source_local_entity_id)}",
        "vector_namespace": VECTOR_NAMESPACE_SOURCE_LOCAL_ENTITIES,
        "source_id": str(entity.get("source_id") or ""),
        "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
        "source_scope_id": entity.get("source_scope_id"),
        "chunk_id": source_ids[0] if source_ids else "",
        "entity_name": entity_name,
        "source_local_entity_id": source_local_entity_id,
        "cluster_id": entity.get("cluster_id"),
        "content": content,
        "metadata": {
            "entity_type": entity.get("entity_type"),
            "canonical_entity_name": entity.get("canonical_entity_name"),
            "source_ids": source_ids,
            "source_doc_ids": source_doc_ids,
            "file_path": entity.get("file_path"),
            "scope": "source_local",
        },
    }


def _entity_cluster_vector_document(cluster: dict[str, Any]) -> dict[str, Any]:
    cluster_id = str(cluster.get("cluster_id") or "")
    display_name = str(cluster.get("canonical_display_name") or cluster_id)
    aliases = cluster.get("aliases") or []
    content = "\n".join(
        part
        for part in (
            display_name,
            ", ".join(aliases),
            str(cluster.get("summary") or ""),
            str(cluster.get("source_separation_note") or ""),
        )
        if part
    )
    source_doc_ids = cluster.get("source_doc_ids") or []
    return {
        "vector_id": f"cluster-{_hash_id(cluster_id)}",
        "vector_namespace": VECTOR_NAMESPACE_ENTITY_CLUSTERS,
        "source_id": cluster_id,
        "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
        "source_scope_id": "",
        "cluster_id": cluster_id,
        "entity_name": display_name,
        "content": content,
        "metadata": {
            "aliases": aliases,
            "source_scope_ids": cluster.get("source_scope_ids") or [],
            "source_doc_ids": source_doc_ids,
            "global_pseudo_degree": cluster.get("global_pseudo_degree"),
            "evidence_degree": cluster.get("evidence_degree"),
            "scope": "pseudo_cluster",
        },
    }


def _source_local_relationship_vector_document(relation: dict[str, Any]) -> dict[str, Any]:
    source_local_relation_id = str(relation.get("source_local_relation_id") or "")
    src = str(relation.get("src_id") or "")
    tgt = str(relation.get("tgt_id") or "")
    facts_text = "\n".join(str(fact.get("statement") or "") for fact in (relation.get("facts") or []))
    content = "\n".join(
        part
        for part in (
            str(relation.get("keywords") or ""),
            src,
            tgt,
            str(relation.get("description") or ""),
            facts_text,
        )
        if part
    )
    source_doc_ids = relation.get("source_doc_ids") or []
    source_ids = _split_graph_field(relation.get("source_id"))
    return {
        "vector_id": f"slrel-{_hash_id(source_local_relation_id)}",
        "vector_namespace": VECTOR_NAMESPACE_SOURCE_LOCAL_RELATIONSHIPS,
        "source_id": str(relation.get("source_id") or ""),
        "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
        "source_scope_id": relation.get("source_scope_id"),
        "chunk_id": source_ids[0] if source_ids else "",
        "relation_id": relation.get("relation_id"),
        "source_local_relation_id": source_local_relation_id,
        "pseudo_relation_id": relation.get("pseudo_relation_id"),
        "src_id": src,
        "tgt_id": tgt,
        "content": content,
        "metadata": {
            "keywords": relation.get("keywords"),
            "source_ids": source_ids,
            "source_doc_ids": source_doc_ids,
            "file_path": relation.get("file_path"),
            "scope": "source_local",
        },
    }


def _pseudo_relationship_vector_document(relation: dict[str, Any]) -> dict[str, Any]:
    pseudo_relation_id = str(relation.get("pseudo_relation_id") or "")
    source_doc_ids = relation.get("source_doc_ids") or []
    content = "\n".join(
        part
        for part in (
            str(relation.get("keywords") or ""),
            str(relation.get("src_display_name") or ""),
            str(relation.get("tgt_display_name") or ""),
            str(relation.get("description") or ""),
        )
        if part
    )
    return {
        "vector_id": f"pseudorel-{_hash_id(pseudo_relation_id)}",
        "vector_namespace": VECTOR_NAMESPACE_PSEUDO_RELATIONSHIPS,
        "source_id": pseudo_relation_id,
        "source_doc_id": source_doc_ids[0] if source_doc_ids else "",
        "source_scope_id": "",
        "pseudo_relation_id": pseudo_relation_id,
        "src_id": relation.get("src_display_name"),
        "tgt_id": relation.get("tgt_display_name"),
        "content": content,
        "metadata": {
            "source_scope_ids": relation.get("source_scope_ids") or [],
            "source_doc_ids": source_doc_ids,
            "member_source_local_relation_ids": relation.get("member_source_local_relation_ids") or [],
            "scope": "pseudo_relationship",
        },
    }


def _atomic_fact_vector_document(fact: dict[str, Any]) -> dict[str, Any]:
    source_doc_id = str(fact.get("source_doc_id") or "")
    return {
        "vector_id": f"fact-{_hash_id(fact.get('fact_id'))}",
        "vector_namespace": VECTOR_NAMESPACE_ATOMIC_FACTS,
        "source_id": str(fact.get("source_chunk_id") or fact.get("fact_id") or ""),
        "source_doc_id": source_doc_id,
        "source_scope_id": fact.get("source_scope_id"),
        "chunk_id": fact.get("source_chunk_id"),
        "entity_name": fact.get("subject"),
        "source_local_entity_id": fact.get("source_local_entity_id"),
        "source_local_relation_id": fact.get("source_local_relation_id"),
        "cluster_id": fact.get("cluster_id"),
        "fact_id": fact.get("fact_id"),
        "content": str(fact.get("statement") or ""),
        "metadata": {
            "subject": fact.get("subject"),
            "predicate": fact.get("predicate"),
            "object": fact.get("object"),
            "fact_type": fact.get("fact_type"),
            "confidence": fact.get("confidence"),
            "scope": "source_local_claim",
        },
    }


def _entity_property_vector_document(prop: dict[str, Any]) -> dict[str, Any]:
    source_doc_id = str(prop.get("source_doc_id") or "")
    content = "\n".join(
        part
        for part in (
            str(prop.get("entity_name") or ""),
            str(prop.get("property_name") or ""),
            str(prop.get("property_value") or ""),
            str(prop.get("statement") or ""),
        )
        if part
    )
    return {
        "vector_id": f"prop-{_hash_id(prop.get('property_id'))}",
        "vector_namespace": VECTOR_NAMESPACE_ENTITY_PROPERTIES,
        "source_id": str(prop.get("source_chunk_id") or prop.get("property_id") or ""),
        "source_doc_id": source_doc_id,
        "source_scope_id": prop.get("source_scope_id"),
        "chunk_id": prop.get("source_chunk_id"),
        "entity_name": prop.get("entity_name"),
        "source_local_entity_id": prop.get("source_local_entity_id"),
        "cluster_id": prop.get("cluster_id"),
        "property_id": prop.get("property_id"),
        "content": content,
        "metadata": {
            "property_name": prop.get("property_name"),
            "property_value": prop.get("property_value"),
            "confidence": prop.get("confidence"),
            "scope": "source_local_property",
        },
    }


def _normalize_compact(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip().lower())


def _reference_vector_document_payload(row: sqlite3.Row) -> dict[str, Any]:
    payload = dict(row)
    payload["metadata"] = _json_loads(payload.pop("metadata_json", "{}"), default={})
    return payload


def _source_local_entity_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> dict[str, Any] | None:
    source_local_entity_id = str(hit.get("source_local_entity_id") or "")
    if not source_local_entity_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM reference_source_local_entities WHERE source_local_entity_id = ?",
            (source_local_entity_id,),
        ).fetchone()
        if not row:
            return None
        payload = _source_local_entity_payload(row)
    if source_doc_ids and not source_doc_ids.intersection(set(payload.get("source_doc_ids") or [])):
        return None
    if source_scope_ids and str(payload.get("source_scope_id") or "") not in source_scope_ids:
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _source_local_relationship_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> dict[str, Any] | None:
    source_local_relation_id = str(hit.get("source_local_relation_id") or "")
    if not source_local_relation_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM reference_source_local_relationships WHERE source_local_relation_id = ?",
            (source_local_relation_id,),
        ).fetchone()
        if not row:
            return None
        payload = _source_local_relationship_payload(row)
    if source_doc_ids and not source_doc_ids.intersection(set(payload.get("source_doc_ids") or [])):
        return None
    if source_scope_ids and str(payload.get("source_scope_id") or "") not in source_scope_ids:
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _atomic_fact_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> dict[str, Any] | None:
    fact_id = str(hit.get("fact_id") or "")
    if not fact_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute("SELECT * FROM reference_atomic_facts WHERE fact_id = ?", (fact_id,)).fetchone()
        if not row:
            return None
        payload = _atomic_fact_payload(row)
    if source_doc_ids and payload.get("source_doc_id") not in source_doc_ids:
        return None
    if source_scope_ids and str(payload.get("source_scope_id") or "") not in source_scope_ids:
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _entity_property_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> dict[str, Any] | None:
    property_id = str(hit.get("property_id") or "")
    if not property_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute("SELECT * FROM reference_entity_properties WHERE property_id = ?", (property_id,)).fetchone()
        if not row:
            return None
        payload = _entity_property_payload(row)
    if source_doc_ids and payload.get("source_doc_id") not in source_doc_ids:
        return None
    if source_scope_ids and str(payload.get("source_scope_id") or "") not in source_scope_ids:
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _entity_cluster_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_scope_ids: set[str],
    source_doc_ids: set[str],
) -> dict[str, Any] | None:
    cluster_id = str(hit.get("cluster_id") or "")
    if not cluster_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute("SELECT * FROM reference_entity_clusters WHERE cluster_id = ?", (cluster_id,)).fetchone()
        if not row:
            return None
        payload = _entity_cluster_payload(row)
    if source_scope_ids and not source_scope_ids.intersection(set(payload.get("source_scope_ids") or [])):
        return None
    if source_doc_ids and not source_doc_ids.intersection(set(payload.get("source_doc_ids") or [])):
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _pseudo_relationship_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_scope_ids: set[str],
    source_doc_ids: set[str],
) -> dict[str, Any] | None:
    pseudo_relation_id = str(hit.get("pseudo_relation_id") or "")
    if not pseudo_relation_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM reference_pseudo_relationships WHERE pseudo_relation_id = ?",
            (pseudo_relation_id,),
        ).fetchone()
        if not row:
            return None
        payload = _pseudo_relationship_payload(row)
    if source_scope_ids and not source_scope_ids.intersection(set(payload.get("source_scope_ids") or [])):
        return None
    if source_doc_ids and not source_doc_ids.intersection(set(payload.get("source_doc_ids") or [])):
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _chunk_payload_from_vector_hit(
    db_path: str | Path,
    hit: dict[str, Any],
    *,
    source_doc_ids: set[str] | None = None,
    source_scope_ids: set[str] | None = None,
) -> dict[str, Any] | None:
    chunk_id = str(hit.get("chunk_id") or hit.get("source_id") or "")
    if not chunk_id:
        return None
    with _connect(db_path) as conn:
        row = conn.execute("SELECT * FROM reference_text_chunks WHERE chunk_id = ?", (chunk_id,)).fetchone()
        if not row:
            return None
        payload = dict(row)
        payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    if source_doc_ids and str(payload.get("full_doc_id") or "") not in source_doc_ids:
        return None
    if source_scope_ids and str(payload.get("source_scope_id") or "") not in source_scope_ids:
        return None
    payload["score"] = hit.get("score")
    payload["vector_id"] = hit.get("vector_id")
    return payload


def _source_local_entity_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["facts"] = _json_loads(payload.pop("facts_json", "[]"), default=[])
    payload["attributes"] = _json_loads(payload.pop("attributes_json", "[]"), default=[])
    payload["timeline_claims"] = _json_loads(payload.pop("timeline_claims_json", "[]"), default=[])
    payload["source_doc_ids"] = _json_loads(payload.pop("source_doc_ids_json", "[]"), default=[])
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    payload["source_chunk_ids"] = _split_graph_field(payload.get("source_id"))
    return payload


def _source_local_relationship_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["facts"] = _json_loads(payload.pop("facts_json", "[]"), default=[])
    payload["source_doc_ids"] = _json_loads(payload.pop("source_doc_ids_json", "[]"), default=[])
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    payload["source_chunk_ids"] = _split_graph_field(payload.get("source_id"))
    return payload


def _entity_cluster_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["aliases"] = _json_loads(payload.pop("aliases_json", "[]"), default=[])
    payload["member_source_local_entity_ids"] = _json_loads(payload.pop("member_source_local_entity_ids_json", "[]"), default=[])
    payload["source_scope_ids"] = _json_loads(payload.pop("source_scope_ids_json", "[]"), default=[])
    payload["source_doc_ids"] = _json_loads(payload.pop("source_doc_ids_json", "[]"), default=[])
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    if not payload.get("source_separation_note") and payload.get("ambiguity_notes"):
        payload["source_separation_note"] = payload.get("ambiguity_notes")
    payload.pop("ambiguity_notes", None)
    return payload


def _pseudo_relationship_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["member_source_local_relation_ids"] = _json_loads(payload.pop("member_source_local_relation_ids_json", "[]"), default=[])
    payload["source_scope_ids"] = _json_loads(payload.pop("source_scope_ids_json", "[]"), default=[])
    payload["source_doc_ids"] = _json_loads(payload.pop("source_doc_ids_json", "[]"), default=[])
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    return payload


def _atomic_fact_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    return payload


def _entity_property_payload(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    payload = dict(row)
    payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
    return payload


def _matched_entity_names(entities: tuple[dict[str, Any], ...]) -> set[str]:
    names: set[str] = set()
    for entity in entities:
        for value in (entity.get("entity_id"), entity.get("canonical_name")):
            if str(value or "").strip():
                names.add(str(value).strip())
        for alias in entity.get("aliases") or []:
            if str(alias or "").strip():
                names.add(str(alias).strip())
    return names


def _scene_order_from_id(value: Any) -> int:
    text = str(value or "")
    for part in text.split("_"):
        if part.isdigit():
            return int(part)
    return 0


def _max_scene_order_before(value: Any) -> int:
    order = _scene_order_from_id(value)
    return max(order - 1, 0) if order else 0


def _tokens(text: Any) -> list[str]:
    lowered = str(text or "").lower()
    tokens = re.findall(r"[a-z0-9_]+", lowered)
    cjk_chars = [char for char in lowered if "\u4e00" <= char <= "\u9fff"]
    cjk_bigrams = [
        lowered[index : index + 2]
        for index in range(max(len(lowered) - 1, 0))
        if any("\u4e00" <= char <= "\u9fff" for char in lowered[index : index + 2])
    ]
    return tokens + cjk_chars + cjk_bigrams


def _reference_chroma_metadata(record: dict[str, Any]) -> dict[str, Any]:
    source_doc_ids = _vector_document_source_doc_ids(record)
    source_scope_ids = _vector_document_source_scope_ids(record)
    metadata = {
        "item_id": record.get("item_id"),
        "vector_namespace": record.get("vector_namespace") or "",
        "vector_id": record.get("vector_id") or record.get("doc_id") or "",
        "source_doc_id": record.get("source_doc_id") or "",
        "source_scope_id": record.get("source_scope_id") or "",
        "source_doc_ids_json": json.dumps(source_doc_ids, ensure_ascii=False),
        "source_scope_ids_json": json.dumps(source_scope_ids, ensure_ascii=False),
        "source_doc_count": len(source_doc_ids),
        "source_scope_count": len(source_scope_ids),
        "chunk_id": record.get("chunk_id") or "",
        "entity_name": record.get("entity_name") or "",
        "source_local_entity_id": record.get("source_local_entity_id") or "",
        "cluster_id": record.get("cluster_id") or "",
        "relation_id": record.get("relation_id") or "",
        "source_local_relation_id": record.get("source_local_relation_id") or "",
        "pseudo_relation_id": record.get("pseudo_relation_id") or "",
        "fact_id": record.get("fact_id") or "",
        "property_id": record.get("property_id") or "",
        "item_type": record.get("item_type"),
        "knowledge_scope": record.get("knowledge_scope"),
        "available_from_order": int(record.get("available_from_order") or 0),
    }
    return {key: value for key, value in metadata.items() if value is not None}


def _has_source_local_reference_assets(db_path: str | Path) -> bool:
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        row = conn.execute("SELECT COUNT(*) FROM reference_source_local_entities").fetchone()
        return bool(row and int(row[0]) > 0)


def _resolve_reference_source_doc_ids(
    db_path: str | Path,
    *,
    source_doc_ids: tuple[str, ...] | set[str] | list[str] = (),
    source_paths: tuple[str, ...] | set[str] | list[str] = (),
) -> set[str]:
    explicit = {str(item).strip() for item in source_doc_ids if str(item or "").strip()}
    paths = {str(item).strip() for item in source_paths if str(item or "").strip()}
    if not paths:
        return explicit
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        rows = conn.execute("SELECT doc_id, source_path, relative_path, file_name FROM reference_full_docs").fetchall()
    for row in rows:
        values = {str(row["source_path"] or ""), str(row["relative_path"] or ""), str(row["file_name"] or "")}
        if values.intersection(paths):
            explicit.add(str(row["doc_id"]))
    return explicit


def _resolve_reference_source_scope_ids(
    db_path: str | Path,
    *,
    source_scope_ids: tuple[str, ...] | set[str] | list[str] = (),
    source_doc_ids: set[str] | None = None,
    source_paths: tuple[str, ...] | set[str] | list[str] = (),
) -> set[str]:
    explicit = {str(item).strip() for item in source_scope_ids if str(item or "").strip()}
    doc_ids = {str(item).strip() for item in (source_doc_ids or set()) if str(item or "").strip()}
    paths = {str(item).strip() for item in source_paths if str(item or "").strip()}
    if not doc_ids and not paths:
        return explicit
    with _connect(db_path) as conn:
        _apply_reference_schema(conn)
        rows = conn.execute("SELECT doc_id, source_scope_id, source_path, relative_path, file_name FROM reference_full_docs").fetchall()
    for row in rows:
        doc_id = str(row["doc_id"] or "")
        scope_id = str(row["source_scope_id"] or _default_source_scope_id(doc_id))
        values = {str(row["source_path"] or ""), str(row["relative_path"] or ""), str(row["file_name"] or "")}
        if doc_id in doc_ids or values.intersection(paths):
            explicit.add(scope_id)
    return explicit


def _reference_source_filter_summary(
    query: ReferenceKnowledgeQuery,
    *,
    resolved_source_doc_ids: set[str] | None = None,
    resolved_source_scope_ids: set[str] | None = None,
) -> dict[str, Any]:
    source_doc_ids = sorted(resolved_source_doc_ids if resolved_source_doc_ids is not None else query.source_doc_ids)
    source_scope_ids = sorted(resolved_source_scope_ids if resolved_source_scope_ids is not None else query.source_scope_ids)
    mode = "all"
    if query.source_scope_ids or source_scope_ids:
        mode = "source_scopes"
    if query.source_doc_ids or query.source_paths or source_doc_ids:
        mode = "files"
    return {
        "mode": mode,
        "source_scope_ids": source_scope_ids,
        "source_doc_ids": source_doc_ids,
        "source_paths": list(query.source_paths),
    }


def _supporting_reference_chunks(
    db_path: str | Path,
    *,
    direct_chunks: list[dict[str, Any]],
    entities: list[dict[str, Any]],
    relationships: list[dict[str, Any]],
    atomic_facts: list[dict[str, Any]],
    entity_properties: list[dict[str, Any]],
    top_k: int,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> list[dict[str, Any]]:
    chunk_scores: dict[str, float] = {}
    for chunk in direct_chunks:
        chunk_id = str(chunk.get("chunk_id") or "")
        if chunk_id:
            chunk_scores[chunk_id] = max(chunk_scores.get(chunk_id, 0.0), float(chunk.get("score") or 0.0) + 1.0)
    for entity in entities:
        for chunk_id in entity.get("source_chunk_ids") or _split_graph_field(entity.get("source_id")):
            chunk_scores[str(chunk_id)] = chunk_scores.get(str(chunk_id), 0.0) + 0.8 + float(entity.get("score") or 0.0) * 0.1
    for relation in relationships:
        for chunk_id in relation.get("source_chunk_ids") or _split_graph_field(relation.get("source_id")):
            chunk_scores[str(chunk_id)] = chunk_scores.get(str(chunk_id), 0.0) + 0.7 + float(relation.get("score") or 0.0) * 0.1
    for fact in atomic_facts:
        chunk_id = str(fact.get("source_chunk_id") or "")
        if chunk_id:
            chunk_scores[chunk_id] = chunk_scores.get(chunk_id, 0.0) + 0.6 + float(fact.get("score") or 0.0) * 0.1
    for prop in entity_properties:
        chunk_id = str(prop.get("source_chunk_id") or "")
        if chunk_id:
            chunk_scores[chunk_id] = chunk_scores.get(chunk_id, 0.0) + 0.4 + float(prop.get("score") or 0.0) * 0.1
    if not chunk_scores:
        return direct_chunks[: max(top_k, 0)]
    placeholders = ", ".join("?" for _ in chunk_scores)
    with _connect(db_path) as conn:
        rows = [
            dict(row)
            for row in conn.execute(
                f"SELECT * FROM reference_text_chunks WHERE chunk_id IN ({placeholders})",
                list(chunk_scores),
            )
        ]
    chunks = []
    for row in rows:
        if source_doc_ids and str(row.get("full_doc_id") or "") not in source_doc_ids:
            continue
        if source_scope_ids and str(row.get("source_scope_id") or "") not in source_scope_ids:
            continue
        payload = dict(row)
        payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
        payload["score"] = round(chunk_scores.get(str(payload.get("chunk_id") or ""), 0.0), 4)
        chunks.append(payload)
    merged = {str(chunk.get("chunk_id") or ""): chunk for chunk in direct_chunks if str(chunk.get("chunk_id") or "")}
    for chunk in chunks:
        key = str(chunk.get("chunk_id") or "")
        if not key:
            continue
        current = merged.get(key)
        if current is None or float(chunk.get("score") or 0.0) > float(current.get("score") or 0.0):
            merged[key] = chunk
    return sorted(
        merged.values(),
        key=lambda item: (float(item.get("score") or 0.0), -int(item.get("chunk_order_index") or 0)),
        reverse=True,
    )[: max(top_k, 0)]


def _reference_evidence_chunk_ids(
    *,
    direct_chunks: list[dict[str, Any]],
    entities: list[dict[str, Any]],
    relationships: list[dict[str, Any]],
    atomic_facts: list[dict[str, Any]],
    entity_properties: list[dict[str, Any]],
) -> set[str]:
    chunk_ids = {str(chunk.get("chunk_id") or "") for chunk in direct_chunks}
    for entity in entities:
        chunk_ids.update(_entity_evidence_chunk_ids(entity))
    for relation in relationships:
        chunk_ids.update(_relationship_evidence_chunk_ids(relation))
    for fact in atomic_facts:
        chunk_ids.update(_fact_evidence_chunk_ids(fact))
    for prop in entity_properties:
        chunk_ids.update(_property_evidence_chunk_ids(prop))
    return {chunk_id for chunk_id in chunk_ids if chunk_id}


def _reference_chunk_index_by_ids(
    db_path: str | Path,
    chunk_ids: set[str],
    *,
    source_doc_ids: set[str],
    source_scope_ids: set[str],
) -> dict[str, dict[str, Any]]:
    if not chunk_ids:
        return {}
    placeholders = ", ".join("?" for _ in chunk_ids)
    with _connect(db_path) as conn:
        rows = [
            dict(row)
            for row in conn.execute(
                f"SELECT * FROM reference_text_chunks WHERE chunk_id IN ({placeholders})",
                sorted(chunk_ids),
            )
        ]
    index = {}
    for row in rows:
        if source_doc_ids and str(row.get("full_doc_id") or "") not in source_doc_ids:
            continue
        if source_scope_ids and str(row.get("source_scope_id") or "") not in source_scope_ids:
            continue
        payload = dict(row)
        payload["raw"] = _json_loads(payload.pop("raw_json", "{}"), default={})
        index[str(payload.get("chunk_id") or "")] = payload
    return index


def _attach_reference_evidence_chunks(
    items: list[dict[str, Any]],
    chunk_index: dict[str, dict[str, Any]],
    *,
    chunk_id_getter: Any,
) -> None:
    for item in items:
        item["evidence_chunks"] = [
            _reference_evidence_chunk_payload(chunk, statement=_evidence_statement_for_item(item))
            for chunk_id in chunk_id_getter(item)
            if (chunk := chunk_index.get(str(chunk_id or "")))
        ]


def _entity_evidence_chunk_ids(entity: dict[str, Any]) -> list[str]:
    return _dedupe_ordered(entity.get("source_chunk_ids") or _split_graph_field(entity.get("source_id")))


def _relationship_evidence_chunk_ids(relation: dict[str, Any]) -> list[str]:
    fact_chunk_ids = [
        str(fact.get("source_chunk_id") or "")
        for fact in relation.get("facts") or []
        if str(fact.get("source_chunk_id") or "").strip()
    ]
    return _dedupe_ordered(fact_chunk_ids + list(relation.get("source_chunk_ids") or _split_graph_field(relation.get("source_id"))))


def _fact_evidence_chunk_ids(fact: dict[str, Any]) -> list[str]:
    return [str(fact.get("source_chunk_id") or "")] if str(fact.get("source_chunk_id") or "").strip() else []


def _property_evidence_chunk_ids(prop: dict[str, Any]) -> list[str]:
    return [str(prop.get("source_chunk_id") or "")] if str(prop.get("source_chunk_id") or "").strip() else []


def _evidence_statement_for_item(item: dict[str, Any]) -> str:
    return str(item.get("statement") or item.get("description") or item.get("property_value") or "").strip()


def _reference_evidence_chunk_payload(chunk: dict[str, Any], *, statement: str = "") -> dict[str, Any]:
    content = str(chunk.get("content") or "")
    return {
        "chunk_id": chunk.get("chunk_id"),
        "source_doc_id": chunk.get("full_doc_id") or chunk.get("doc_id"),
        "source_scope_id": chunk.get("source_scope_id"),
        "source_path": chunk.get("source_path"),
        "relative_path": chunk.get("relative_path"),
        "title": chunk.get("title"),
        "heading": chunk.get("heading"),
        "content": content,
    }


def _relationship_facts_for_query(
    *,
    relationships: list[dict[str, Any]],
    atomic_facts: list[dict[str, Any]],
    evidence_chunk_index: dict[str, dict[str, Any]],
    top_k: int,
) -> list[dict[str, Any]]:
    relation_ids = {str(relation.get("source_local_relation_id") or "") for relation in relationships}
    facts = []
    for relation in relationships:
        for fact in relation.get("facts") or []:
            if str(fact.get("source_local_relation_id") or ""):
                enriched = dict(fact)
                enriched.setdefault("src_id", relation.get("src_id"))
                enriched.setdefault("tgt_id", relation.get("tgt_id"))
                enriched.setdefault("keywords", relation.get("keywords"))
                enriched.setdefault("source_doc_ids", relation.get("source_doc_ids") or [])
                facts.append(enriched)
    for fact in atomic_facts:
        if str(fact.get("source_local_relation_id") or ""):
            facts.append(dict(fact))
    deduped = []
    seen = set()
    for fact in facts:
        key = str(fact.get("fact_id") or "")
        if not key or key in seen:
            continue
        if relation_ids and str(fact.get("source_local_relation_id") or "") not in relation_ids:
            continue
        seen.add(key)
        fact["evidence_chunks"] = [
            _reference_evidence_chunk_payload(chunk, statement=str(fact.get("statement") or ""))
            for chunk_id in _fact_evidence_chunk_ids(fact)
            if (chunk := evidence_chunk_index.get(str(chunk_id or "")))
        ]
        deduped.append(fact)
    return sorted(
        deduped,
        key=lambda item: (float(item.get("score") or 0.0), str(item.get("source_chunk_id") or ""), str(item.get("fact_id") or "")),
        reverse=True,
    )[: max(top_k, 0)]


def _empty_reference_evidence_board() -> dict[str, Any]:
    return {
        "source_filter": {"mode": "all", "source_scope_ids": [], "source_doc_ids": [], "source_paths": []},
        "matched_clusters": [],
        "source_local_entities": [],
        "source_local_relationships": [],
        "relationship_facts": [],
        "atomic_facts": [],
        "entity_properties": [],
        "supporting_chunks": [],
        "pseudo_relationships": [],
        "graph_insights": [],
        "source_separation_notices": [],
    }


def _reference_evidence_board(
    *,
    source_filter: dict[str, Any],
    matched_clusters: list[dict[str, Any]],
    entities: list[dict[str, Any]],
    relationships: list[dict[str, Any]],
    relationship_facts: list[dict[str, Any]],
    atomic_facts: list[dict[str, Any]],
    entity_properties: list[dict[str, Any]],
    chunks: list[dict[str, Any]],
    pseudo_relationships: list[dict[str, Any]],
) -> dict[str, Any]:
    source_separation_notices = []
    for cluster in matched_clusters:
        note = str(cluster.get("source_separation_note") or "").strip()
        if note:
            source_separation_notices.append(
                {
                    "cluster_id": cluster.get("cluster_id"),
                    "canonical_display_name": cluster.get("canonical_display_name"),
                    "notice": note,
                }
            )
    graph_insights = []
    for cluster in sorted(matched_clusters, key=lambda item: int(item.get("global_pseudo_degree") or 0), reverse=True):
        graph_insights.append(
            {
                "type": "entity_cluster_degree",
                "cluster_id": cluster.get("cluster_id"),
                "canonical_display_name": cluster.get("canonical_display_name"),
                "global_pseudo_degree": cluster.get("global_pseudo_degree"),
                "evidence_degree": cluster.get("evidence_degree"),
                "source_coverage": cluster.get("source_coverage"),
                "degree_scope": "source_filtered" if source_filter.get("mode") != "all" else "global",
            }
        )
    return {
        "source_filter": source_filter,
        "matched_clusters": matched_clusters,
        "source_local_entities": entities,
        "source_local_relationships": relationships,
        "relationship_facts": relationship_facts,
        "atomic_facts": atomic_facts,
        "entity_properties": entity_properties,
        "supporting_chunks": chunks,
        "pseudo_relationships": pseudo_relationships,
        "graph_insights": graph_insights,
        "source_separation_notices": source_separation_notices,
    }


def _lightrag_entity_context_item_type(entity_type: Any) -> str:
    value = normalize_entity_type(str(entity_type or "")).strip().lower()
    if value in {"character", "person"}:
        return "character_profile"
    if value in {"location", "place"}:
        return "location_doc"
    if value in {"organization", "org"}:
        return "organization_fact"
    if value in {"event", "timeline", "time"}:
        return "timeline_doc"
    return "world_bible"


def _source_doc_ids_for_context_chunks(chunk_ids: list[str], chunk_by_id: dict[str, dict[str, Any]]) -> list[str]:
    values = []
    for chunk_id in chunk_ids:
        row = chunk_by_id.get(str(chunk_id or ""))
        if row and row.get("full_doc_id"):
            values.append(str(row.get("full_doc_id")))
    if values:
        return _dedupe_ordered(values)
    return _source_doc_ids_from_chunk_ids(chunk_ids)


def _context_source_docs(source_doc_ids: list[str], doc_by_id: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    docs = []
    for doc_id in source_doc_ids:
        row = doc_by_id.get(str(doc_id or ""))
        if not row:
            continue
        docs.append(
            {
                "doc_id": row.get("doc_id"),
                "source_path": row.get("source_path"),
                "relative_path": row.get("relative_path"),
                "file_name": row.get("file_name"),
                "title": row.get("title"),
            }
        )
    return docs


def _packet_lightrag_entity_item(entity: dict[str, Any]) -> dict[str, Any]:
    evidence_chunks = entity.get("evidence_chunks") or []
    return {
        "item_id": f"external_entity:{entity.get('source_local_entity_id') or entity.get('entity_name')}",
        "doc_id": ",".join(entity.get("source_doc_ids") or []),
        "chunk_id": entity.get("source_id"),
        "item_type": "character_profile" if entity.get("entity_type") == "character" else "world_bible",
        "subject": entity.get("entity_name"),
        "statement": entity.get("description"),
        "evidence": entity.get("description"),
        "authority": 0.7,
        "knowledge_scope": "author_only",
        "known_to": [],
        "available_from": "unknown",
        "timeline_hint": "",
        "confidence": 0.75,
        "score": entity.get("score"),
        "source_role": "external_reference_entity",
        "entity_type": entity.get("entity_type"),
        "source_scope_id": entity.get("source_scope_id"),
        "source_local_entity_id": entity.get("source_local_entity_id"),
        "cluster_id": entity.get("cluster_id"),
        "source_doc_ids": entity.get("source_doc_ids") or [],
        "evidence_chunks": evidence_chunks,
    }


def _packet_lightrag_relationship_item(relation: dict[str, Any]) -> dict[str, Any]:
    src = str(relation.get("src_id") or "")
    tgt = str(relation.get("tgt_id") or "")
    evidence_chunks = relation.get("evidence_chunks") or []
    return {
        "item_id": f"external_relation:{relation.get('source_local_relation_id') or relation.get('relation_id')}",
        "doc_id": ",".join(relation.get("source_doc_ids") or []),
        "chunk_id": relation.get("source_id"),
        "item_type": "relationship_fact",
        "subject": f"{src} - {tgt}",
        "statement": relation.get("description"),
        "evidence": relation.get("description"),
        "authority": 0.7,
        "knowledge_scope": "author_only",
        "known_to": [],
        "available_from": "unknown",
        "timeline_hint": "",
        "confidence": 0.75,
        "score": relation.get("score"),
        "source_role": "external_reference_relation",
        "keywords": relation.get("keywords"),
        "source_scope_id": relation.get("source_scope_id"),
        "source_local_relation_id": relation.get("source_local_relation_id"),
        "pseudo_relation_id": relation.get("pseudo_relation_id"),
        "source_doc_ids": relation.get("source_doc_ids") or [],
        "facts": relation.get("facts") or [],
        "evidence_chunks": evidence_chunks,
    }


def _packet_lightrag_fact_item(fact: dict[str, Any]) -> dict[str, Any]:
    evidence_chunks = fact.get("evidence_chunks") or []
    return {
        "item_id": f"external_fact:{fact.get('fact_id')}",
        "doc_id": fact.get("source_doc_id"),
        "chunk_id": fact.get("source_chunk_id"),
        "item_type": "relationship_fact" if fact.get("source_local_relation_id") else "world_bible",
        "subject": fact.get("subject"),
        "statement": fact.get("statement"),
        "evidence": fact.get("statement"),
        "authority": 0.72,
        "knowledge_scope": "author_only",
        "known_to": [],
        "available_from": "unknown",
        "timeline_hint": fact.get("statement") if _looks_temporal_text(fact.get("statement")) else "",
        "confidence": fact.get("confidence"),
        "score": fact.get("score"),
        "source_role": "external_reference_fact",
        "source_scope_id": fact.get("source_scope_id"),
        "source_doc_ids": [fact.get("source_doc_id")] if fact.get("source_doc_id") else [],
        "source_local_entity_id": fact.get("source_local_entity_id"),
        "source_local_relation_id": fact.get("source_local_relation_id"),
        "cluster_id": fact.get("cluster_id"),
        "object": fact.get("object"),
        "predicate": fact.get("predicate"),
        "fact_type": fact.get("fact_type"),
        "evidence_chunks": evidence_chunks,
    }


def _packet_lightrag_property_item(prop: dict[str, Any]) -> dict[str, Any]:
    evidence_chunks = prop.get("evidence_chunks") or []
    return {
        "item_id": f"external_property:{prop.get('property_id')}",
        "doc_id": prop.get("source_doc_id"),
        "chunk_id": prop.get("source_chunk_id"),
        "item_type": "character_profile" if prop.get("property_name") == "entity_type" and prop.get("property_value") == "character" else "world_bible",
        "subject": prop.get("entity_name"),
        "statement": prop.get("statement"),
        "evidence": prop.get("statement"),
        "authority": 0.7,
        "knowledge_scope": "author_only",
        "known_to": [],
        "available_from": "unknown",
        "timeline_hint": "",
        "confidence": prop.get("confidence"),
        "score": prop.get("score"),
        "source_role": "external_reference_property",
        "source_scope_id": prop.get("source_scope_id"),
        "source_doc_ids": [prop.get("source_doc_id")] if prop.get("source_doc_id") else [],
        "source_local_entity_id": prop.get("source_local_entity_id"),
        "cluster_id": prop.get("cluster_id"),
        "property_name": prop.get("property_name"),
        "property_value": prop.get("property_value"),
        "evidence_chunks": evidence_chunks,
    }


def _packet_lightrag_chunk_item(chunk: dict[str, Any]) -> dict[str, Any]:
    title = str(chunk.get("heading") or chunk.get("title") or chunk.get("full_doc_id") or "external chunk")
    content = str(chunk.get("content") or "")
    return {
        "item_id": f"external_chunk:{chunk.get('chunk_id')}",
        "doc_id": chunk.get("full_doc_id"),
        "chunk_id": chunk.get("chunk_id"),
        "item_type": "author_note",
        "subject": title,
        "statement": content,
        "evidence": content,
        "authority": 0.6,
        "knowledge_scope": "author_only",
        "known_to": [],
        "available_from": "unknown",
        "timeline_hint": "",
        "confidence": 0.7,
        "score": chunk.get("score"),
        "source_role": "external_reference_chunk",
        "source_doc_ids": [chunk.get("full_doc_id")] if chunk.get("full_doc_id") else [],
        "source_scope_id": chunk.get("source_scope_id"),
    }


def _entity_item_visible_to_character(item: dict[str, Any], matched_names: set[str]) -> bool:
    subject = str(item.get("subject") or "").lower()
    names = {name.lower() for name in matched_names}
    return bool(subject and subject in names)


def _relation_item_visible_to_character(item: dict[str, Any], matched_names: set[str]) -> bool:
    subject = str(item.get("subject") or "").lower()
    names = {name.lower() for name in matched_names}
    return any(name and name in subject for name in names)


def _fact_item_visible_to_character(item: dict[str, Any], matched_names: set[str]) -> bool:
    subject = str(item.get("subject") or "").lower()
    statement = str(item.get("statement") or "").lower()
    names = {name.lower() for name in matched_names}
    return any(name and (name in subject or name in statement) for name in names)


def _looks_temporal_text(value: Any) -> bool:
    text = str(value or "")
    return bool(re.search(r"\b(?:19|20)\d{2}\b|之前|之后|以前|以后|先于|晚于|时间线|timeline|before|after", text, flags=re.IGNORECASE))


def _dedupe_packet_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen = set()
    result = []
    for item in items:
        key = str(item.get("item_id") or "")
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def _merge_descriptions(name: str, descriptions: list[str]) -> str:
    values = [str(item).strip() for item in descriptions if str(item or "").strip()]
    if not values:
        return f"{name}"
    return GRAPH_FIELD_SEP.join(_dedupe_ordered(values))


def _merge_keywords(values: Any) -> str:
    keywords = []
    for value in values:
        for part in re.split(r"[,，;；]+", str(value or "")):
            text = part.strip()
            if text:
                keywords.append(text)
    return ", ".join(_dedupe_ordered(keywords))


def _dedupe_ordered(values: Any) -> list[str]:
    seen = set()
    result = []
    for value in values:
        text = str(value or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def _dedupe_relation_pairs(values: list[list[str]]) -> list[list[str]]:
    seen = set()
    result = []
    for value in values:
        if len(value) < 2:
            continue
        src, tgt = _sorted_relation_pair(value[0], value[1])
        if not src or not tgt:
            continue
        key = (src, tgt)
        if key in seen:
            continue
        seen.add(key)
        result.append([src, tgt])
    return result


def _sorted_relation_pair(src: Any, tgt: Any) -> tuple[str, str]:
    source = str(src or "").strip()
    target = str(tgt or "").strip()
    if source <= target:
        return source, target
    return target, source


def _reference_relation_id(src: Any, tgt: Any) -> str:
    source, target = _sorted_relation_pair(src, tgt)
    return _stable_reference_id("refrel", source, target)


def _default_source_scope_id(doc_id: Any) -> str:
    value = str(doc_id or "").strip()
    return f"scope_{_hash_id(value)}" if value else "scope_unknown"


def _canonical_reference_entity_name(value: Any) -> str:
    text = str(value or "").strip()
    return re.sub(r"\s+", " ", text)


def _reference_entity_cluster_id(entity_name: Any, entity_type: Any = None) -> str:
    type_key = normalize_entity_type(entity_type or "concept")
    return _stable_reference_id("refcluster", type_key, _normalize_compact(_canonical_reference_entity_name(entity_name)))


def _reference_source_local_entity_id(source_scope_id: Any, entity_name: Any) -> str:
    return _stable_reference_id("refsle", source_scope_id, _normalize_compact(_canonical_reference_entity_name(entity_name)))


def _reference_source_local_relation_id(source_scope_id: Any, relation_id: Any) -> str:
    return _stable_reference_id("refslr", source_scope_id, relation_id)


def _reference_pseudo_relation_id(src_cluster_id: Any, tgt_cluster_id: Any) -> str:
    source, target = _sorted_relation_pair(src_cluster_id, tgt_cluster_id)
    return _stable_reference_id("refpseudorel", source, target)


def _entity_fact_from_description(
    *,
    source_local_entity_id: str,
    source_scope_id: str,
    cluster_id: str,
    entity_name: str,
    description: str,
    source_chunk_id: str,
    source_doc_id: str,
) -> dict[str, Any]:
    statement = str(description or "").strip()
    return {
        "fact_id": _stable_reference_id("reffact", source_local_entity_id, source_chunk_id, statement),
        "source_scope_id": source_scope_id,
        "source_local_entity_id": source_local_entity_id,
        "source_local_relation_id": "",
        "cluster_id": cluster_id,
        "subject": entity_name,
        "predicate": "described_as",
        "object": "",
        "statement": statement,
        "fact_type": "entity_description_claim",
        "confidence": 0.7,
        "source_chunk_id": source_chunk_id,
        "source_doc_id": source_doc_id,
    }


def _relationship_fact_from_description(
    *,
    source_local_relation_id: str,
    source_scope_id: str,
    src: str,
    tgt: str,
    description: str,
    source_chunk_id: str,
    source_doc_id: str,
) -> dict[str, Any]:
    statement = str(description or "").strip()
    return {
        "fact_id": _stable_reference_id("reffact", source_local_relation_id, source_chunk_id, statement),
        "source_scope_id": source_scope_id,
        "source_local_entity_id": "",
        "source_local_relation_id": source_local_relation_id,
        "cluster_id": "",
        "subject": src,
        "predicate": "related_to",
        "object": tgt,
        "statement": statement,
        "fact_type": "relationship_description_claim",
        "confidence": 0.7,
        "source_chunk_id": source_chunk_id,
        "source_doc_id": source_doc_id,
    }


def _reference_cluster_summary(display_name: str, source_scope_ids: list[str]) -> str:
    if len(source_scope_ids) <= 1:
        return f"{display_name} appears in one external source scope."
    return (
        f"{display_name} appears in {len(source_scope_ids)} external source scopes. "
        "This is a pseudo cluster for retrieval and graph navigation; source-local claims remain separate."
    )


def _stable_reference_id(prefix: str, *parts: Any) -> str:
    return f"{prefix}_{_hash_id('|'.join(str(part) for part in parts))}"


def _hash_id(value: Any) -> str:
    return hashlib.sha1(str(value or "").encode("utf-8")).hexdigest()


def _split_graph_field(value: Any) -> list[str]:
    return [part for part in str(value or "").split(GRAPH_FIELD_SEP) if part]


def _source_doc_ids_from_chunk_ids(chunk_ids: list[str]) -> list[str]:
    values = []
    for chunk_id in chunk_ids:
        match = re.match(r"(.+)_chunk_\d+$", str(chunk_id or ""))
        if match:
            values.append(match.group(1))
    return _dedupe_ordered(values)


def _source_doc_ids_for_chunk_ids(db_path: str | Path, chunk_ids: list[str]) -> list[str]:
    if not chunk_ids:
        return []
    placeholders = ", ".join("?" for _ in chunk_ids)
    with _connect(db_path) as conn:
        rows = conn.execute(
            f"SELECT DISTINCT full_doc_id FROM reference_text_chunks WHERE chunk_id IN ({placeholders})",
            chunk_ids,
        ).fetchall()
    return _dedupe_ordered(str(row["full_doc_id"]) for row in rows)


def _batches(records: list[dict[str, Any]], batch_size: int) -> list[list[dict[str, Any]]]:
    return [records[index : index + batch_size] for index in range(0, len(records), batch_size)]


def _select_records(records: list[dict[str, Any]], *, start: int, limit: int | None) -> list[dict[str, Any]]:
    offset = max(start - 1, 0)
    return records[offset:] if limit is None else records[offset : offset + max(limit, 0)]


def _read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return payload


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    with Path(path).open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            payload = json.loads(line)
            if isinstance(payload, dict):
                records.append(payload)
            else:
                raise ValueError(f"JSONL rows must be objects: {path}")
    return records


def _load_reference_llm_response_cache(kg_dir: Path) -> list[dict[str, Any]]:
    trace_path = kg_dir / "trace.jsonl"
    if not trace_path.exists():
        return []
    records = []
    for trace in _read_jsonl(trace_path):
        chunk_id = str(trace.get("chunk_id") or "").strip()
        if not chunk_id:
            continue
        raw_path = _optional_trace_artifact_path(kg_dir, trace.get("raw_output_path"))
        parsed_path = _optional_trace_artifact_path(kg_dir, trace.get("parsed_path"))
        raw_text = raw_path.read_text(encoding="utf-8") if raw_path else ""
        parsed = _read_json(parsed_path) if parsed_path else {}
        records.append(
            {
                "cache_id": _stable_reference_id("ref_llm_cache", chunk_id, raw_text),
                "chunk_id": chunk_id,
                "doc_id": trace.get("doc_id"),
                "status": trace.get("status") or parsed.get("status"),
                "provider": parsed.get("provider"),
                "model": parsed.get("model"),
                "raw_text": raw_text,
                "parsed": parsed,
                "trace": trace,
                "created_at": datetime.now(timezone.utc).isoformat(),
            }
        )
    return records


def _optional_trace_artifact_path(kg_dir: Path, value: Any) -> Path | None:
    text = str(value or "").strip()
    if not text:
        return None
    path = Path(text)
    if not path.is_absolute():
        path = kg_dir / path
    if not path.exists() or not path.is_file():
        return None
    return path


def _write_jsonl(path: Path, records: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False) + "\n")


def _format_policy(policy: Any) -> str:
    if isinstance(policy, list):
        return "\n".join(f"- {item}" for item in policy)
    return str(policy or "")


def _safe_path_id(value: Any) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_.-]+", "_", str(value or "item"))
    return safe.strip("._") or "item"


def _json_loads(value: Any, *, default: Any) -> Any:
    try:
        return json.loads(str(value))
    except (TypeError, ValueError, json.JSONDecodeError):
        return default


def _connect(path: str | Path) -> sqlite3.Connection:
    conn = sqlite3.connect(str(path))
    conn.row_factory = sqlite3.Row
    return conn


def _import_chromadb():
    try:
        import chromadb
    except ImportError as exc:
        raise RuntimeError(
            "chromadb is not installed. Use the screenplay conda environment or install the optional vector DB dependency."
        ) from exc
    return chromadb
